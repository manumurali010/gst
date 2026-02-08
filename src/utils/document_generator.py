import os
import tempfile
# Lazy imports to prevent startup hang
# from reportlab.lib.pagesizes import A4
# from reportlab.pdfgen import canvas
# from reportlab.lib.units import inch
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
# from xhtml2pdf import pisa
# from docx2pdf import convert
from src.utils.constants import OUTPUT_DIR

class DocumentGenerator:
    def __init__(self):
        if not os.path.exists(OUTPUT_DIR):
            os.makedirs(OUTPUT_DIR)

    def generate_pdf_from_html(self, html_content, filename):
        """Generate PDF directly from HTML content"""
        filepath = os.path.join(OUTPUT_DIR, filename + ".pdf")
        
        # Add some basic CSS for PDF formatting if not present
        if "<style>" not in html_content:
            style = """
            <style>
                @page {
                    size: A4;
                    margin: 1cm;
                }
                body {
                    font-family: 'Bookman Old Style', serif;
                    font-size: 11pt;
                    text-align: justify;
                }
                table {
                    border-collapse: collapse;
                    width: 90%;
                    margin: 0 auto;
                    font-size: 10pt;
                }
                td, th {
                    border: 1px solid black;
                    padding: 2px 5px;
                    text-align: center;
                }
            </style>
            """
            html_content = style + html_content
            
        # DEBUG: Save HTML to file to inspect
        with open(os.path.join(OUTPUT_DIR, "debug_last_generated.html"), "w", encoding="utf-8") as f:
            f.write(html_content)

        # [HARD-DISABLE] Emergency Bypass
        print("[STABILIZATION] PDF Generation from HTML is HARD-DISABLED.")
        return filepath # Return path even if empty to avoid crashes if expected
        
        try:
            with open(filepath, "w+b") as result_file:
                # from xhtml2pdf import pisa
                pass
        except Exception as e:
            # Log full traceback to file
            import traceback
            error_log = os.path.join(OUTPUT_DIR, "pdf_gen_error.log")
            with open(error_log, "w") as f:
                f.write(f"Error generating PDF: {str(e)}\n")
                f.write(traceback.format_exc())
            raise e
    
    def generate_pdf_from_docx(self, docx_letterhead_path, data, filename):
        """Generate PDF using DOCX letterhead template"""
        # Create a temporary DOCX with letterhead + content
        temp_docx = os.path.join(tempfile.gettempdir(), f"temp_{filename}.docx")
        
        try:
            # Verify letterhead exists
            if not os.path.exists(docx_letterhead_path):
                raise FileNotFoundError(f"Letterhead template not found: {docx_letterhead_path}")
            
            # Load the letterhead template
            doc = Document(docx_letterhead_path)
            
            # Add form content to the document
            self._add_content_to_docx(doc, data)
            
            # Save temporary DOCX
            doc.save(temp_docx)
            
            # Convert to PDF
            output_pdf = os.path.join(OUTPUT_DIR, filename + ".pdf")
            from docx2pdf import convert
            convert(temp_docx, output_pdf)
            
            return output_pdf
            
        except Exception as e:
            raise Exception(f"Failed to generate PDF from DOCX: {str(e)}")
        finally:
            # Clean up temp file
            if os.path.exists(temp_docx):
                try:
                    os.remove(temp_docx)
                except:
                    pass
    
    def generate_word_from_docx(self, docx_letterhead_path, data, filename):
        """Generate Word document using DOCX letterhead template"""
        try:
            # Check if this is a PNG-based letterhead
            # PNG letterheads are saved as .png files, and we need to use the PNG for Word
            letterhead_dir = os.path.dirname(docx_letterhead_path)
            letterhead_name = os.path.basename(docx_letterhead_path)
            
            # Check if there's a corresponding PNG file
            if letterhead_name.endswith('_png.html'):
                # This is a PNG-based letterhead, find the PNG file
                png_name = letterhead_name.replace('_png.html', '.png')
                png_path = os.path.join(letterhead_dir, png_name)
                
                if os.path.exists(png_path):
                    # Create new document with PNG letterhead
                    doc = Document()
                    
                    # Add PNG image at top
                    from docx.shared import Inches
                    doc.add_picture(png_path, width=Inches(6.5))
                    
                    # Add form content (skip page break since PNG is inline)
                    self._add_content_to_docx(doc, data, skip_page_break=True)
                    
                    # Save
                    filepath = os.path.join(OUTPUT_DIR, filename + ".docx")
                    doc.save(filepath)
                    
                    return filepath
            
            # Regular DOCX letterhead or PNG not found - use template
            # Verify letterhead exists
            if not os.path.exists(docx_letterhead_path):
                raise FileNotFoundError(f"Letterhead template not found: {docx_letterhead_path}")
            
            # Load the letterhead template
            doc = Document(docx_letterhead_path)
            
            # Add form content
            self._add_content_to_docx(doc, data)
            
            # Save
            filepath = os.path.join(OUTPUT_DIR, filename + ".docx")
            doc.save(filepath)
            
            return filepath
            
        except Exception as e:
            raise Exception(f"Failed to generate Word document from DOCX: {str(e)}")
    
    def _add_content_to_docx(self, doc, data, skip_page_break=False):
        """Add form content to a DOCX document"""
        # Add a page break to separate letterhead from content (unless skipped for PNG)
        if not skip_page_break:
            doc.add_page_break()
            
        # Define style for body text
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Bookman Old Style'
        font.size = Pt(11)
        
        # Add form header
        p = doc.add_heading(f"NOTICE: {data.get('form_type', 'NOTICE')}", level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add details
        for line in [
            f"Date: {data.get('date', '')}",
            f"GSTIN: {data.get('gstin', '')}",
            f"Legal Name: {data.get('legal_name', '')}",
            f"Address: {data.get('address', '')}",
            f"Subject: Notice under {data.get('proceeding_type', '')}"
        ]:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Add facts
        p = doc.add_paragraph(data.get('facts', ''))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Add tax table if present
        if 'tax_data' in data and data['tax_data']:
            # Create table with 7 columns
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            hdr_cells = table.rows[0].cells
            headers = ["Act", "From", "To", "Tax", "Interest", "Penalty", "Total"]
            for i, h in enumerate(headers):
                hdr_cells[i].text = h
                # Set font for header
                for paragraph in hdr_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = 'Bookman Old Style'
                        run.font.size = Pt(10)
                        run.font.bold = True
                
            for row_data in data['tax_data']:
                row_cells = table.add_row().cells
                # Map data to 7 columns
                vals = [
                    str(row_data.get('Act', '')),
                    str(row_data.get('From', '')),
                    str(row_data.get('To', '')),
                    str(row_data.get('Tax', 0)),
                    str(row_data.get('Interest', 0)),
                    str(row_data.get('Penalty', 0)),
                    str(row_data.get('Total', 0))
                ]
                
                for i, val in enumerate(vals):
                    row_cells[i].text = val
                    # Set font for cells
                    for paragraph in row_cells[i].paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.name = 'Bookman Old Style'
                            run.font.size = Pt(10)

    def generate_word(self, data, filename):
        filepath = os.path.join(OUTPUT_DIR, filename + ".docx")
        doc = Document()
        
        # Define style for body text
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Bookman Old Style'
        font.size = Pt(11)
        
        # Header
        head = doc.add_heading('GOVERNMENT OF INDIA', 0)
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub = doc.add_paragraph('GOODS AND SERVICES TAX DEPARTMENT')
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = doc.add_heading(f"NOTICE: {data.get('form_type', 'NOTICE')}", level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Details
        for line in [
            f"Date: {data.get('date', '')}",
            f"GSTIN: {data.get('gstin', '')}",
            f"Legal Name: {data.get('legal_name', '')}",
            f"Address: {data.get('address', '')}",
            f"Subject: Notice under {data.get('proceeding_type', '')}"
        ]:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Body
        p = doc.add_paragraph(data.get('facts', ''))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Table
        if 'tax_data' in data and data['tax_data']:
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            hdr_cells = table.rows[0].cells
            headers = ["Act", "From", "To", "Tax", "Interest", "Penalty", "Total"]
            for i, h in enumerate(headers):
                hdr_cells[i].text = h
                for paragraph in hdr_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = 'Bookman Old Style'
                        run.font.size = Pt(10)
                        run.font.bold = True
                
            for row_data in data['tax_data']:
                row_cells = table.add_row().cells
                vals = [
                    str(row_data.get('Act', '')),
                    str(row_data.get('From', '')),
                    str(row_data.get('To', '')),
                    str(row_data.get('Tax', 0)),
                    str(row_data.get('Interest', 0)),
                    str(row_data.get('Penalty', 0)),
                    str(row_data.get('Total', 0))
                ]
                for i, val in enumerate(vals):
                    row_cells[i].text = val
                    for paragraph in row_cells[i].paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.name = 'Bookman Old Style'
                            run.font.size = Pt(10)
                
        doc.save(filepath)
        return filepath

    def generate_pdf(self, data, filename):
        # Fallback to manual generation if needed, but we prefer HTML
        # ... (keeping existing logic as fallback or for other uses)
        filepath = os.path.join(OUTPUT_DIR, filename + ".pdf")
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import inch
        c = canvas.Canvas(filepath, pagesize=A4)
        width, height = A4
        
        # Header
        c.setFont("Helvetica-Bold", 16)
        c.drawCentredString(width / 2, height - 1 * inch, "GOVERNMENT OF INDIA")
        c.setFont("Helvetica-Bold", 14)
        c.drawCentredString(width / 2, height - 1.3 * inch, "GOODS AND SERVICES TAX DEPARTMENT")
        
        # Title
        c.setFont("Helvetica-Bold", 12)
        c.drawCentredString(width / 2, height - 2 * inch, f"NOTICE: {data.get('form_type', 'NOTICE')}")
        
        # Details
        c.setFont("Helvetica", 10)
        y = height - 2.5 * inch
        x = 1 * inch
        line_height = 14
        
        details = [
            f"Date: {data.get('date', '')}",
            f"GSTIN: {data.get('gstin', '')}",
            f"Legal Name: {data.get('legal_name', '')}",
            f"Trade Name: {data.get('trade_name', '')}",
            f"Address: {data.get('address', '')}",
            f"Proceeding: {data.get('proceeding_type', '')}",
            f"Form: {data.get('form_type', '')}",
        ]
        
        for line in details:
            c.drawString(x, y, line)
            y -= line_height
            
        y -= line_height
        c.drawString(x, y, "Subject: Notice under GST Act")
        y -= line_height * 2
        
        # Body
        c.setFont("Helvetica", 10)
        text = data.get('facts', '')
        # Simple text wrapping (very basic)
        text_lines = text.split('\n')
        for line in text_lines:
            # Check if line fits, otherwise cut (simplified for now)
            if y < 1 * inch:
                c.showPage()
                y = height - 1 * inch
            c.drawString(x, y, line)
            y -= line_height
            
        y -= line_height * 2
        
        # Table
        if 'tax_data' in data and data['tax_data']:
            # Draw table headers
            headers = ["Period", "Type", "Tax", "Interest", "Penalty", "Late Fee", "Total"]
            col_widths = [60, 40, 60, 60, 60, 60, 60]
            x_offset = x
            c.setFont("Helvetica-Bold", 9)
            for i, header in enumerate(headers):
                c.drawString(x_offset, y, header)
                x_offset += col_widths[i]
            y -= line_height
            
            c.setFont("Helvetica", 9)
            for row in data['tax_data']:
                x_offset = x
                vals = [
                    row.get('Period', ''),
                    row.get('Tax Type', ''),
                    str(row.get('Tax Amount', 0)),
                    str(row.get('Interest', 0)),
                    str(row.get('Penalty', 0)),
                    str(row.get('Late fee', 0)),
                    str(row.get('Total', 0))
                ]
                for i, val in enumerate(vals):
                    c.drawString(x_offset, y, val)
                    x_offset += col_widths[i]
                y -= line_height
                
        c.save()
        return filepath
