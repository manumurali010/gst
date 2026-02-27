from PyQt6.QtWidgets import (QWidget, QVBoxLayout, QHBoxLayout, QLabel, QPushButton, 
                             QListWidget, QStackedWidget, QSplitter, QScrollArea, QTextEdit, QTextBrowser,
                             QMessageBox, QFrame, QCheckBox, QTableWidget, QTableWidgetItem, QHeaderView, QDateEdit, QComboBox, QLineEdit, QFileDialog, QDialog, QGridLayout, QSpacerItem, QSizePolicy, QGraphicsDropShadowEffect, QToolTip)
from PyQt6.QtCore import Qt, QDate, pyqtSignal, QRect
from PyQt6 import QtCore
from PyQt6.QtGui import QPixmap, QShortcut, QKeySequence, QIcon, QResizeEvent, QColor, QCursor
from src.database.db_manager import DatabaseManager
from PyQt6.QtWebEngineWidgets import QWebEngineView
from src.utils.preview_generator import PreviewGenerator
from src.ui.collapsible_box import CollapsibleBox
from src.ui.rich_text_editor import RichTextEditor
from src.ui.components.modern_card import ModernCard
from src.ui.issue_card import IssueCard
from src.ui.adjudication_setup_dialog import AdjudicationSetupDialog
from src.ui.components.side_nav_card import SideNavCard # Canonical import
from src.ui.components.finalization_panel import FinalizationPanel
from src.services.asmt10_generator import ASMT10Generator
from src.services.ph_intimation_generator import PHIntimationGenerator
from src.ui.styles import Theme, Styles
from src.utils.constants import WorkflowStage
from src.utils.number_utils import safe_int
import os
import json
import copy
from jinja2 import Template, Environment, FileSystemLoader
import datetime
import traceback
import base64
import re

# --- HELPER CLASSES FOR DRC-01A REFACTOR ---

class LegalReference:
    """Canonical representation of a legal provision for deduplication and sorting."""
    def __init__(self, raw_text):
        self.raw = raw_text # Preserve original text
        self.type = "Other" # Section, Rule, Notification
        self.act = "Unknown"  # e.g., CGST Act, SGST Act
        self.major = 0
        self.minor = ""
        self.canonical = ""
        self.parse()

    @property
    def canonical_id(self):
        """Stable, unique identifier for deduplication."""
        # Normalize minor for ID: remove spaces, lowercase
        clean_minor = re.sub(r'\s+', '', str(self.minor)).lower()
        return f"{self.type}:{self.act}:{self.major}:{clean_minor}"

    def parse(self):
        text = self.raw.strip()
        
        # 1. Act Detection (Simple heuristic)
        if "SGST" in text.upper(): self.act = "SGST Act"
        elif "CGST" in text.upper(): self.act = "CGST Act"
        elif "IGST" in text.upper(): self.act = "IGST Act"
        # If no act found, stay as "Unknown" as per refinement

        # 2. Canonical Normalization of prefix
        # Use word boundaries \b to prevent matching "Sec" inside "Section"
        if not text.lower().startswith("section "):
            text = re.sub(r'^(Sec\b|S\b|u/s|under\s+section)\.?\s*', 'Section ', text, flags=re.I)
        
        if not text.lower().startswith("rule "):
            text = re.sub(r'^(Rule\b|R\b)\.?\s*', 'Rule ', text, flags=re.I)
        
        # 3. Section Pattern: Section 7(1)(a)
        sec_match = re.search(r'Section\s+(\d+)(.*)', text, re.I)
        if sec_match:
            self.type = "Section"
            self.major = int(sec_match.group(1))
            self.minor = sec_match.group(2).strip()
            self.canonical = f"Section {self.major}{self.minor}"
            return

        # 4. Rule Pattern: Rule 117
        rule_match = re.search(r'Rule\s+(\d+)(.*)', text, re.I)
        if rule_match:
            self.type = "Rule"
            self.major = int(rule_match.group(1))
            self.minor = rule_match.group(2).strip()
            self.canonical = f"Rule {self.major}{self.minor}"
            return
            
        self.canonical = text

    def __eq__(self, other):
        if not isinstance(other, LegalReference): return False
        return self.canonical_id == other.canonical_id

    def __hash__(self):
        return hash(self.canonical_id)

    def __lt__(self, other):
        # Logical Sort Priority
        type_priority = {"Section": 0, "Rule": 1, "Other": 2}
        if self.type != other.type:
            return type_priority.get(self.type, 2) < type_priority.get(other.type, 2)
        if self.act != other.act:
            return str(self.act) < str(other.act)
        if self.major != other.major:
            return self.major < other.major
        return self.minor < other.minor

class ASMT10ReferenceDialog(QDialog):
    """
    Modal Dialog for viewing ASMT-10 as a reference while drafting.
    """
    def __init__(self, parent, html):
        super().__init__(parent)
        self.setWindowTitle("🔒 Finalised ASMT-10 (Read-Only Reference)")
        self.resize(1000, 800)
        
        layout = QVBoxLayout(self)
        layout.setContentsMargins(10, 10, 10, 10)
        
        # Scroll Area for preview
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        container = QWidget()
        container_layout = QVBoxLayout(container)
        container_layout.setSpacing(20)
        container_layout.setAlignment(Qt.AlignmentFlag.AlignHCenter)
        scroll.setWidget(container)
        layout.addWidget(scroll)
        
        # Close Button
        btn_close = QPushButton("Close Reference")
        btn_close.clicked.connect(self.accept)
        btn_close.setStyleSheet("""
            QPushButton {
                background-color: #f8f9fa;
                border: 1px solid #dcdde1;
                padding: 8px 16px;
                border-radius: 4px;
            }
            QPushButton:hover {
                background-color: #e9ecef;
            }
        """)
        layout.addWidget(btn_close)
        
        # Render HTML to images
        if html:
            images = PreviewGenerator.generate_preview_image(html, all_pages=True)
            if images:
                for img_bytes in images:
                    pixmap = PreviewGenerator.get_qpixmap_from_bytes(img_bytes)
                    if pixmap:
                        lbl = QLabel()
                        # Fixed width for modal
                        lbl.setPixmap(pixmap.scaledToWidth(900, Qt.TransformationMode.SmoothTransformation))
                        lbl.setStyleSheet("border: 1px solid #ccc; background-color: white;")
                        container_layout.addWidget(lbl)
            else:
                container_layout.addWidget(QLabel("Preview Generation Failed"))
        else:
            container_layout.addWidget(QLabel("No ASMT-10 Data Found"))

class ProceedingsWorkspace(QWidget):
    # [PHASE 3] Strict Transition Matrix
    # Defines the ONLY allowed allowed predecessors for a target stage.
    TRANSITION_MATRIX = {
        WorkflowStage.ASMT10_ISSUED: [WorkflowStage.ASMT10_DRAFT],
        WorkflowStage.DRC01A_ISSUED: [WorkflowStage.DRC01A_DRAFT],
        
        # Unified SCN Entry: Scrutiny (ASMT10) OR Direct (DRC01A) OR Skipped (Direct w/ Skip)
        WorkflowStage.SCN_DRAFT: [
            WorkflowStage.DRC01A_ISSUED, 
            WorkflowStage.ASMT10_ISSUED, 
            WorkflowStage.DRC01A_DRAFT # Protected by drc01a_skipped check
        ],
        
        WorkflowStage.SCN_ISSUED:    [WorkflowStage.SCN_DRAFT],
        WorkflowStage.PH_SCHEDULED:  [WorkflowStage.SCN_ISSUED, WorkflowStage.PH_SCHEDULED], # Allow re-scheduling
        
        # Structured PH Lifecycle
        WorkflowStage.PH_COMPLETED:  [WorkflowStage.PH_SCHEDULED],
        WorkflowStage.ORDER_ISSUED:  [WorkflowStage.PH_COMPLETED]
    }

    # Deterministic Status Mapping (Enum -> String)
    STAGE_TO_STATUS_MAP = {
        WorkflowStage.ASMT10_ISSUED: "ASMT-10 Issued", # Actually tracked in asmt10_status, but we map for consistency if needed
        WorkflowStage.DRC01A_ISSUED: "DRC-01A Issued",
        WorkflowStage.SCN_ISSUED:    "SCN Issued",
        WorkflowStage.PH_SCHEDULED:  "PH Intimated",
        WorkflowStage.PH_COMPLETED:  "PH Concluded",
        WorkflowStage.ORDER_ISSUED:  "Order Issued"
    }
    def __init__(self, navigate_callback, sidebar=None, proceeding_id=None):
        super().__init__()
        print("ProceedingsWorkspace: init start")
        self.navigate_callback = navigate_callback
        self.sidebar = sidebar # Store sidebar reference
        self.db = DatabaseManager()
        self.db.init_sqlite()
        print("ProceedingsWorkspace: DB initialized")
        
        self.proceeding_id = proceeding_id
        self.proceeding_data = {}
        self.is_hydrated = False
        
        self.asmt10_zoom_level = 1.0
        self.asmt10_show_letterhead = True
        
        self.active_scn_step = 0
        self.preview_initialized = False
        self.scn_workflow_phase = "METADATA" # Authority state: METADATA | DRAFTING


        # [NEW] PH Generator & State
        self.ph_generator = PHIntimationGenerator()
        self.ph_entries = [] # List of dicts
        self.ph_editing_index = -1
        
        # UI State Flags
        self.scn_issues_initialized = False # Track if we've loaded issues once to avoid overwrite
        
        print("ProceedingsWorkspace: calling init_ui")
        self.init_ui()
        print("ProceedingsWorkspace: init_ui done")
        
        
        if self.proceeding_id:
            self.load_proceeding(self.proceeding_id)

    def transition_to(self, target_stage: WorkflowStage):
        """
        Atomic, Strict State Transition Engine.
        Updates both 'workflow_stage' and legacy 'status'.
        """
        current_stage = self.get_current_stage()
        
        print(f"Workflow Transition Request: {current_stage.name} -> {target_stage.name}")
        
        # 1. Backward Transition Guard (Strict Forward Flow)
        # Exception: Re-scheduling PH (PH -> PH) is allowed
        is_rescheduling_ph = (current_stage == WorkflowStage.PH_SCHEDULED and target_stage == WorkflowStage.PH_SCHEDULED)
        
        if target_stage < current_stage and not is_rescheduling_ph:
             raise ValueError(f"Illegal Backward Transition: Cannot move from {current_stage.name} back to {target_stage.name}")

        # 2. Matrix Validation
        allowed_predecessors = self.TRANSITION_MATRIX.get(target_stage, [])
        if current_stage not in allowed_predecessors:
             raise ValueError(f"Invalid Transition: {current_stage.name} is not a valid predecessor for {target_stage.name}. Allowed: {[s.name for s in allowed_predecessors]}")

        # 3. Engine-Level Skip Enforcement (DRC01A_DRAFT -> SCN_DRAFT)
        if target_stage == WorkflowStage.SCN_DRAFT and current_stage == WorkflowStage.DRC01A_DRAFT:
             drc_skipped = self.proceeding_data.get('drc01a_skipped', False)
             if not drc_skipped:
                  raise ValueError("Illegal Skip: Cannot move to SCN_DRAFT from DRC01A_DRAFT unless 'drc01a_skipped' is explicitly set.")

        # 4. Atomic DB Update
        try:
             # Map status string
             status_str = self.STAGE_TO_STATUS_MAP.get(target_stage)
             
             update_payload = {
                 "workflow_stage": int(target_stage)
             }
             if status_str:
                  update_payload["status"] = status_str
             
             # Get current version for optimistic locking
             v_no = self.proceeding_data.get('version_no')
             
             # Execute
             success = self.db.update_proceeding(self.proceeding_id, update_payload, version_no=v_no)
             if not success:
                  raise Exception("Database update failed (Unknown reason).")
             
             print(f"Transition Success: {current_stage.name} -> {target_stage.name}")
             
             # 5. Post-Commit State & UI Refresh
             # Increment version locally since it was incremented in DB
             if v_no is not None:
                  self.proceeding_data['version_no'] = v_no + 1
             
             self.proceeding_data['workflow_stage'] = int(target_stage)
             if status_str:
                  self.proceeding_data['status'] = status_str
                  
             # Refresh UI elements
             self.update_summary_tab()
             self.check_existing_documents()
             self.evaluate_scn_workflow_phase()
             self.apply_context_layout(target_stage.get_context_key())
             
        except Exception as e:
             # If ConcurrencyError or other DB error occurs, state remains unchanged
             print(f"Transition Aborted: {e}")
             raise e

    def get_current_stage(self) -> WorkflowStage:
        """
        Robustly determine current workflow stage using the dedicated column.
        Default: ASMT10_DRAFT (Base Stage) if undefined.
        """
        if not self.proceeding_data:
             # Default to initial stage if no data loaded yet (e.g. during UI init)
             return WorkflowStage.ASMT10_DRAFT
            
        stage_val = self.proceeding_data.get('workflow_stage')
        if stage_val is not None:
            try:
                return WorkflowStage(int(stage_val))
            except ValueError:
                pass
                
        # Fallback for legacy cases (should be migrated, but be safe)
        status = self.proceeding_data.get('status', '')
        if "Order Issued" in status: return WorkflowStage.ORDER_ISSUED
        if "PH Intimated" in status: return WorkflowStage.PH_SCHEDULED
        if "SCN Issued" in status: return WorkflowStage.SCN_ISSUED
        if "DRC-01A Issued" in status: return WorkflowStage.DRC01A_ISSUED
        if "ASMT-10 Issued" in status: return WorkflowStage.ASMT10_ISSUED
        
        return WorkflowStage.ASMT10_DRAFT

    def init_ui(self):
        print("ProceedingsWorkspace: init_ui start")
        self.setObjectName("ProceedingsWorkspace")
        
        # Styles are handled globally in styles.py via main_window.py
        
        self.layout = QHBoxLayout(self)
        self.layout.setContentsMargins(0, 0, 0, 0)
        
        # 2. Main Workspace Splitter (Context-Driven)
        self.workspace_splitter = QSplitter(Qt.Orientation.Horizontal)
        self.workspace_splitter.setHandleWidth(1)
        self.workspace_splitter.setStyleSheet("QSplitter::handle { background-color: #e0e0e0; }")
        
        # Center Pane: Content (Stacked)
        self.content_stack = QStackedWidget()
        
        # Tab 0: Summary
        print("ProceedingsWorkspace: creating summary tab")
        self.summary_tab = self.create_summary_tab()
        self.content_stack.addWidget(self.summary_tab)
        
        # Tab 1: DRC-01A Editor
        print("ProceedingsWorkspace: creating drc01a tab")
        self.drc01a_tab = self.create_drc01a_tab()
        self.content_stack.addWidget(self.drc01a_tab)

        # Tab 1.5: ASMT-10 (Read-Only)
        print("ProceedingsWorkspace: creating asmt10 tab")
        self.asmt10_tab = self.create_asmt10_tab()
        self.content_stack.addWidget(self.asmt10_tab)
        
        # Tab 2: SCN (Show Cause Notice)
        print("ProceedingsWorkspace: creating scn tab")
        self.scn_tab = self.create_scn_tab()
        self.content_stack.addWidget(self.scn_tab)
        
        # Tab 3: PH Intimation (Personal Hearing)
        print("ProceedingsWorkspace: creating ph_intimation tab")
        self.ph_tab = self.create_ph_intimation_tab()
        self.content_stack.addWidget(self.ph_tab)
        
        # Tab 4: Order
        print("ProceedingsWorkspace: creating order_tab")
        self.order_tab = self.create_order_tab()
        self.content_stack.addWidget(self.order_tab)
        print("ProceedingsWorkspace: order_tab added")
        
        # Placeholders for Documents and Timeline
        for i, name in enumerate(["Documents", "Timeline"]):
            lbl = QLabel(f"{name} - Coming Soon")
            lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
            lbl.setStyleSheet("font-size: 12pt; color: #7f8c8d;")
            self.content_stack.addWidget(lbl)
            
        # 2a. Central Container (Header + Content Stack)
        self.central_container_widget = QWidget()
        self.central_container_layout = QVBoxLayout(self.central_container_widget)
        self.central_container_layout.setContentsMargins(0, 0, 0, 0)
        self.central_container_layout.setSpacing(0)

        # Header
        self.central_header = QWidget()
        self.central_header.setFixedHeight(40)
        self.central_header.setStyleSheet("background-color: #f8f9fa; border-bottom: 1px solid #e0e0e0;")
        header_layout = QHBoxLayout(self.central_header)
        header_layout.setContentsMargins(15, 0, 15, 0)
        
        self.context_title_lbl = QLabel("") # Dynamic Title
        self.context_title_lbl.setStyleSheet("font-weight: bold; color: #5f6368;")
        
        self.context_metadata_lbl = QLabel("")
        self.context_metadata_lbl.setStyleSheet("color: #7f8c8d; font-size: 10pt; font-weight: 500;")
        # Prevent long names from stretching the window bounds by ignoring size hint bounds
        self.context_metadata_lbl.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Preferred)
        
        header_layout.addWidget(self.context_title_lbl)
        header_layout.addSpacing(20)
        header_layout.addWidget(self.context_metadata_lbl, 1) # Allows it to consume space without expanding window
        
        self.central_container_layout.addWidget(self.central_header)
        self.central_container_layout.addWidget(self.content_stack)

        self.layout.addWidget(self.central_container_widget)
        
        # Initial Context
        self.apply_context_layout("summary")

    def add_drc01a_issue_card(self, template, data=None):
        """Reusable helper to add an issue card to the DRC-01A draft"""
        from src.ui.issue_card import IssueCard
        card = IssueCard(template, data=data, mode="DRC-01A")
        
        # Connect signals
        card.removeClicked.connect(lambda: self.remove_issue_card(card))
        card.valuesChanged.connect(self.calculate_grand_totals)
        
        self.issues_layout.addWidget(card)
        self.issue_cards.append(card)
        
        # Trigger initial calculation
        card.calculate_values()
        self.calculate_grand_totals()
        return card

    def restore_drc01a_draft_state(self):
        """
        Hydrate DRC-01A from 'case_issues' (Source of Truth).
        Handles Legacy Migration from 'additional_details' if case_issues is empty.
        """
        try:
            # 1. Fetch Authoritative Records
            records = self.db.get_case_issues(self.proceeding_id, stage='DRC-01A')
            
            # 2. Legacy Migration Fallback
            if not records:
                details = self.proceeding_data.get('additional_details', {})
                if 'issues' in details and details['issues']:
                    print("DRC-01A Hydration: Migrating Legacy Draft...")
                    legacy_issues = details['issues']
                    
                    # Persist as structured records
                    migrated_list = []
                    for issue in legacy_issues:
                        # Ensure 'data' wrapper if missing in legacy
                        migrated_list.append({
                            'issue_id': issue.get('issue_id'),
                            'data': issue.get('data', issue) 
                        })
                    
                    self.db.save_case_issues(self.proceeding_id, migrated_list, stage='DRC-01A')
                    
                    # Clear legacy artifact to ensure Single Source forward
                    details.pop('issues')
                    self.db.update_proceeding(self.proceeding_id, {'additional_details': details})
                    self.proceeding_data['additional_details'] = details # Local Sync
                    
                    # Refetch now that they are in DB
                    records = self.db.get_case_issues(self.proceeding_id, stage='DRC-01A')
            
            # 3. Hydrate Cards
            if records:
                print(f"DRC-01A Hydration: Restoring {len(records)} issues from DB.")
                for record in records:
                    issue_id = record['issue_id']
                    data_payload = record.get('data', {})
                    
                    # Resolve Template: Master > Name > Snapshot
                    template = self.issue_templates.get(issue_id)
                    if not template:
                         # Fetch deep if not in simple list
                         template = self.db.get_issue(issue_id)
                    
                    if template:
                        # Deepcopy to avoid mutating master
                        import copy
                        card_template = copy.deepcopy(template)
                        self.add_drc01a_issue_card(card_template, data=data_payload)
                    else:
                        print(f"Warning: Template not found for issue {issue_id}")

            # 4. Restore Sections Violated
            metadata = self.proceeding_data.get('additional_details', {}).get('drc01a_metadata', {})
            if 'sections_violated_html' in metadata:
                self.sections_editor.setHtml(metadata['sections_violated_html'])

        except Exception as e:
            print(f"Error restoring DRC-01A draft: {e}")
            import traceback
            traceback.print_exc()

    def handle_sidebar_action(self, action):
        """Switch tabs based on sidebar action with STRICT Workflow Enforcement"""
        action_map = {
            "summary": 0,
            "drc01a": 1,
            "asmt10": 2, # Stacks at index 2 (summary=0, drc01a=1, asmt10=2, scn=3...)
            "scn": 3,
            "ph": 4,
            "order": 5,
            "documents": 6,
            "timeline": 7
        }
        
        # [PHASE 2] Strict Workflow State Enforcement
        current_stage = self.get_current_stage()
        
        # 1. DRC-01A Logic
        # (Generally accessible for viewing, but edits might be restricted in the tab itself via ui state)
        # However, the task says "Editable only in DRC01A_DRAFT".
        # We handle "Editing" restrictions inside the tab (e.g. hiding buttons), 
        # but here we allow navigation to view the "Issued" state.
        
        # 2. SCN Access Logic
        # Rule: SCN drafting/viewing relies on DRC-01A being issued (or skipped).
        if action == "scn":
             # Exception: If DRC-01A is skipped (Direct Adjudication or configured so)
             # We need to check drc01a_skipped.
             is_skipped = self.proceeding_data.get('drc01a_skipped', False)
             # Also allow if strictly Direct Adjudication without Scrutiny origin? 
             # For now, rely on stage. 
             # Access allowed if Stage >= DRC01A_ISSUED OR Skipped
             allowed = (current_stage >= WorkflowStage.DRC01A_ISSUED) or is_skipped
             
             # Fallback for Direct Adjudication which might not set DRC01A_ISSUED if it skips it?
             # If skipped, stage might be SCRUTINY_ACTIVE or similar? 
             # Actually, creating direct adjudication sets stage to SCRUTINY_ACTIVE?
             # Let's assume strictness.
             if not allowed and current_stage < WorkflowStage.DRC01A_ISSUED:
                  QMessageBox.warning(self, "Workflow Restricted", "SCN Drafting is locked until DRC-01A is issued.")
                  return

        # 3. PH Access Logic
        # Rule: PH Intimation requires SCN Issued
        if action == "ph":
             if current_stage < WorkflowStage.SCN_ISSUED:
                  QMessageBox.warning(self, "Workflow Restricted", "Personal Hearing Intimation is locked until SCN is issued.")
                  return

        # 4. Order Access Logic
        # Rule: Order requires PH Scheduled (or SCN Issued if PH is waived?)
        if action == "order":
             if current_stage < WorkflowStage.PH_SCHEDULED:
                  # Maybe allow if SCN issued and enough time passed? 
                  # Strict rule: PH must be at least typically scheduled/waived.
                  # For now, require PH_SCHEDULED.
                  QMessageBox.warning(self, "Workflow Restricted", "Adjudication Order generation is locked until PH is scheduled.")
                  return


        if action in action_map:
            index = action_map[action]
            self.content_stack.setCurrentIndex(index)
            
            # Apply Layout context
            self.apply_context_layout(action)
            
            # Auto-load issues when switching to SCN tab
            if action == "scn":
                self.load_scn_issue_templates()

    def apply_context_layout(self, context_key):
        """
        Enforce single legal context layout rules.
        """
        # Update UI Title
        titles = {
            "scn": "Show Cause Notice Drafting",
            "drc01a": "DRC-01A Drafting",
            "summary": "Case Summary & Cockpit",
            "asmt10": "ASMT-10 Reference",
            "ph": "Personal Hearing Intimation", 
            "order": "Adjudication Order"
        }
        self.context_title_lbl.setText(titles.get(context_key, ""))
        
        # [PHASE 3] Dynamic Workflow Enforcement
        current_stage = self.get_current_stage()
        
        # 1. PH Conclude Button Visibility
        # Visible only if PH is intimated but not yet concluded
        if hasattr(self, 'ph_conclude_btn'):
             self.ph_conclude_btn.setVisible(current_stage == WorkflowStage.PH_SCHEDULED)
             
        # 2. Order Finalize Button Lock
        # Enabled only if PH is completed
        if hasattr(self, 'order_finalize_btn'):
             self.order_finalize_btn.setEnabled(current_stage == WorkflowStage.PH_COMPLETED)
             if current_stage < WorkflowStage.PH_COMPLETED:
                  self.order_finalize_btn.setToolTip("Personal Hearing must be concluded before finalizing the order.")
             else:
                  self.order_finalize_btn.setToolTip("")

        # SCN Specific Page 1 Navigation check (Metadata)
        is_scn = (context_key == "scn")
        if is_scn and hasattr(self, 'active_scn_step'):
            # This logic is mostly handled in on_scn_page_changed
            pass

    def create_asmt10_tab(self):
        """Create Read-Only ASMT-10 Tab with Professional Document Framing"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)
        
        # 1. Document Toolbar
        self.asmt10_toolbar = QFrame()
        self.asmt10_toolbar.setFixedHeight(50)
        self.asmt10_toolbar.setStyleSheet("""
            QFrame {
                background-color: white;
                border-bottom: 1px solid #e0e6ed;
            }
        """)
        toolbar_layout = QHBoxLayout(self.asmt10_toolbar)
        toolbar_layout.setContentsMargins(20, 0, 20, 0)
        
        title_vbox = QVBoxLayout()
        title_vbox.setSpacing(2)
        title_vbox.setAlignment(Qt.AlignmentFlag.AlignVCenter)
        
        self.asmt10_title_lbl = QLabel("🔒 Finalised ASMT-10 — Read-Only Reference")
        self.asmt10_title_lbl.setStyleSheet("font-weight: bold; font-size: 10pt; color: #1e293b;")
        title_vbox.addWidget(self.asmt10_title_lbl)
        
        self.asmt10_meta_lbl = QLabel("OC No. - | Date: -")
        self.asmt10_meta_lbl.setStyleSheet("font-size: 8pt; color: #64748b;")
        title_vbox.addWidget(self.asmt10_meta_lbl)
        
        toolbar_layout.addLayout(title_vbox)
        toolbar_layout.addStretch()
        
        # Letterhead Toggle
        self.asmt10_lh_cb = QCheckBox("Show Letterhead")
        self.asmt10_lh_cb.setChecked(True)
        self.asmt10_lh_cb.stateChanged.connect(self._on_asmt10_lh_toggled)
        toolbar_layout.addWidget(self.asmt10_lh_cb)
        
        toolbar_layout.addSpacing(20)
        
        # Zoom Controls (Now controls font size / scale in TextBrowser if possible, 
        # or we just remove them because TextBrowser handles standard zoom via Ctrl+Wheel)
        # Keeping them for consistency, but mapping to zoomIn/zoomOut
        zoom_layout = QHBoxLayout()
        zoom_layout.setSpacing(5)
        
        for icon, cmd in [("➖", self._zoom_asmt10_out), ("100%", self._zoom_asmt10_reset), ("➕", self._zoom_asmt10_in)]:
            btn = QPushButton(icon)
            btn.setFixedSize(40, 30)
            btn.setCursor(Qt.CursorShape.PointingHandCursor)
            btn.setStyleSheet("""
                QPushButton {
                    background-color: #f8fafc;
                    border: 1px solid #e2e8f0;
                    border-radius: 4px;
                    font-size: 9pt;
                }
                QPushButton:hover { background-color: #f1f5f9; }
            """)
            btn.clicked.connect(cmd)
            zoom_layout.addWidget(btn)
        
        toolbar_layout.addLayout(zoom_layout)
        layout.addWidget(self.asmt10_toolbar)
        
        # 2. Native HTML Preview (QWebEngineView)
        # Upgraded to support high-fidelity A4 simulation and JS pagination
        self.asmt10_browser = QWebEngineView()
        self.asmt10_browser.setStyleSheet("background-color: #525659;") 
        
        layout.addWidget(self.asmt10_browser)
        
        return widget

    def render_asmt10_preview(self, source_scrutiny_id):
        """
        Fetch structured Scrutiny Data and render ASMT-10 Preview.
        Strictly Read-Only and Runtime Generated.
        """
        scrutiny_data = self.db.get_scrutiny_case_data(source_scrutiny_id)
        
        if not scrutiny_data:
            self.asmt10_browser.setHtml("<h3 style='color:white; text-align:center; margin-top:50px;'>Source Scrutiny Data not found.</h3>")
            return

        # Update Toolbar Meta
        oc = scrutiny_data.get('oc_number', 'DRAFT')
        dt = scrutiny_data.get('asmt10_finalised_on', '-')
        self.asmt10_meta_lbl.setText(f"OC No. {oc} | Date: {dt}")
        
        try:
            case_info = {
                'case_id': scrutiny_data.get('case_id'),
                'financial_year': scrutiny_data.get('financial_year'),
                'section': scrutiny_data.get('section'),
                'notice_date': dt,
                'oc_number': oc,
                'last_date_to_reply': scrutiny_data.get('last_date_to_reply', 'N/A')
            }
            
            taxpayer = scrutiny_data.get('taxpayer_details', {})
            if not isinstance(taxpayer, dict): taxpayer = {}
            
            # [ENRICHMENT] Strictly preview-only: Fetch fresh taxpayer details
            # Ensure we use the correct key for GSTIN (handle both 'gstin' and 'taxpayer_gstin')
            gstin = scrutiny_data.get('gstin') or scrutiny_data.get('taxpayer_gstin')
            if not gstin and taxpayer:
                 gstin = taxpayer.get('GSTIN') or taxpayer.get('gstin')
                 
            if gstin:
                fresh_taxpayer = self.db.get_taxpayer(gstin)
                if fresh_taxpayer:
                    # Update preview dictionary with fresh data if missing or N/A
                    if taxpayer.get('Legal Name', 'N/A') in ['N/A', '', None]:
                        taxpayer['Legal Name'] = fresh_taxpayer.get('Legal Name', taxpayer.get('Legal Name', 'N/A'))
                    
                    if taxpayer.get('Address', 'Address not available') in ['Address not available', '', None]:
                         taxpayer['Address'] = (fresh_taxpayer.get('Address') or 
                                               fresh_taxpayer.get('Address of Principal Place of Business') or 
                                               taxpayer.get('Address', 'Address not available'))
                    
                # Also update gstin in case_info for the header
                case_info['gstin'] = gstin
            
            # Fix double fetch logic in original code
            issues = scrutiny_data.get('selected_issues', [])
            
            # Handle new parser format which returns {"metadata":..., "issues": [...]}
            if isinstance(issues, dict) and 'issues' in issues:
                 issues = issues['issues']
            
            generator = ASMT10Generator()
            full_data = case_info.copy()
            full_data['taxpayer_details'] = taxpayer
            full_data['gstin'] = gstin
            
            # [ALIGNMENT] Use "professional" style mode for high-fidelity A4 parity
            html_content = generator.generate_html(full_data, issues, for_preview=True, show_letterhead=self.asmt10_show_letterhead, style_mode="professional")
            
            # Update the browser directly
            self.asmt10_browser.setHtml(html_content)
                
        except Exception as e:
            print(f"ASMT-10 Render Error: {e}")
            self.asmt10_browser.setHtml(f"<h3 style='color:red; text-align:center;'>Error rendering ASMT-10: {e}</h3>")

    def _on_asmt10_lh_toggled(self, state):
        self.asmt10_show_letterhead = (state == Qt.CheckState.Checked.value or state == True)
        source_id = self.proceeding_data.get('source_scrutiny_id') or self.proceeding_data.get('scrutiny_id')
        if source_id:
            self.render_asmt10_preview(source_id)

    def _zoom_asmt10_in(self):
        self.asmt10_browser.setZoomFactor(self.asmt10_browser.zoomFactor() + 0.1)

    def _zoom_asmt10_out(self):
        self.asmt10_browser.setZoomFactor(max(0.1, self.asmt10_browser.zoomFactor() - 0.1))

    def _zoom_asmt10_reset(self):
        self.asmt10_browser.setZoomFactor(1.0)

    def _show_preview_error(self, message):
         self.asmt10_browser.setHtml(f"<h3 style='color:red; text-align:center;'>{message}</h3>")

    def _clear_asmt10_preview(self):
        """Helper to safely clear the ASMT-10 page container"""
        if not hasattr(self, 'asmt10_page_layout'):
            return
        while self.asmt10_page_layout.count():
            child = self.asmt10_page_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()

    def _ensure_dict(self, data):
        """Helper to safely ensure data is a dict, handling double-serialized JSON strings."""
        if isinstance(data, dict):
            return data
        if not data:
            return {}
        
        # If it's a string, try to parse
        if isinstance(data, str):
            try:
                parsed = json.loads(data)
                # Recursion to handle double-encoding (e.g. '"{\"a\":1}"')
                return self._ensure_dict(parsed)
            except Exception:
                # If parsing fails, return empty dict or log error? 
                # Better to return empty dict to prevent app crash on access
                print(f"Warning: Failed to parse JSON data: {data}")
                return {}
        
        return {}

    def _clear_scn_workspace_state(self):
        """Reset SCN workspace to pristine state before loading new data."""
        if hasattr(self, 'scn_issue_cards'):
            for card in self.scn_issue_cards:
                card.setParent(None)
                if hasattr(self, 'scn_issues_layout'):
                    self.scn_issues_layout.removeWidget(card)
                card.deleteLater()
            self.scn_issue_cards = []
            
        # Reset Workflow Flags
        self.scn_issues_initialized = False
        self.scn_workflow_phase = "METADATA"
        self.preview_initialized = False

    def hydrate_proceeding_data(self):
        """
        Idempotently normalize self.proceeding_data['additional_details'].
        Must handle str/dict/null and corrupted JSON safely.
        """
        raw = self.proceeding_data.get("additional_details")
        parsed = {}

        if isinstance(raw, str):
            if not raw.strip():
                parsed = {}
            else:
                try:
                    import json
                    parsed = json.loads(raw)
                    # Handle potential double-serialization
                    if isinstance(parsed, str):
                        try:
                            parsed = json.loads(parsed)
                        except: pass
                except (json.JSONDecodeError, TypeError):
                    print(f"Warning: Failed to hydrate additional_details JSON. Falling back to {{}}.")
                    parsed = {}
        elif isinstance(raw, dict):
            parsed = raw
        else:
            parsed = {}

        # Type Enforcement & Normalization
        if not isinstance(parsed.get("ph_entries"), list):
            parsed["ph_entries"] = []

        if not isinstance(parsed.get("scn_metadata"), dict):
            parsed["scn_metadata"] = {}

        if not isinstance(parsed.get("drc01a_metadata"), dict):
            parsed["drc01a_metadata"] = {}

        # Assign back atomically
        self.proceeding_data["additional_details"] = parsed

    def hydrate_scn_grounds_data(self):
        """
        Ensure SCN grounds data exists and handle migration/defaults. (Safe Method)
        Strategy:
        - If 'scn_grounds' missing -> Initialize default (Manual Override ON for safety).
        - Try to populate ASMT-10 details from snapshot if available.
        - Updates self.proceeding_data in-place (memory only).
        """
        try:
            import json
            details = self.proceeding_data.get('additional_details', {})
            if not details: details = {}
            # details is already dict due to hydrate_proceeding_data, but be safe
            if isinstance(details, str): return 

            grounds = details.get('scn_grounds')
            
            # --- PHASE 20/21: RECOVERY LOGIC for Blanched Cases ---
            # If grounds exists BUT it has the 'Stale Default' list while files are now available,
            # we should re-calculate to recover from the Metadata Blanching bug.
            
            def _get_current_file_doc_list(details_dict):
                f_paths = details_dict.get('file_paths', {})
                d_list = []
                if 'tax_liability_yearly' in f_paths: d_list.append("Tax Liability Excel")
                if any(k.startswith('gstr3b') for k in f_paths): d_list.append("GSTR-3B")
                if any(k.startswith('gstr1') for k in f_paths): d_list.append("GSTR-1")
                if any(k.startswith('gstr2a') for k in f_paths): d_list.append("GSTR-2A")
                if any(k.startswith('gstr2b') for k in f_paths): d_list.append("GSTR-2B")
                if any(k.startswith('gstr9') for k in f_paths): d_list.append("GSTR-9")
                if any(k.startswith('gstr9c') for k in f_paths): d_list.append("GSTR-9C")
                return d_list

            # [PHASE 4] Dynamic Draft Typing (Registry is Truth)
            source_type = self.proceeding_data.get('source_type', 'SCRUTINY').lower()
            
            # Initialize if missing (Legacy or New Case)
            if not grounds:
                print(f"SCN Grounds: Hydrating default structure ({source_type}).")
                doc_list = _get_current_file_doc_list(details)
                
                # Fallback default if no files
                if not doc_list:
                    # [PHASE 4] Smart Defaults: Field Audits (Adjudication) might not have these
                    if source_type == 'adjudication':
                        doc_list = [] # Start blank for direct adj
                    else:
                        doc_list = ["GSTR-1", "GSTR-3B", "GSTR-2A"]

                # Create Structure
                grounds = {
                    "version": 1,
                    "type": source_type, # [FIX] Dynamic Type
                    "manual_override": True,
                    "manual_text": "",
                    "data": {
                        "financial_year": self.proceeding_data.get('financial_year', '-'),
                        "docs_verified": doc_list, 
                        # ASMT-10 Ref only for Scrutiny
                        "asmt10_ref": {
                            "oc_no": "", 
                            "date": "",
                            "officer_designation": "Proper Officer",
                            "office_address": ""
                        } if source_type == 'scrutiny' else {},
                        "reply_ref": {
                            "received": False,
                            "date": None
                        }
                    }
                }
                details['scn_grounds'] = grounds
                self.proceeding_data['additional_details'] = details
            else:
                # [PHASE 4] Registry Repair (Truth Enforcement)
                current_type = grounds.get('type')
                if current_type != source_type:
                    print(f"SCN Grounds: Repairing type mismatch ('{current_type}' -> '{source_type}')")
                    grounds['type'] = source_type
                    # If migrating TO adjudication, maybe clear asmt10_ref? 
                    # For now, we just fix the type tag to ensure correct rendering downstream.
                    details['scn_grounds'] = grounds
                    self.proceeding_data['additional_details'] = details

                # [RECOVERY] Check if it's a blanched draft
                current_docs = grounds.get('data', {}).get('docs_verified', [])
                STALE_DEFAULT = ["GSTR-1", "GSTR-3B", "GSTR-2A"]
                
                if current_docs == STALE_DEFAULT:
                    fresh_list = _get_current_file_doc_list(details)
                    # If we found more files now (due to Deep Merge Fix), update!
                    if fresh_list and fresh_list != STALE_DEFAULT:
                        print("SCN Grounds: Recovering blanched doc list.")
                        grounds['data']['docs_verified'] = fresh_list
                        # Update the pointer just in case
                        details['scn_grounds'] = grounds
                        self.proceeding_data['additional_details'] = details
                
        except Exception as e:
            print(f"Error hydrating SCN grounds: {e}")
            import traceback
            traceback.print_exc()

    def load_proceeding(self, pid):
        # [FIX] State Persistence: Clear existing workspace first
        self._clear_scn_workspace_state()
        
        self.is_hydrated = False
        self.proceeding_id = pid
        self.proceeding_data = self.db.get_proceeding(pid)
        
        if not self.proceeding_data:
            # Critical Fix: Ensure data is never None
            self.proceeding_data = {}
            QMessageBox.critical(self, "Error", "Proceeding not found!")
            pass # Clear preview logic removed
            return

        # [PHASE 15] Centralized Hydration
        self.hydrate_proceeding_data()
        self.hydrate_scn_grounds_data() # [NEW] Hydrate SCN Grounds
        
        # [NEW] Bind Data to Grounds Form if UI is ready
        if hasattr(self, 'scn_grounds_form'):
            details = self.proceeding_data.get('additional_details', {})
            grounds_data = details.get('scn_grounds')
            
            # [PHASE 18] UI Overlay (Generation-Time Linkage)
            # Deep copy to avoid mutating stored JSON in memory
            import copy
            if grounds_data:
                ui_view = copy.deepcopy(grounds_data)
                
                # Overlay Authoritative Identifiers (ONLY for Scrutiny)
                source_type = self.proceeding_data.get('source_type', 'SCRUTINY')
                
                if source_type == 'SCRUTINY' and 'data' in ui_view:
                    if 'asmt10_ref' not in ui_view['data']:
                        ui_view['data']['asmt10_ref'] = {}
                    
                    # Fetch authoritative IDs
                    auth_oc = self.proceeding_data.get('oc_number', '')
                    auth_date = self.proceeding_data.get('notice_date', '')
                    
                    ui_view['data']['asmt10_ref']['oc_no'] = auth_oc
                    ui_view['data']['asmt10_ref']['date'] = auth_date
                    
                self.scn_grounds_form.set_data(ui_view)
            else:
                 self.scn_grounds_form.set_data(None)
            
            # [NEW] Introductory Paragraph auto-regeneration logic
            self.scn_grounds_form.auto_regenerate()
            
        self.proceeding_data['taxpayer_details'] = self._ensure_dict(self.proceeding_data.get('taxpayer_details'))
        

            
        # Fetch associated documents
        self.documents = self.db.get_documents(pid)
            
        # Update UI
        self.update_summary_tab()
        
        # [PHASE 3] Workspace Branching
        source_type = self.proceeding_data.get('source_type', 'SCRUTINY')
        self._configure_workspace_mode(source_type)

    def _configure_workspace_mode(self, source_type):
        """
        Configure UI states based on Case Source Type (Registry-Driven).
        """
        pre_scn_mode = 'ASMT10' # Default
        
        if source_type == 'SCRUTINY':
            # Scrutiny Mode -> ASMT10 Origin
            pre_scn_mode = 'ASMT10'
            
            source_scrutiny_id = self.proceeding_data.get('source_scrutiny_id') or self.proceeding_data.get('scrutiny_id')
            if source_scrutiny_id:
                self.render_asmt10_preview(source_scrutiny_id)
            

            
        elif source_type == 'ADJUDICATION':
            # Adjudication Mode -> Check Source
            source_scrutiny_id = self.proceeding_data.get('source_scrutiny_id')
            
            if source_scrutiny_id:
                # Scrutiny-Origin Adjudication
                pre_scn_mode = 'ASMT10'
                self.render_asmt10_preview(source_scrutiny_id)
            else:
                # Direct Adjudication
                pre_scn_mode = 'DRC01A'

                
                # Attach DRC-01A Logic (Ensure panel exists/is visible)
                if hasattr(self, 'drc01a_tab'):
                    # Logic to enable DRC-01A editing would go here
                    pass
        
        # [REVERTED] Sidebar is managed by MainWindow -> Sidebar.set_mode().
        # We do not rebuild it here to avoid destroying inner sidebars (SCN/PH).
        
        
        # Restore Draft State (Context Aware)
        # [PHASE 5] Strict Hydration per Document Type
        if pre_scn_mode == 'DRC01A':
             self.restore_drc01a_draft_state()
        elif pre_scn_mode == 'ASMT10':
             pass # ASMT-10 is read-only rendering, handled by render_asmt10_preview
        
        # [PHASE 15] Hydrate SCN Initialization Flag from normalized details
        self.scn_issues_initialized = False
        details = self.proceeding_data.get('additional_details', {})
        scn_meta = details.get('scn_metadata', {})
        self.scn_issues_initialized = scn_meta.get('scn_issues_initialized') or details.get('scn_issues_initialized', False)
        print(f"ProceedingsWorkspace: SCN Hydration State = {self.scn_issues_initialized}")

        # Check for existing generated documents to toggle View Mode
        self.check_existing_documents()
        
        # --- Adjudication Setup Check ---
        if self.proceeding_data.get('is_adjudication') and not self.proceeding_data.get('adjudication_section'):
            print("Adjudication Setup Required: Launching Dialog")
            dlg = AdjudicationSetupDialog(self)
            if dlg.exec():
                # Save Section
                section = dlg.selected_section
                success = self.db.update_adjudication_case(pid, {'adjudication_section': section})
                if success:
                    # Update local data
                    self.proceeding_data['adjudication_section'] = section
                    # Also set initiating_section for compatibility if needed, but preference is distinct field
                    self.proceeding_data['initiating_section'] = section 
                    QMessageBox.information(self, "Setup Complete", f"Case initialized under Section {section}.")
                else:
                    QMessageBox.critical(self, "Error", "Failed to save adjudication section.")
            else:
                QMessageBox.warning(self, "Setup Incomplete", "SCN drafting is disabled until section is selected.")
                # We should probably disable tabs here or mark as incomplete
        
        # Guard: validate section after setup dialog — disable DRC-01A tab if still missing
        resolved_sec = self.proceeding_data.get('adjudication_section') or self.proceeding_data.get('initiating_section')
        if self.proceeding_data.get('is_adjudication') and hasattr(self, 'content_stack'):
            drc01a_tab_idx = 2  # Assumed index – skip silently if out of range
            sec_missing = not resolved_sec or not str(resolved_sec).strip()
            if hasattr(self, 'tab_bar'):
                self.tab_bar.setTabEnabled(drc01a_tab_idx, not sec_missing)
        
        # Hydrate Persistent Top Metadata Header
        legal_name = self.proceeding_data.get('legal_name')
        if not legal_name or not str(legal_name).strip():
            legal_name = "Unknown Taxpayer"
            
        fy = self.proceeding_data.get('financial_year')
        if not fy or not str(fy).strip():
            fy = "Unknown FY"
            
        section = self.proceeding_data.get('adjudication_section') or self.proceeding_data.get('initiating_section')
        if not section or not str(section).strip():
            section = "Unknown Section"
            
        # Manually clamp name length to guarantee no massive window expanding forces
        max_name_len = 60
        display_name = (str(legal_name)[:max_name_len] + '...') if len(str(legal_name)) > max_name_len else str(legal_name)
        
        self.context_metadata_lbl.setText(f"  {display_name}   |   FY {fy}   |   Section {section}")
        
        self.is_hydrated = True
        


    def create_side_nav_layout(self, items, page_changed_callback=None, use_scroll=True):
        """
        Creates a side navigation layout (Accordion/Master-Detail).
        items: list of tuples (title, icon_char, widget)
        page_changed_callback: optional function(index) called when page switches
        Returns: main_widget containing the layout
        """
        main_widget = QWidget()
        main_layout = QHBoxLayout(main_widget)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)
        
        # 1. Left Sidebar
        sidebar = QFrame()
        sidebar.setFixedWidth(250)
        sidebar.setStyleSheet("background-color: white; border-right: 1px solid #e0e0e0;")
        sidebar_layout = QVBoxLayout(sidebar)
        sidebar_layout.setContentsMargins(0, 20, 0, 20)
        sidebar_layout.setSpacing(5)
        
        nav_cards = []
        content_stack = QStackedWidget()
        
        # Store cards on the widget so they can be accessed externally
        main_widget.nav_cards = nav_cards 
        
        def switch_page(index):
            # Guard
            if index < len(nav_cards) and not nav_cards[index].is_enabled():
                return

            # Update Active State
            for i, card in enumerate(nav_cards):
                card.set_active(i == index)
            content_stack.setCurrentIndex(index)
            
            # Call callback if provided
            if page_changed_callback:
                page_changed_callback(index)

        for i, (title, icon, page_widget) in enumerate(items):
            # Create Nav Card
            card = SideNavCard(i, icon, title)
            card.clicked.connect(switch_page)
            sidebar_layout.addWidget(card)
            nav_cards.append(card)
            
            # Add Page to Stack
            # Wrap page in scroll area if needed, but usually the page itself handles it
            # Let's ensure uniform styling for pages
            page_container = QWidget()
            page_layout = QVBoxLayout(page_container)
            page_layout.setContentsMargins(20, 20, 20, 20)
            
            # Header for the page
            header = QLabel(title)
            header.setStyleSheet("font-size: 15pt; font-weight: bold; color: #2c3e50; margin-bottom: 15px;")
            page_layout.addWidget(header)
            
            page_layout.addWidget(page_widget, 1) # Force stretch to fill container
            
            if use_scroll:
                # Scroll Area for the page content
                scroll = QScrollArea()
                scroll.setWidgetResizable(True)
                scroll.setWidget(page_container)
                scroll.setFrameShape(QFrame.Shape.NoFrame)
                scroll.setStyleSheet("background-color: #f8f9fa;")
                content_stack.addWidget(scroll)
            else:
                page_container.setStyleSheet("background-color: #f8f9fa;")
                content_stack.addWidget(page_container)
            
        sidebar_layout.addStretch()
        
        # Set first item active
        if nav_cards:
            switch_page(0)
            
        main_layout.addWidget(sidebar)
        main_layout.addWidget(content_stack)
        
        return main_widget


    def create_summary_tab(self):
        """
        Create Read-Only Summary Tab with Refined Case Cockpit Layout.
        Focus: Legal Authority and Procedural Context.
        """
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.Shape.NoFrame)
        scroll.setStyleSheet("background-color: #f8fafc;")

        self.summary_content = QWidget()
        self.summary_layout = QVBoxLayout(self.summary_content)
        self.summary_layout.setContentsMargins(20, 20, 20, 20)
        self.summary_layout.setSpacing(15)

        # 1. Header Card (Case Identity)
        self.header_card = QFrame()
        self.header_card.setObjectName("HeaderCard")
        self.header_card.setStyleSheet("""
            #HeaderCard {
                background-color: white;
                border: 1px solid #e2e8f0;
                border-radius: 12px;
                padding: 20px;
            }
        """)
        header_layout = QHBoxLayout(self.header_card)
        
        info_layout = QVBoxLayout()
        self.lbl_legal_name = QLabel("-")
        self.lbl_legal_name.setStyleSheet("font-size: 18pt; font-weight: bold; color: #1e293b;")
        info_layout.addWidget(self.lbl_legal_name)
        
        self.lbl_gstin = QLabel("-")
        self.lbl_gstin.setStyleSheet("font-family: monospace; font-size: 10pt; color: #64748b;")
        info_layout.addWidget(self.lbl_gstin)
        header_layout.addLayout(info_layout)
        
        header_layout.addStretch()
        
        # FY Badge
        self.fy_badge = QLabel("-")
        self.fy_badge.setStyleSheet("""
            background-color: #f1f5f9;
            color: #475569;
            border-radius: 12px;
            padding: 5px 15px;
            font-weight: bold;
            font-size: 10pt;
            border: 1px solid #e2e8f0;
        """)
        header_layout.addWidget(self.fy_badge, 0, Qt.AlignmentFlag.AlignTop)
        
        self.summary_layout.addWidget(self.header_card)

        # 2. Status Banner (Dominant Element)
        self.status_banner = QLabel("-")
        self.status_banner.setWordWrap(True)
        self.status_banner.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.status_banner.setFixedHeight(50)
        self.status_banner.setStyleSheet("""
            font-weight: 800;
            font-size: 11pt;
            border-radius: 8px;
            padding: 12px;
            border: 1px solid transparent;
        """)
        self.summary_layout.addWidget(self.status_banner)

        # 3. Legal Basis (Chips)
        self.chips_layout = QHBoxLayout()
        self.chips_layout.setSpacing(10)
        
        self.chip_origin = QLabel("-")
        self.chip_section = QLabel("-")
        
        for chip in [self.chip_origin, self.chip_section]:
            chip.setStyleSheet("""
                background-color: #f1f5f9;
                color: #475569;
                border-radius: 15px;
                padding: 4px 12px;
                font-size: 8pt;
                font-weight: 600;
                border: 1px solid #e2e8f0;
            """)
            self.chips_layout.addWidget(chip)
        
        self.chips_layout.addStretch()
        self.summary_layout.addLayout(self.chips_layout)

        # 4. Authority Block (Issuance Metadata) - Softened Visuals
        auth_container = QWidget()
        auth_layout = QVBoxLayout(auth_container)
        auth_layout.setContentsMargins(0, 5, 0, 5)
        auth_layout.setSpacing(0)
        
        # Header Row (Subtle)
        h_row = QWidget()
        h_layout = QHBoxLayout(h_row)
        h_layout.setContentsMargins(15, 10, 15, 10)
        for text in ["Document", "O.C. Number / Reference", "Date", "Status"]:
            lbl = QLabel(text)
            lbl.setStyleSheet("font-weight: bold; color: #94a3b8; font-size: 8pt; text-transform: uppercase; letter-spacing: 0.5px;")
            h_layout.addWidget(lbl, 1)
        auth_layout.addWidget(h_row)
        
        # Subtle Divider Line
        divider = QFrame()
        divider.setFrameShape(QFrame.Shape.HLine)
        divider.setStyleSheet("color: #f1f5f9;")
        auth_layout.addWidget(divider)
        
        # Rows
        self.auth_rows = []
        for i, doc_name in enumerate(["ASMT-10", "SCN", "Order"]):
            row = QFrame()
            row_layout = QHBoxLayout(row)
            row_layout.setContentsMargins(15, 12, 15, 12)
            
            name_lbl = QLabel(doc_name)
            ref_lbl = QLabel("—")
            date_lbl = QLabel("—")
            stat_lbl = QLabel("—")
            
            for lbl in [name_lbl, ref_lbl, date_lbl, stat_lbl]:
                lbl.setStyleSheet("font-size: 10pt; color: #1e293b;")
                row_layout.addWidget(lbl, 1)
            
            # Special highlighting for ASMT-10 (First Row)
            if i == 0:
                name_lbl.setStyleSheet("font-weight: bold; font-size: 10pt; color: #1e293b;")
            
            auth_layout.addWidget(row)
            self.auth_rows.append({
                'name': name_lbl, 'ref': ref_lbl, 'date': date_lbl, 'status': stat_lbl, 'frame': row
            })
            
            # Row Divider
            if i < 2:
                r_div = QFrame()
                r_div.setFrameShape(QFrame.Shape.HLine)
                r_div.setStyleSheet("color: #f8fafc;")
                auth_layout.addWidget(r_div)
            
        self.summary_layout.addWidget(auth_container)

        # 5. Next Action Hint
        self.next_action_hint = QLabel("")
        self.next_action_hint.setStyleSheet("color: #64748b; font-style: italic; font-size: 10pt; margin-top: 20px; font-weight: 500;")
        self.next_action_hint.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.summary_layout.addWidget(self.next_action_hint)

        self.summary_layout.addStretch()
        scroll.setWidget(self.summary_content)
        return scroll

    def update_summary_tab(self):
        if not self.proceeding_data:
            return
            
        data = self.proceeding_data
        
        # 1. Header Card
        self.lbl_legal_name.setText(data.get('legal_name', 'Unknown Taxpayer'))
        self.lbl_gstin.setText(data.get('gstin', '-'))
        self.fy_badge.setText(f"FY {data.get('financial_year', '-')}")
        
        # 2. Status Banner (Strict logic derived from workflow_stage)
        current_stage = self.get_current_stage()
        
        # Issuance facts (for authority block below)
        source_id = data.get('source_scrutiny_id') or data.get('scrutiny_id')
        asmt_issued = source_id and data.get('oc_number')
        scn_issued = data.get('scn_number')
        ord_issued = data.get('order_number')
        
        banner_text = "Status: - "
        banner_style = ""
        
        if current_stage == WorkflowStage.ORDER_ISSUED:
            banner_text = "⚖️ Order Issued"
            banner_style = "background-color: #f0fdf4; color: #166534; border: 1px solid #bbf7d0;"
        elif current_stage == WorkflowStage.PH_SCHEDULED:
            banner_text = "📅 PH Scheduled"
            banner_style = "background-color: #fffbeb; color: #92400e; border: 1px solid #fde68a;"
        elif current_stage == WorkflowStage.SCN_ISSUED:
            banner_text = "⚖️ Show Cause Notice Issued"
            banner_style = "background-color: #f0fdf4; color: #166534; border: 1px solid #bbf7d0;"
        elif current_stage == WorkflowStage.SCN_DRAFT:
            banner_text = "✏️ SCN Draft in Progress"
            banner_style = "background-color: #fffbeb; color: #92400e; border: 1px solid #fde68a;"
        elif current_stage == WorkflowStage.ASMT10_ISSUED:
            banner_text = "🔒 ASMT-10 Finalised — SCN Pending"
            banner_style = "background-color: #eff6ff; color: #1e40af; border: 1px solid #bfdbfe;"
        elif current_stage == WorkflowStage.DRC01A_ISSUED:
            banner_text = "📄 DRC-01A Issued — Awaiting Response"
            banner_style = "background-color: #eff6ff; color: #1e40af; border: 1px solid #bfdbfe;"
        elif current_stage == WorkflowStage.DRC01A_DRAFT:
            banner_text = "✏️ DRC-01A Draft"
            banner_style = "background-color: #fffbeb; color: #92400e; border: 1px solid #fde68a;"
        else:
            banner_text = f"Status: {data.get('status', 'Active')}"
            banner_style = "background-color: #f1f5f9; color: #475569; border: 1px solid #e2e8f0;"
            
        self.status_banner.setText(banner_text)
        self.status_banner.setStyleSheet(self.status_banner.styleSheet() + banner_style)
        
        # 3. Legal Basis Chips
        origin_text = "Scrutiny-origin Adjudication" if source_id else "Direct Adjudication"
        self.chip_origin.setText(origin_text)
        
        adj_section = data.get('adjudication_section') or "Section Pending"
        self.chip_section.setText(f"Section {adj_section}" if adj_section != "Section Pending" else adj_section)
        
        # 4. Authority Block (Hierarchy & Emphasis)
        # Row 0: ASMT-10 (The Legal Foundation)
        row_asmt = self.auth_rows[0]
        
        is_direct = data.get('source_type') == 'ADJUDICATION' and not source_id
        if is_direct:
             row_asmt['frame'].setVisible(False)
        else:
             row_asmt['frame'].setVisible(True)
             asmt_oc = data.get('oc_number') if source_id else "—"
             asmt_date = data.get('asmt10_finalised_on') or data.get('notice_date') if source_id else "—"
             
             if asmt_issued:
                 row_asmt['ref'].setText(asmt_oc)
                 row_asmt['date'].setText(asmt_date or "—")
                 row_asmt['status'].setText("✔ Issued")
                 row_asmt['status'].setStyleSheet("color: #166534; font-weight: 800; font-size: 8pt;")
                 row_asmt['frame'].setStyleSheet("background-color: #f0fdf4; border-radius: 6px; border-bottom: 1px solid #dcfce7;")
                 row_asmt['name'].setStyleSheet("font-weight: 800; color: #1e293b;")
             else:
                 for k in ['ref', 'date', 'status']: row_asmt[k].setText("—")
                 row_asmt['frame'].setStyleSheet("opacity: 0.5;")
                 row_asmt['status'].setStyleSheet("color: #94a3b8;")

        # Row 1: SCN
        row_scn = self.auth_rows[1]
        if scn_issued:
            row_scn['ref'].setText(scn_issued)
            row_scn['date'].setText(data.get('scn_date') or "—")
            row_scn['status'].setText("✔ Issued")
            row_scn['status'].setStyleSheet("color: #166534; font-weight: bold;")
            row_scn['frame'].setStyleSheet("background-color: #ffffff; border-bottom: 1px solid #f1f5f9;")
            row_scn['name'].setStyleSheet("font-weight: bold; color: #1e293b;")
        else:
            row_scn['ref'].setText("—")
            row_scn['date'].setText("—")
            row_scn['status'].setText("Not Issued")
            row_scn['status'].setStyleSheet("color: #94a3b8;")
            row_scn['name'].setStyleSheet("color: #94a3b8;")
            row_scn['frame'].setStyleSheet("opacity: 0.5;") # Muted until issued

        # Row 2: Order
        row_ord = self.auth_rows[2]
        if ord_issued:
            row_ord['ref'].setText(ord_issued)
            row_ord['date'].setText(data.get('order_date') or "—")
            row_ord['status'].setText("✔ Issued")
            row_ord['status'].setStyleSheet("color: #166534; font-weight: bold;")
            row_ord['frame'].setStyleSheet("background-color: #ffffff; border-bottom: 1px solid #f1f5f9;")
            row_ord['name'].setStyleSheet("font-weight: bold; color: #1e293b;")
        else:
            for k in ['ref', 'date']: row_ord[k].setText("—")
            row_ord['status'].setText("—")
            row_ord['status'].setStyleSheet("color: #94a3b8;")
            row_ord['name'].setStyleSheet("color: #94a3b8;")
            row_ord['frame'].setStyleSheet("opacity: 0.5;") # Muted

        # 5. Next Action Guidance (Informational)
        hint = ""
        if not ord_issued:
            if asmt_issued and not scn_issued:
                hint = "Next Action: Draft Show Cause Notice (SCN)"
            elif scn_issued and not ord_issued:
                hint = "Next Action: Issue Personal Hearing Intimation"
        
        self.next_action_hint.setText(hint)

    def create_drc01a_tab(self):
        # ROOT CONTAINER (Returned to content_stack)
        tab_root = QWidget()
        layout = QVBoxLayout(tab_root)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        print("ProceedingsWorkspace: create_drc01a_tab start (4-Step Refactor)")
        # Initialize
        self.issue_cards = []
        self.ACT_PRIORITY = ["IGST", "CGST", "SGST", "UTGST", "Cess"]
        
        # Debounce timer for preview refresh
        self.drc01a_refresh_timer = QtCore.QTimer()
        self.drc01a_refresh_timer.setSingleShot(True)
        self.drc01a_refresh_timer.timeout.connect(self._run_refresh_step4)

        # --- DRAFT CONTAINER ---
        self.drc01a_draft_container = QWidget()
        draft_layout = QVBoxLayout(self.drc01a_draft_container)
        draft_layout.setContentsMargins(0, 0, 0, 0)
        
        # --- STEP 1: Reference Details ---
        ref_widget = QWidget()
        ref_layout = QVBoxLayout(ref_widget)
        ref_layout.setContentsMargins(20, 20, 20, 20)
        
        ref_card = QWidget()
        ref_card.setStyleSheet("background-color: white; border: 1px solid #e0e0e0; border-radius: 8px;")
        ref_inner = QVBoxLayout(ref_card)
        
        ref_grid = QGridLayout()
        ref_grid.setSpacing(15)
        
        ref_grid.addWidget(QLabel("OC No (Intimation):"), 0, 0)
        self.oc_number_input = QLineEdit()
        ref_grid.addWidget(self.oc_number_input, 0, 1)
        
        suggest_btn = QPushButton("Get Next")
        suggest_btn.clicked.connect(lambda: self.suggest_next_oc(self.oc_number_input))
        ref_grid.addWidget(suggest_btn, 0, 2)
        
        ref_grid.addWidget(QLabel("OC Date:"), 0, 3)
        self.oc_date_input = QDateEdit()
        self.oc_date_input.setCalendarPopup(True)
        self.oc_date_input.setDate(QDate.currentDate())
        ref_grid.addWidget(self.oc_date_input, 0, 4)

        ref_grid.addWidget(QLabel("Issuing Officer:"), 1, 0)
        self.officer_combo = QComboBox()
        self.load_active_officers()
        ref_grid.addWidget(self.officer_combo, 1, 1, 1, 4)
        
        ref_inner.addLayout(ref_grid)
        ref_layout.addWidget(ref_card)
        ref_layout.addStretch()

        # --- STEP 2: Issues Involved ---
        issues_widget = QWidget()
        issues_vbox = QVBoxLayout(issues_widget)
        issues_vbox.setContentsMargins(0,0,0,0)
        issues_vbox.setSpacing(0)
        
        # Toolbar
        toolbar = QWidget()
        toolbar.setStyleSheet("background-color: #f8f9fa; border-bottom: 1px solid #e0e0e0;")
        tb_layout = QHBoxLayout(toolbar)
        self.issue_combo = QComboBox()
        self.issue_combo.addItem("Select Issue Template...", None)
        self.load_issue_templates()
        tb_layout.addWidget(QLabel("<b>Add Issue:</b>"))
        tb_layout.addWidget(self.issue_combo, 1)
        add_btn = QPushButton("Insert")
        add_btn.setStyleSheet("background-color: #3498db; color: white;")
        add_btn.clicked.connect(self.insert_selected_issue)
        tb_layout.addWidget(add_btn)
        issues_vbox.addWidget(toolbar)
        
        # Scroll Area for Issues
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.Shape.NoFrame)
        self.issues_container = QWidget()
        self.issues_layout = QVBoxLayout(self.issues_container)
        self.issues_layout.setAlignment(Qt.AlignmentFlag.AlignTop)
        self.issues_layout.setSpacing(20)
        self.issues_layout.setContentsMargins(20, 20, 20, 20)
        scroll.setWidget(self.issues_container)
        issues_vbox.addWidget(scroll, 1) # Take space

        # Bottom Total Card (Always Visible in Step 2)
        self.step2_total_card = QFrame()
        self.step2_total_card.setStyleSheet("background-color: #2c3e50; color: white; border-bottom-left-radius: 0px; border-bottom-right-radius: 0px;")
        st2_layout = QHBoxLayout(self.step2_total_card)
        st2_layout.setContentsMargins(20, 15, 20, 15)
        st2_layout.addWidget(QLabel("<b>Total Liability:</b>"))
        self.lbl_total_tax = QLabel("₹ 0")
        self.lbl_total_tax.setStyleSheet("font-size: 16px; font-weight: bold; color: #f1c40f;")
        st2_layout.addWidget(self.lbl_total_tax)
        st2_layout.addStretch()
        issues_vbox.addWidget(self.step2_total_card)

        # --- STEP 3: Compliance & Summary ---
        comp_widget = QWidget()
        comp_layout = QVBoxLayout(comp_widget)
        comp_layout.setContentsMargins(20, 20, 20, 20)
        comp_layout.setSpacing(20)

        # Dates & Summary Table Grid
        comp_top_card = QWidget()
        comp_top_card.setStyleSheet("background: white; border: 1px solid #e0e0e0; border-radius: 8px;")
        ctc_layout = QVBoxLayout(comp_top_card)
        
        date_grid = QGridLayout()
        date_grid.addWidget(QLabel("<b>Reply Deadline:</b>"), 0, 0)
        self.reply_deadline_input = QDateEdit()
        self.reply_deadline_input.setCalendarPopup(True)
        self.reply_deadline_input.setDate(QDate.currentDate().addDays(15))
        date_grid.addWidget(self.reply_deadline_input, 0, 1)
        
        date_grid.addWidget(QLabel("<b>Payment Deadline:</b>"), 0, 2)
        self.payment_deadline_input = QDateEdit()
        self.payment_deadline_input.setCalendarPopup(True)
        self.payment_deadline_input.setDate(QDate.currentDate().addDays(15))
        date_grid.addWidget(self.payment_deadline_input, 0, 3)
        ctc_layout.addLayout(date_grid)
        
        comp_layout.addWidget(comp_top_card)

        # Tax Summary Table
        self.tax_summary_table = QTableWidget()
        self.tax_summary_table.setColumnCount(6)
        self.tax_summary_table.setHorizontalHeaderLabels(["Act", "Period", "Tax", "Interest", "Penalty", "Total"])
        self.tax_summary_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.tax_summary_table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self.tax_summary_table.setStyleSheet("background: white; border: 1px solid #e0e0e0; border-radius: 8px;")
        comp_layout.addWidget(QLabel("<b>Demand Summary:</b>"))
        comp_layout.addWidget(self.tax_summary_table)

        # Step 3 Detailed Breakdown Label
        self.step3_breakdown_lbl = QLabel("")
        self.step3_breakdown_lbl.setStyleSheet("color: #2c3e50; font-weight: bold; font-size: 11pt;")
        comp_layout.addWidget(self.step3_breakdown_lbl)
        comp_layout.addStretch()

        # --- STEP 4: Actions & Finalize ---
        actions_widget = QWidget()
        actions_layout = QVBoxLayout(actions_widget)
        
        self.step4_browser = QWebEngineView()
        self.step4_browser.setMinimumHeight(150)
        self.step4_browser.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        actions_layout.addWidget(self.step4_browser)
        
        # Action Buttons
        btn_row = QHBoxLayout()
        btn_row.addStretch()
        refresh_btn = QPushButton("🔄 Refresh Preview")
        refresh_btn.clicked.connect(self.trigger_drc01a_refresh)
        btn_row.addWidget(refresh_btn)
        
        pdf_btn = QPushButton("⬇️ Draft PDF")
        pdf_btn.clicked.connect(lambda: self.export_drc01a(format="pdf"))
        btn_row.addWidget(pdf_btn)
        
        save_btn = QPushButton("Save Draft")
        save_btn.clicked.connect(self.save_drc01a)
        btn_row.addWidget(save_btn)
        
        finalize_btn = QPushButton("Proceed to Finalize")
        finalize_btn.setStyleSheet("background-color: #27ae60; color: white; font-weight: bold;")
        finalize_btn.clicked.connect(self.show_drc01a_finalization_panel)
        btn_row.addWidget(finalize_btn)
        btn_row.addStretch()
        actions_layout.addLayout(btn_row)

        # Helper for wrapping non-scrollable widgets
        def wrap_scroll(w):
            s = QScrollArea()
            s.setWidgetResizable(True)
            s.setFrameShape(QFrame.Shape.NoFrame)
            s.setWidget(w)
            return s

        # Side Nav Connection
        nav_items = [
            ("Reference Details", "1", wrap_scroll(ref_widget)),
            ("Issues Involved", "2", issues_widget), # Already has inner scroll
            ("Compliance & Summary", "3", wrap_scroll(comp_widget)),
            ("Actions & Finalize", "4", actions_widget) # WebEngine handles scroll
        ]
        side_nav = self.create_side_nav_layout(nav_items, use_scroll=False)
        draft_layout.addWidget(side_nav)

        # Containers
        self.drc01a_view_container = QWidget()
        self.drc01a_view_container.hide()
        view_layout = QVBoxLayout(self.drc01a_view_container)
        view_layout.addWidget(QLabel("<b>DRC-01A Issued</b>"))
        self.drc01a_browser = QWebEngineView()
        view_layout.addWidget(self.drc01a_browser, 1)

        # Final layout connection to ROOT
        layout.addWidget(self.drc01a_draft_container)
        layout.addWidget(self.drc01a_view_container)
        
        self.restore_drc01a_draft_state()
        
        print("ProceedingsWorkspace: create_drc01a_tab done")
        return tab_root

    def load_active_officers(self):
        """Load active officers into the combo box"""
        if getattr(self, 'officer_combo', None):
            self.officer_combo.clear()
            self.officer_combo.addItem("Select Issuing Officer...", None)
            officers = self.db.get_active_officers()
            for off in officers:
                display_text = f"{off['name']} ({off['designation']}, {off['jurisdiction']})"
                self.officer_combo.addItem(display_text, off['id'])
                
            if self.proceeding_data:
                saved_id = self.proceeding_data.get('issuing_officer_id')
                if saved_id:
                    index = self.officer_combo.findData(saved_id)
                    if index >= 0:
                        self.officer_combo.setCurrentIndex(index)

    def create_drc01a_finalization_panel(self):
        print("ProceedingsWorkspace: create_drc01a_finalization_panel start")
        container = QWidget()
        layout = QVBoxLayout(container)
        layout.setContentsMargins(0, 0, 0, 0)
        
        from src.ui.components.finalization_panel import FinalizationPanel
        self.drc_fin_panel = FinalizationPanel()
        
        # Connect Actions
        self.drc_fin_panel.cancel_btn.clicked.connect(lambda: self.toggle_view_mode("drc01a", False))
        self.drc_fin_panel.finalize_btn.clicked.connect(self.finalize_drc01a_issuance)
        
        # Download Actions
        self.drc_fin_panel.pdf_btn.clicked.connect(lambda: self.export_drc01a(format="pdf"))
        self.drc_fin_panel.docx_btn.clicked.connect(lambda: self.export_drc01a(format="docx"))
        
        self.drc_fin_panel.save_btn.clicked.connect(self.save_drc01a)
        self.drc_fin_panel.refresh_btn.clicked.connect(self.show_drc01a_finalization_panel)
        
        layout.addWidget(self.drc_fin_panel)
        self.drc01a_finalization_layout = layout # Keep reference if needed
        return container

    def trigger_drc01a_refresh(self):
        """Debounced preview refresh"""
        self.drc01a_refresh_timer.start(500)

    def _run_refresh_step4(self):
        """Actual preview generation for Step 4"""
        try:
            model = self._get_drc01a_model()
            html = self.render_drc01a_html(model)
            self.step4_browser.setHtml(html)
            
            # Update tax summary table in Step 3
            self._update_drc01a_tax_summary_table(model.get('tax_rows'))
            
        except Exception as e:
            print(f"Error refreshing DRC-01A preview: {e}")
            traceback.print_exc()

    def _update_drc01a_tax_summary_table(self, tax_rows):
        """Populate the Step 3 summary table"""
        if tax_rows is None:
            raise ValueError("tax_rows missing from model during table update.")
            
        self.tax_summary_table.setRowCount(0)
        if not tax_rows: return
        
        from src.utils.formatting import format_indian_number
        
        self.tax_summary_table.setRowCount(len(tax_rows))
        for row_idx, row_data in enumerate(tax_rows):
            # 1. Act & Period
            self.tax_summary_table.setItem(row_idx, 0, QTableWidgetItem(str(row_data.get('act', ''))))
            self.tax_summary_table.setItem(row_idx, 1, QTableWidgetItem(str(row_data.get('period', ''))))
            
            # 2. Monetary Columns
            monetary_cols = [
                (2, 'tax'),
                (3, 'interest'),
                (4, 'penalty'),
                (5, 'total')
            ]
            
            for col_idx, key in monetary_cols:
                val = row_data.get(key, 0)
                formatted_val = format_indian_number(val)
                item = QTableWidgetItem(str(formatted_val))
                item.setTextAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)
                self.tax_summary_table.setItem(row_idx, col_idx, item)

    def refresh_step3_preview(self):
        """Legacy placeholder - redirected to debounced trigger"""
        self.trigger_drc01a_refresh()

    def show_drc01a_finalization_panel(self):
        """Review and Show Finalization Panel"""
        # 1. Validate Issues Exist
        if not self.issue_cards:
             QMessageBox.warning(self, "Draft Empty", "Please add at least one issue before finalizing.")
             return
             
        # 2. Validate Metadata Inputs
        errors = []
        if not self.oc_number_input.text().strip():
            errors.append("OC Number is mandatory.")
        if getattr(self, 'officer_combo', None) and not self.officer_combo.currentData():
            errors.append("Proper Officer selection is mandatory.")
            
        # Date Logic Check (Reply Date > OC Date)
        oc_date = self.oc_date_input.date()
        reply_date = self.reply_deadline_input.date()
        payment_date = self.payment_deadline_input.date()
        
        if reply_date <= oc_date:
            errors.append("Reply Date must be after OC Date.")
        
        if payment_date <= oc_date:
            errors.append("Payment Date must be after OC Date.")
            
        if errors:
            QMessageBox.warning(self, "Validation Error", "\n".join(errors))
            return
            
        try:
            # 3. Build Model
            drc_model = self._get_drc01a_model()
            
            if not drc_model:
                 raise ValueError("Failed to generate data model. Processing data might be missing.")

            # [STRICT] Zero Liability Check
            if drc_model.get('grand_total_val', 0) <= 0:
                QMessageBox.critical(self, "Liability Error", "Cannot issue DRC-01A with Zero Liability.\nPlease add tax amounts to issues.")
                return

            # 4. Render Preview (Single Source)
            html_preview = self.render_drc01a_html(drc_model)
            
            # 5. Load into Panel
            self.drc_fin_panel.load_data(
                proceeding_data=self.proceeding_data,
                issues_list=self.issue_cards, 
                doc_type="DRC-01A",
                doc_no=drc_model['oc_no'],
                doc_date=drc_model['oc_date'],
                ref_no="-" 
            )
            
            # 6. Set HTML Preview (Browser)
            self.drc_fin_panel.set_preview_html(html_preview)

            # 7. Switch View
            self.drc01a_draft_container.hide()
            self.drc01a_view_container.hide()
            self.drc01a_finalization_container.show()
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            QMessageBox.critical(self, "Preview Validation Error", f"Failed to generate preview:\n{str(e)}")
        
    def finalize_drc01a_issuance(self):
        """
        Execute Strict Issuance Logic:
        1. Save Final State to DB
        2. Generate & Save HTML Snapshot (Freeze)
        3. Generate PDF from Snapshot
        4. Update Workflow Stage -> DRC01A_ISSUED
        """
        reply = QMessageBox.question(self, "Confirm Issuance", 
                                   "Are you sure you want to issue DRC-01A?\n\nThis action is IRREVERSIBLE. The document will be locked.",
                                   QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        
        if reply == QMessageBox.StandardButton.No:
            return

        try:
             # 1. Save Final State
             self.save_drc01a() 
             
             # 2. Generate Final HTML Snapshot
             drc_model = self._get_drc01a_model()
             html_content = self.render_drc01a_html(drc_model)
             
             # Create Snapshot Directory
             import os
             base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
             output_dir = os.path.join(base_dir, "data", "generated_documents")
             os.makedirs(output_dir, exist_ok=True)
             
             # Save HTML File
             safe_oc = "".join([c for c in drc_model['oc_no'] if c.isalnum() or c in ('-','_')])
             filename_base = f"DRC-01A_{safe_oc}"
             html_filename = f"{filename_base}.html"
             html_path = os.path.join(output_dir, html_filename)
             
             with open(html_path, 'w', encoding='utf-8') as f:
                 f.write(html_content)
                 
             print(f"Snapshot saved: {html_path}")

             # 3. Generate PDF (From Snapshot Logic)
             from PyQt6.QtPrintSupport import QPrinter
             from PyQt6.QtGui import QTextDocument
             
             pdf_filename = f"{filename_base}.pdf"
             pdf_path = os.path.join(output_dir, pdf_filename)
             
             printer = QPrinter(QPrinter.PrinterMode.HighResolution)
             printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
             printer.setOutputFileName(pdf_path)
             
             doc = QTextDocument()
             doc.setHtml(html_content)
             doc.print(printer)
             
             # 4. Update Workflow State (Atomic Transaction)
             if self.transition_to(WorkflowStage.DRC01A_ISSUED):
                 QMessageBox.information(self, "Success", f"DRC-01A Issued Successfully.\n\nSaved to: {pdf_path}")
                 
                 try:
                    os.startfile(pdf_path)
                 except: pass

                 # 5. Refresh UI to show View Mode (Locked)
                 self.drc01a_finalization_container.hide()
                 self.toggle_view_mode("drc01a", True) 
             else:
                 QMessageBox.critical(self, "Error", "Failed to update workflow stage.")
                 
        except Exception as e:
            QMessageBox.critical(self, "Issuance Error", f"An error occurred: {e}")
            print(f"Finalization Error: {e}")
            import traceback
            traceback.print_exc()


    def _get_drc01a_model(self):
        """
        SSoT: Aggregate all data for DRC-01A rendering and export.
        Uses pure integers for monetary aggregation.
        Aggregates normalized legal references from individual issues.
        """
        if not self.proceeding_data:
            return {
                'tax_rows': [],
                'grand_total_liability': "0"
            }

        tp = self.proceeding_data.get('taxpayer_details', {})
        if isinstance(tp, str):
            try: tp = json.loads(tp)
            except: tp = {}
        elif tp is None:
            tp = {}

        # 1. Taxpayer & Case Info
        model = {
            'gstin': self.proceeding_data.get('gstin', '') or tp.get('GSTIN', ''),
            'legal_name': self.proceeding_data.get('legal_name', '') or tp.get('Legal Name', ''),
            'trade_name': self.proceeding_data.get('trade_name', '') or tp.get('Trade Name', ''),
            'address': self.proceeding_data.get('address', '') or tp.get('Address', ''),
            'case_id': self.proceeding_data.get('case_id', '') or '',
            'oc_no': self.oc_number_input.text(),
            'oc_date': self.oc_date_input.date().toString("dd/MM/yyyy"),
            'financial_year': self.proceeding_data.get('financial_year', '') or '',
            'initiating_section': '',  # placeholder; resolved below
            'form_type': self.proceeding_data.get('form_type', '') or 'DRC-01A',
            'reply_date': self.reply_deadline_input.date().toString("dd/MM/yyyy"),
            'payment_date': self.payment_deadline_input.date().toString("dd/MM/yyyy"),
            'issue_date': self.oc_date_input.date().toString("dd/MM/yyyy")
        }
        
        # Resolved section: prefer adjudication_section over initiating_section
        model['initiating_section'] = (
            self.proceeding_data.get('adjudication_section') or
            self.proceeding_data.get('initiating_section') or ''
        )

        # 2. Strict Adjudication Section Validation with Normalization
        raw_sec = str(model['initiating_section']).strip()
        print(f"[DRC-01A] Section resolve -> raw_sec='{raw_sec}'")
        
        # Normalize common patterns to a bare number string: '73', '74'
        import re
        sec_match = re.search(r'(73|74)', raw_sec)
        if not raw_sec:
            raise ValueError("DRC-01A: adjudication_section is blank. Cannot build model. Ensure section is set during case setup.")
        elif sec_match:
            model['section_base'] = sec_match.group(1)  # '73' or '74'
        else:
            raise ValueError(f"Invalid section_base for DRC-01A derived from: '{raw_sec}'. Must contain '73' or '74'.")
            
        model['section_title'] = f"section {model['section_base']}(5)"
        model['section_body'] = f"{model['section_base']}(5)"

        # 3. Aggregate Legal Provisions (Normalized Layer)
        normalized_provisions = self._aggregate_legal_references()
        if normalized_provisions:
            model['sections_violated_html'] = "<ul>" + "".join([f"<li>{p}</li>" for p in normalized_provisions]) + "</ul>"
        else:
            model['sections_violated_html'] = "<i>Not specified</i>"

        # 4. Narrative Content from Issues
        issues_html = ""
        for card in self.issue_cards:
            if card.is_included:
                issues_html += card.generate_html()
                issues_html += "<br><hr style='border: 1px dashed #eee;'><br>"
        model['issues_html'] = issues_html

        # 5. Financial Aggregation (Decimal Layer)
        # Using Decimal for precise summation across all included issue cards
        total_breakdown = {} # Act -> {Tax: D, Interest: D, Penalty: D, Total: D}
        
        # 5. Financial Aggregation (Integer Layer)
        # Standardizing on Integers for precision and consistency across system
        total_breakdown = {} # Act -> {Tax: int, Interest: int, Penalty: int, Total: int}
        
        for card in self.issue_cards:
            if not card.is_included: continue
            
            card_breakdown = card.get_tax_breakdown() # Returns card level integer mapping
            for act, vals in card_breakdown.items():
                if act not in total_breakdown:
                    total_breakdown[act] = {
                        'Tax': 0, 'Interest': 0, 
                        'Penalty': 0, 'Total': 0
                    }
                
                # Aggregate values safely ensuring integer types
                for key in ['tax', 'interest', 'penalty']:
                    display_key = key.capitalize() # 'tax' -> 'Tax'
                    v = vals.get(key, 0)
                    total_breakdown[act][display_key] += v
                    total_breakdown[act]['Total'] += v

        # Prepare Tax Table Rows
        tax_rows = []
        grand_total = 0
        
        # Sort by ACT_PRIORITY
        for act in self.ACT_PRIORITY:
            if act in total_breakdown:
                vals = total_breakdown[act]
                row_tot = vals['Total']
                grand_total += row_tot
                
                tax_rows.append({
                    'act': act,
                    'period': model['financial_year'], # Summary level uses FY
                    'tax': vals['Tax'],
                    'interest': vals['Interest'],
                    'penalty': vals['Penalty'],
                    'total': row_tot
                })
        
        model['tax_rows'] = tax_rows
        model['grand_total_liability'] = grand_total
        
        # Period derivation for metadata
        model['tax_period_from'] = f"01/04/{model['financial_year'][:4]}" if model['financial_year'] else ""
        model['tax_period_to'] = f"31/03/{model['financial_year'][5:]}" if len(model['financial_year']) > 5 else ""
        
        # 6. Strict Advice Paragraph Generation
        payment_date = model['payment_date']
        base = model['section_base']
        
        if base == "73":
            model['advice_paragraph'] = f"You are hereby advised to pay the amount of tax as ascertained above alongwith the amount of applicable interest in full by {payment_date}, failing which Show Cause Notice will be issued under section 73(1)."
        elif base == "74":
            model['advice_paragraph'] = f"You are hereby advised to pay the amount of tax as ascertained above alongwith the amount of applicable interest and penalty under section 74(5) by {payment_date}, failing which Show Cause Notice will be issued under section 74(1)."

        # HTML Table String Injection
        model['tax_table_html'] = self.generate_tax_table_html(tax_rows) 
        
        # 7. Extract Officer Data (Deterministic Snapshot Logic)
        import json
        snapshot_json = self.proceeding_data.get('issuing_officer_snapshot')
        officer_data = {}
        if snapshot_json:
            try:
                parsed_snap = json.loads(snapshot_json)
                if 'DRC-01A' in parsed_snap:
                    officer_data = parsed_snap['DRC-01A']
                elif 'name' in parsed_snap:
                    officer_data = parsed_snap # Fallback for old flat format
            except json.JSONDecodeError:
                pass
        
        if not officer_data:
            officer_id = self.proceeding_data.get('issuing_officer_id') or (getattr(self, 'officer_combo', None) and self.officer_combo.currentData())
            officer_data = self.db.get_officer_by_id(officer_id) if officer_id else {}

        model['officer_name'] = officer_data.get('name', 'Proper Officer')
        model['designation'] = officer_data.get('designation', 'Superintendent')
        model['jurisdiction'] = officer_data.get('jurisdiction', 'Paravur Range')
        
        return model

    def _aggregate_legal_references(self):
        """
        Aggregate and normalize legal provisions from all issue cards.
        [DEPRECATED] Legal provisions are now consolidated into the 'Brief Facts' narrative.
        """
        return []

    def validate_drc01a_model(self, model):
        """
        Validation Layer for model integrity.
        Returns (bool, msg)
        """
        if not model:
            return False, "Data model is empty."
        
        required = {
            'gstin': "GSTIN",
            'legal_name': "Legal Name",
            'oc_no': "OC Number",
            'oc_date': "OC Date",
            'initiating_section': "Adjudication Section"
        }
        
        for key, label in required.items():
            if not model.get(key) or model.get(key) == "____":
                return False, f"Mandatory field '{label}' is missing or incomplete."

        # Logical checks
        if not model.get('tax_rows'):
            return False, "At least one row of tax quantification is required."

        # Date consistency (Basic check)
        try:
            d_oc = datetime.datetime.strptime(model['oc_date'], "%d/%m/%Y") # Corrected format string
            d_reply = datetime.datetime.strptime(model['reply_date'], "%d/%m/%Y") # Corrected format string
            if d_reply < d_oc:
                return False, "Reply date cannot be earlier than OC date."
        except ValueError: # Catch specific ValueError for strptime
            pass # Formats might differ or be partial

        return True, "Success"

    def render_drc01a_html(self, model):
        """
        PURE FUNCTION: Render HTML using the data model.
        No UI reads allowed.
        """
        if not model:
            return "<h3>Data Validation Failed</h3>"

        try:
            from jinja2 import Environment, FileSystemLoader
            import os
            import datetime # Import datetime for potential use in template or context

            # Locate Template
            current_dir = os.path.dirname(os.path.abspath(__file__))
            # Work up to src/templates
            import os

            # Locate Template
            current_dir = os.path.dirname(os.path.abspath(__file__)) # src/ui
            src_dir = os.path.dirname(current_dir) # src
            root_dir = os.path.dirname(src_dir) # gst (Project Root)
            template_dir = os.path.join(root_dir, "templates")

            print(f"DEBUG: Loading Templates from Root: {template_dir}")
            
            target_template = os.path.join(template_dir, 'drc01a.html')
            if not os.path.exists(target_template):
                 print(f"CRITICAL: drc01a.html NOT FOUND in {template_dir}")
                 # Fallback for dev environment oddities
                 if os.path.exists("D:/gst/templates/drc01a.html"):
                     template_dir = "D:/gst/templates"
                     target_template = "D:/gst/templates/drc01a.html"

            print(f"DEBUG: Loading DRC-01A via TemplateEngine")

            from src.utils.template_engine import TemplateEngine
            from src.utils.formatting import format_indian_number
            
            # Defense against None type aggregation
            total = model.get('grand_total_liability', 0)
            formatted_total = format_indian_number(total) if total is not None else "0"
            
            import re
            
            # --- 2. Backend String Sanitization ---
            advice_paragraph = model.get('advice_paragraph', '')
            advice_paragraph = re.sub(r'\s+', ' ', advice_paragraph).strip()
            
            raw_issues_html = model.get('issues_html', '')
            
            # --- 3. Structural Cleanup of issues_html ---
            sanitized_issues = re.sub(r'\s+', ' ', raw_issues_html).strip()
            
            # Remove inline styles ONLY from <p> and <span> to preserve <table> borders and grid styling
            sanitized_issues = re.sub(r'(<(?:p|span)[^>]*?)\s+style="[^"]*"', r'\1', sanitized_issues, flags=re.IGNORECASE)
            
            # Remove empty paragraphs
            sanitized_issues = re.sub(r'<p>\s*</p>', '', sanitized_issues)
            
            # Collapse multiple <br>
            sanitized_issues = re.sub(r'(<br\s*/?>\s*){2,}', '<br>', sanitized_issues)
            
            # Trim outer whitespace
            sanitized_issues = sanitized_issues.strip()
            
            context = {
                'gstin': model['gstin'],
                'legal_name': model['legal_name'],
                'trade_name': f"({model['trade_name']})" if model.get('trade_name') and model['trade_name'] != model['legal_name'] else "",
                'address': model['address'],
                'case_id': model['case_id'],
                'oc_no': model['oc_no'],
                'oc_date': model['oc_date'],
                'financial_year': model['financial_year'],
                'form_type': model['form_type'],
                'section_base': model['section_base'],
                'tax_table_html': model.get('tax_table_html', ''),
                'issues_html': sanitized_issues,
                'grand_total': formatted_total,
                'advice_paragraph': advice_paragraph,
                'last_date_reply': model['reply_date'],
                'officer_name': model.get('officer_name', 'Proper Officer'),
                'designation': model.get('designation', 'Superintendent'),
                'jurisdiction': model.get('jurisdiction', 'Paravur Range')
            }
            
            # 8. Strict Contract Validation
            required_keys = [
                'gstin', 'legal_name', 'trade_name', 'address', 'case_id', 'oc_no', 'oc_date', 
                'financial_year', 'form_type', 'section_base', 'tax_table_html', 'issues_html', 
                'grand_total', 'advice_paragraph', 'last_date_reply', 'officer_name', 'designation', 'jurisdiction'
            ]
            
            for key in required_keys:
                if key not in context or context[key] is None:
                    raise KeyError(f"CRITICAL: Missing required contract key for DRC-01A rendering: {key}")

            # Letterhead injection (if needed, based on a model flag or config)
            # For now, assume template handles placeholder or it's done post-render if needed.
            # If `self.show_letterhead_cb` is still used, it needs to be passed via model.
            # For this pure function, we'll assume the model has a 'show_letterhead' flag.
            # If not, the template itself should handle the placeholder.
            # For now, we'll remove the letterhead logic from here and assume the template is self-contained.

            rendered = template.render(**context)
            print("--- HTML OUTPUT START (First 200 chars) ---")
            print(rendered[:200])
            print("--- HTML OUTPUT END ---")
            print(f"DEBUG: Rendered HTML Length: {len(rendered)}")
            return rendered
            
        except Exception as e:
            print(f"Error rendering DRC-01A: {e}\n{traceback.format_exc()}")
            return f"<h3>Rendering Error: {str(e)}</h3>"

    def export_drc01a(self, format="pdf"):
        """
        Synchronous export logic with validation and error handling.
        """
        from PyQt6.QtWidgets import QFileDialog, QMessageBox
        from PyQt6.QtGui import QCursor
        from PyQt6.QtCore import Qt
        import os
        import traceback
        from src.utils.document_generator import DocumentGenerator # Assuming this exists
        from src.utils.preview_generator import PreviewGenerator # Assuming this exists

        # 1. Regenerate model & Validate
        model = self._get_drc01a_model()
        valid, msg = self.validate_drc01a_model(model)
        if not valid:
            QMessageBox.warning(self, "Export Validation", msg)
            return

        # 2. File Dialog
        filter_str = "PDF Files (*.pdf)" if format == "pdf" else "Word Files (*.docx)"
        default_file = f"DRC-01A_{model['legal_name'].replace(' ', '_')}_{model['oc_no'].replace('/', '_')}"
        path, _ = QFileDialog.getSaveFileName(self, f"Save Draft {format.upper()}", default_file, filter_str)
        
        if not path:
            return # Cancelled silently

        # 3. Execution with UI blocking
        self.setCursor(QCursor(Qt.CursorShape.WaitCursor))
        try:
            if format == "pdf":
                html = self.render_drc01a_html(model)
                success, err_msg = PreviewGenerator.generate_pdf(html, path)
                if not success:
                    raise RuntimeError(err_msg)
            else:
                # Word Export (SSoT Model -> DocumentGenerator)
                generator = DocumentGenerator()
                word_path = generator.generate_word({
                    'form_type': 'DRC-01A',
                    'date': model['oc_date'],
                    'gstin': model['gstin'],
                    'legal_name': model['legal_name'],
                    'address': model['address'],
                    'proceeding_type': f"Intimation under {model['initiating_section']}",
                    'facts': f"Issue: {model['initiating_section']}\n(Note: Full narrative is in the preview)",
                    'tax_data': model['tax_rows']
                }, os.path.basename(path).replace(".docx", ""))
                
                # Move to actual path if generator used OUTPUT_DIR
                if word_path and os.path.exists(word_path) and word_path != path:
                    import shutil
                    shutil.move(word_path, path)

            QMessageBox.information(self, "Success", f"Draft {format.upper()} exported successfully to:\n{path}")
            
        except Exception as e:
            print(f"Export Error: {e}\n{traceback.format_exc()}")
            QMessageBox.critical(self, "Export Error", f"Failed to export {format.upper()}:\n{str(e)}")
        finally:
            self.unsetCursor()

    def generate_drc01a_html(self):
        """Legacy ref to renderer (Keep for compatibility during transition)"""
        model = self._get_drc01a_model()
        return self.render_drc01a_html(model)
    def load_sections(self):
        """Populate Adjudication Sections"""
        if not hasattr(self, 'section_combo'): return
        
        current_val = self.section_combo.currentText()
        self.section_combo.blockSignals(True)
        self.section_combo.clear()
        
        sections = [
            "Section 73: Non-fraud cases",
            "Section 74: Fraud/Suppression cases", 
            "Section 76: Tax collected but not paid",
            "Section 122: General Penalties",
            "Section 125: General Penalty",
            "Section 129: E-Way Bill / Detention",
            "Section 130: Confiscation"
        ]
        
        # Add current value if not in list (Legacy Support)
        if current_val and current_val not in sections and current_val != "Select Section...":
             self.section_combo.addItem(current_val)
             
        self.section_combo.addItems(sections)
        
        # Try to match proceeding data
        proc_sec = self.proceeding_data.get('adjudication_section', '')
        if proc_sec:
            # Normalize lookup (e.g. "73" -> match "Section 73...")
            for idx in range(self.section_combo.count()):
                text = self.section_combo.itemText(idx)
                if proc_sec in text:
                    self.section_combo.setCurrentIndex(idx)
                    break
        self.section_combo.blockSignals(False)

    def load_issue_templates(self):
        """Load issue templates from DB (Active Issues)"""
        try:
            # Fetch active issues from DB
            issues = self.db.get_active_issues()
            
            # Convert list to dict keyed by issue_id
            self.issue_templates = {}
            self.issue_combo.clear()
            self.issue_combo.addItem("Select Issue...", None)
            
            for issue in issues:
                self.issue_templates[issue['issue_id']] = issue
                self.issue_combo.addItem(issue['issue_name'], issue['issue_id'])
                
        except Exception as e:
            print(f"Error loading issue templates: {e}")



    def add_section_to_editor(self):
        """Insert selected section into the editor"""
        section_data = self.section_combo.currentData()
        if not section_data:
            return
            
        title = section_data.get('title', '')
        # content = section_data.get('content', '') # Optional: Add content too?
        
        # Append to editor
        current_html = self.sections_editor.toHtml()
        new_html = f"{current_html}<p><b>{title}</b></p>"
        self.sections_editor.setHtml(new_html)

    def insert_selected_issue(self):
        """Insert selected issue template as a new card"""
        issue_id = self.issue_combo.currentData()
        if not issue_id:
            return
            
        template = self.issue_templates.get(issue_id)
        if not template:
            return
            
        from src.ui.issue_card import IssueCard
        card = IssueCard(template, mode="DRC-01A")
        
        # Connect signals
        card.removeClicked.connect(lambda: self.remove_issue_card(card))
        card.valuesChanged.connect(self.calculate_grand_totals)
        
        self.issues_layout.addWidget(card)
        self.issue_cards.append(card)
        
        # Trigger initial calculation and preview
        card.calculate_values()
        self.calculate_grand_totals()
        

    def remove_issue_card(self, card):
        self.issues_layout.removeWidget(card)
        card.deleteLater()
        if card in self.issue_cards:
            self.issue_cards.remove(card)
        self.calculate_grand_totals()

    def calculate_grand_totals(self, _=None):
        """Dynamic Aggregation for DRC-01A Draft (Derived from Issues) - Enforced Integers"""
        total_tax = 0
        total_interest = 0
        total_penalty = 0
        
        # Aggregate from all issue cards
        for card in self.issue_cards:
            if not card.is_included: continue
            
            breakdown = card.get_tax_breakdown()
            for values in breakdown.values():
               total_tax += values.get('tax', 0)
               total_interest += values.get('interest', 0)
               total_penalty += values.get('penalty', 0)
        
        grand_total = total_tax + total_interest + total_penalty
        
        # Update UI Elements
        from src.utils.formatting import format_indian_number
        
        # Step 2: Total Liability Label
        if hasattr(self, 'lbl_total_tax'):
            self.lbl_total_tax.setText(f"₹ {format_indian_number(grand_total)}")
            
        # Step 3: Global Breakdown Label
        if hasattr(self, 'step3_breakdown_lbl'):
             summary_text = (f"Tax: ₹ {format_indian_number(total_tax)} | "
                            f"Interest: ₹ {format_indian_number(total_interest)} | "
                            f"Penalty: ₹ {format_indian_number(total_penalty)}")
             self.step3_breakdown_lbl.setText(summary_text)

        # Step 3: Refresh the summary table
        # Guard: skip model call if section is not resolved yet (avoids ValueError)
        resolved = self.proceeding_data.get('adjudication_section') or self.proceeding_data.get('initiating_section', '')
        if not resolved or not str(resolved).strip():
            print("[DRC-01A] calculate_grand_totals: skipping tax table refresh — adjudication_section not yet set.")
            return
        try:
            tax_rows = self._get_drc01a_model().get('tax_rows', [])
            self._update_drc01a_tax_summary_table(tax_rows)
        except ValueError as e:
            print(f"[DRC-01A] calculate_grand_totals: suppressed model error: {e}")


    def create_scn_tab(self):
        print("ProceedingsWorkspace: create_scn_tab start")
        """Create Show Cause Notice tab"""
        
        # --- DRAFT CONTAINER ---
        self.scn_draft_container = QWidget()
        self.scn_draft_container.setObjectName("SCNDraftContainer")
        draft_layout = QVBoxLayout(self.scn_draft_container)
        draft_layout.setContentsMargins(0, 0, 0, 0)
        
        # Track initialization state for Issue Adoption
        self.scn_issues_initialized = False
        
        # Preparing Pages
        
        # 1. Reference Details (Reverted to Single-Column centered layout)
        ref_widget = QWidget()
        ref_layout = QVBoxLayout(ref_widget)
        ref_layout.setContentsMargins(20, 20, 20, 20)
        ref_layout.setSpacing(20)
        
        self.ref_card = ModernCard(title="Reference Details", collapsible=False)
        self.ref_card.setFixedWidth(500) # Slightly wider for drafting comfort
        
        # Section 1 – Notice Identification
        header_style = "font-size: 11pt; font-weight: bold; color: #2c3e50;"
        sec1_header = QLabel("Notice Identification")
        sec1_header.setStyleSheet(header_style)
        self.ref_card.addLayout(QVBoxLayout()) # Internal layout access
        self.ref_card.content_layout.addWidget(sec1_header)
        
        # Grid for inputs
        grid_widget = QWidget()
        grid = QGridLayout(grid_widget)
        grid.setContentsMargins(0, 5, 0, 0)
        grid.setSpacing(12)
        
        label_style = "color: #5f6368; font-weight: 500; font-size: 9pt;"
        input_style = "padding: 8px; border: 1px solid #dadce0; border-radius: 4px; font-size: 10pt;"
        
        # SCN No
        scn_no_label = QLabel("SCN No.")
        scn_no_label.setStyleSheet(label_style)
        self.scn_no_input = QLineEdit()
        self.scn_no_input.setPlaceholderText("e.g. SCN/2026/001")
        self.scn_no_input.setStyleSheet(input_style)
        self.scn_no_input.textChanged.connect(self.evaluate_scn_workflow_phase)
        
        grid.addWidget(scn_no_label, 0, 0)
        grid.addWidget(self.scn_no_input, 0, 1, 1, 2)
        
        # O.C. No (with Inline Auto-Generate)
        oc_label = QLabel("O.C. No.")
        oc_label.setStyleSheet(label_style)
        self.scn_oc_input = QLineEdit()
        self.scn_oc_input.setPlaceholderText("e.g. 123/2026")
        self.scn_oc_input.setStyleSheet(input_style)
        self.scn_oc_input.textChanged.connect(self._on_scn_oc_changed)
        
        # Auto-Generate Button (Utility)
        self.btn_auto_oc = QPushButton("Auto-Generate")
        self.btn_auto_oc.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_auto_oc.setStyleSheet("""
            QPushButton {
                background-color: #f8fafc;
                border: 1px solid #cbd5e1;
                color: #475569;
                padding: 6px 10px;
                border-radius: 4px;
                font-size: 8pt;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #f1f5f9;
                border-color: #3498db;
                color: #3498db;
            }
        """)
        self.btn_auto_oc.clicked.connect(self._on_auto_generate_oc)
        
        grid.addWidget(oc_label, 1, 0)
        grid.addWidget(self.scn_oc_input, 1, 1)
        grid.addWidget(self.btn_auto_oc, 1, 2)
        
        # SCN Date
        date_label = QLabel("SCN Date")
        date_label.setStyleSheet(label_style)
        self.scn_date_input = QDateEdit()
        self.scn_date_input.setCalendarPopup(True)
        self.scn_date_input.setDate(QDate.currentDate())
        self.scn_date_input.setMinimumDate(QDate.currentDate())
        self.scn_date_input.setStyleSheet(input_style + " padding: 6px;")
        self.scn_date_input.dateChanged.connect(self.evaluate_scn_workflow_phase)
        
        grid.addWidget(date_label, 2, 0)
        grid.addWidget(self.scn_date_input, 2, 1, 1, 2)
        
        # [NEW] Issuing Officer Selection
        officer_label = QLabel("Issuing Officer")
        officer_label.setStyleSheet(label_style)
        self.scn_officer_combo = QComboBox()
        self.scn_officer_combo.setStyleSheet(input_style + " padding: 6px;")
        
        # Load active officers into SCN combo box specifically
        self.scn_officer_combo.addItem("Select Issuing Officer...", None)
        officers = self.db.get_active_officers()
        
        # Hydrate default value natively
        current_officer_id = self.proceeding_data.get('issuing_officer_id')
        default_index = 0
        
        for idx, off in enumerate(officers, start=1):
            display_text = f"{off['name']} ({off['designation']}, {off['jurisdiction']})"
            self.scn_officer_combo.addItem(display_text, off['id'])
            if off['id'] == current_officer_id:
                default_index = idx
                
        if default_index > 0:
            self.scn_officer_combo.setCurrentIndex(default_index)
            
        grid.addWidget(officer_label, 3, 0)
        grid.addWidget(self.scn_officer_combo, 3, 1, 1, 2)

        # Spacing to Grounds Module
        self.ref_card.content_layout.addWidget(grid_widget)
        
        # [NEW] Grounds Configuration Module
        # Identity Logic (Provenance Label moved here)
        self.oc_provenance_lbl = QLabel("")
        self.oc_provenance_lbl.setStyleSheet("font-size: 8pt; color: #3498db; font-style: italic; margin-bottom: 10px;")
        self.oc_provenance_lbl.hide()
        self.ref_card.content_layout.addWidget(self.oc_provenance_lbl)

        # Instantiate Form
        from src.ui.components.grounds_forms import get_grounds_form
        self.scn_grounds_form = get_grounds_form("scrutiny")
        self.ref_card.content_layout.addWidget(self.scn_grounds_form)
        
        self.ref_card.content_layout.addStretch()
        
        # Action Buttons Hierarchy
        btn_layout = QVBoxLayout()
        btn_layout.setSpacing(10)
        btn_layout.setContentsMargins(0, 20, 0, 0)
        
        # Save & Proceed (Primary)
        self.btn_save_scn_ref = QPushButton("Save & Proceed")
        self.btn_save_scn_ref.setStyleSheet("""
            QPushButton {
                background-color: #1a73e8;
                color: white;
                padding: 12px;
                border-radius: 6px;
                font-weight: bold;
                font-size: 10pt;
            }
            QPushButton:hover {
                background-color: #1557b0;
            }
            QPushButton:disabled {
                background-color: #dadce0;
                color: #70757a;
            }
        """)
        self.btn_save_scn_ref.clicked.connect(self.save_scn_metadata)
        
        # [PHASE B] Standardized Draft Save (Step 1)
        self.btn_draft_step1 = self._create_save_draft_btn()
        
        # Footer layout for buttons
        footer_btn_layout = QHBoxLayout()
        footer_btn_layout.addWidget(self.btn_draft_step1)
        footer_btn_layout.addWidget(self.btn_save_scn_ref)
        btn_layout.addLayout(footer_btn_layout)
        
        # Action Buttons Hierarchy (End)
        
        self.ref_card.addLayout(btn_layout)
        
        # Center the card in the layout
        ref_layout.addStretch()
        ref_layout.addWidget(self.ref_card, 0, Qt.AlignmentFlag.AlignHCenter)
        ref_layout.addStretch()

        # 2. Issues Involved
        issues_widget = QWidget()
        issues_layout = QVBoxLayout(issues_widget)
        issues_layout.setContentsMargins(0,0,0,0)
        
        self.scn_issue_cards = []
        
        scn_issue_selection_layout = QHBoxLayout()
        scn_issue_label = QLabel("Select Issue:")
        self.scn_issue_combo = QComboBox()
        self.scn_issue_combo.addItem("Select an issue...", None)
        print("ProceedingsWorkspace: loading scn issue templates")
        self.load_scn_issue_templates()
        print("ProceedingsWorkspace: scn issue templates loaded")
        
        scn_refresh_issues_btn = QPushButton("🔄")
        scn_refresh_issues_btn.setToolTip("Refresh issue list")
        scn_refresh_issues_btn.clicked.connect(self.load_scn_issue_templates)
        
        scn_insert_issue_btn = QPushButton("Insert Issue Template")
        scn_insert_issue_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        scn_insert_issue_btn.setStyleSheet("""
            QPushButton {
                background-color: #1a73e8; 
                color: white; 
                font-weight: bold; 
                padding: 6px 15px; 
                border-radius: 4px;
                font-size: 9pt;
            }
            QPushButton:hover {
                background-color: #1557b0;
            }
        """)
        scn_insert_issue_btn.clicked.connect(self.insert_scn_issue)
        
        scn_reset_btn = QPushButton("Reset Adoption")
        scn_reset_btn.setStyleSheet("background-color: #f39c12; color: white; border-radius: 4px; padding: 5px 10px; font-size: 8pt; font-weight: bold;")
        scn_reset_btn.clicked.connect(self.reset_scn_adoption)
        
        scn_issue_selection_layout.addWidget(scn_issue_label)
        scn_issue_selection_layout.addWidget(self.scn_issue_combo)
        scn_issue_selection_layout.addWidget(scn_refresh_issues_btn)
        scn_issue_selection_layout.addWidget(scn_insert_issue_btn)
        scn_issue_selection_layout.addWidget(scn_reset_btn)
        
        # [PHASE B] Standardized Draft Save (Step 2)
        btn_draft_step2 = self._create_save_draft_btn()
        scn_issue_selection_layout.addWidget(btn_draft_step2)
        
        scn_issue_selection_layout.addStretch()
        issues_layout.addLayout(scn_issue_selection_layout)
        
        # Proper SCN container
        self.scn_issues_container = QWidget()
        self.scn_issues_layout = QVBoxLayout(self.scn_issues_container)
        self.scn_issues_layout.setAlignment(Qt.AlignmentFlag.AlignTop)
        self.scn_issues_layout.setSpacing(15)
        
        issues_layout.addWidget(self.scn_issues_container)
        issues_layout.addStretch()

        # 3. Demand & Contraventions
        demand_widget = QWidget()
        demand_layout = QVBoxLayout(demand_widget)
        demand_layout.setContentsMargins(0,0,0,0)
        
        demand_header_layout = QHBoxLayout()
        demand_header_layout.addStretch()
        regenerate_btn = QPushButton("Sync with Issues")
        regenerate_btn.setToolTip("Auto-generate demand text tiles based on added issues")
        regenerate_btn.setStyleSheet("padding: 5px; font-size: 8pt; background-color: #3498db; color: white; border-radius: 4px;")
        regenerate_btn.clicked.connect(lambda: self.sync_demand_tiles() if not self.is_scn_phase1() else None)
        
        # [PHASE B] Standardized Draft Save (Step 3)
        btn_draft_step3 = self._create_save_draft_btn()
        demand_header_layout.addWidget(btn_draft_step3)
        
        demand_header_layout.addWidget(regenerate_btn)
        demand_layout.addLayout(demand_header_layout)
        
        self.demand_tiles_widget = QWidget()
        self.demand_tiles_layout = QVBoxLayout(self.demand_tiles_widget)
        self.demand_tiles_layout.setSpacing(10)
        self.demand_tiles_layout.setContentsMargins(0, 0, 0, 0)
        
        self.demand_tiles = []
        
        demand_layout.addWidget(self.demand_tiles_widget)
        demand_layout.addStretch()
        
        # 4. Reliance Placed
        reliance_widget = QWidget()
        rel_layout = QVBoxLayout(reliance_widget)
        rel_layout.setContentsMargins(0,0,0,0)
        
        self.reliance_editor = RichTextEditor("List documents here (e.g., 1. INS-01 dated...)")
        self.reliance_editor.setMinimumHeight(120)
        

        rel_header = QHBoxLayout()
        rel_header.addStretch()
        # [PHASE B] Standardized Draft Save (Step 4)
        btn_draft_step4 = self._create_save_draft_btn()
        rel_header.addWidget(btn_draft_step4)
        rel_layout.addLayout(rel_header)

        rel_layout.addWidget(self.reliance_editor)
        
        # 5. Copy Submitted To
        copy_widget = QWidget()
        copy_layout = QVBoxLayout(copy_widget)
        copy_layout.setContentsMargins(0,0,0,0)
        
        self.copy_to_editor = RichTextEditor("List authorities here...")
        self.copy_to_editor.setMinimumHeight(120)
        

        copy_header = QHBoxLayout()
        copy_header.addStretch()
        # [PHASE B] Standardized Draft Save (Step 5)
        btn_draft_step5 = self._create_save_draft_btn()
        copy_header.addWidget(btn_draft_step5)
        copy_layout.addLayout(copy_header)

        copy_layout.addWidget(self.copy_to_editor)
        
        # 6. Finalization & Preview (Deterministic Step 6)
        self.scn_finalization_container = self.create_scn_finalization_panel()

        # Helper for wrapping non-scrollable widgets
        def wrap_scroll(w):
            s = QScrollArea()
            s.setWidgetResizable(True)
            s.setFrameShape(QFrame.Shape.NoFrame)
            s.setWidget(w)
            return s

        # Build Side Nav
        nav_items = [
            ("Reference Details", "1", wrap_scroll(ref_widget)),
            ("Issue Adoption", "2", issues_widget),
            ("Demand & Contraventions", "3", wrap_scroll(demand_widget)),
            ("Reliance Placed", "4", reliance_widget), # Editor handles scroll
            ("Copy Submitted To", "5", copy_widget), # Editor handles scroll
            ("Actions & Finalize", "✓", self.scn_finalization_container)
        ]
        
        self.scn_side_nav = self.create_side_nav_layout(nav_items, page_changed_callback=self.on_scn_page_changed, use_scroll=False)
        # Store for access in validation
        self.scn_nav_cards = self.scn_side_nav.nav_cards
        
        draft_layout.addWidget(self.scn_side_nav)
        
        # Lock Steps 2-6 initially
        for i in range(1, 6):
            self.scn_nav_cards[i].set_enabled(False)
        
        # --- VIEW CONTAINER ---
        self.scn_view_container = QWidget()
        self.scn_view_container.hide()
        view_layout = QVBoxLayout(self.scn_view_container)
        
        view_title = QLabel("<b>SCN Generated</b>")
        view_title.setStyleSheet("font-size: 14pt; color: #27ae60; margin-bottom: 20px;")
        view_title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        view_layout.addWidget(view_title)
        
        view_msg = QLabel("A Show Cause Notice has already been generated and issued for this case.")
        view_msg.setAlignment(Qt.AlignmentFlag.AlignCenter)
        view_layout.addWidget(view_msg)
        
        summary_lbl = QLabel("Document Generated Successfully.\nClick 'Edit / Revise Draft' to make changes.")
        summary_lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        summary_lbl.setStyleSheet("color: #7f8c8d; font-size: 10pt; margin: 20px;")
        view_layout.addWidget(summary_lbl)
        
        edit_btn = QPushButton("Edit / Revise Draft")
        edit_btn.setStyleSheet("background-color: #f39c12; color: white; padding: 10px; font-weight: bold;")
        edit_btn.clicked.connect(lambda: self.toggle_view_mode("scn", False))
        view_layout.addWidget(edit_btn)
        
        view_layout.addStretch()

        # Wrapper
        wrapper = QWidget()
        wrapper_layout = QVBoxLayout(wrapper)
        wrapper_layout.setContentsMargins(0,0,0,0)
        wrapper_layout.addWidget(self.scn_draft_container)
        wrapper_layout.addWidget(self.scn_view_container)
        
        print("ProceedingsWorkspace: create_scn_tab done")
        return wrapper

    def validate_scn_metadata(self) -> bool:
        """Returns True only if SCN No, OC No, and Date are present"""
        if not hasattr(self, 'scn_no_input'):
            return False
            
        scn_no = self.scn_no_input.text().strip()
        oc_no = self.scn_oc_input.text().strip()
        scn_date = self.scn_date_input.date()
        
        # Valid if SCN No and OC No are not empty, and date is valid
        return bool(scn_no) and bool(oc_no) and scn_date.isValid()

    def save_scn_metadata(self):
        """Authoritative metadata persistence for Step 1. Metadata is always saved, even if incomplete."""
        if not self.proceeding_id:
            return
            
        try:
            # 1. Extract values
            scn_no = self.scn_no_input.text().strip()
            oc_no = self.scn_oc_input.text().strip()
            scn_date = self.scn_date_input.date().toString("yyyy-MM-dd")
            
            # [NEW] Extract and Validate Grounds Data
            grounds_data = None
            if hasattr(self, 'scn_grounds_form'):
                grounds_data = self.scn_grounds_form.get_data()
            
            # 2. Update additional_details without touching other sections
            current_details = self.proceeding_data.get('additional_details', {})
            if isinstance(current_details, str):
                try: current_details = json.loads(current_details)
                except: current_details = {}
            
            current_details.update({
                "scn_number": scn_no,
                "scn_oc_number": oc_no,
                "scn_date": scn_date
            })
            
            if grounds_data:
                current_details['scn_grounds'] = grounds_data
            
            # [NEW] Establish Snapshot Payload
            officer_id = None
            officer_snapshot = None
            if hasattr(self, 'scn_officer_combo'):
                officer_id = self.scn_officer_combo.currentData()
                if officer_id:
                    officer_data = self.db.get_officer_by_id(officer_id)
                    if officer_data:
                        import json
                        existing_snap_str = self.proceeding_data.get('issuing_officer_snapshot')
                        existing_snap = {}
                        if existing_snap_str:
                            try:
                                existing_snap = json.loads(existing_snap_str)
                                if 'name' in existing_snap and 'SCN' not in existing_snap and 'DRC-01A' not in existing_snap:
                                    existing_snap = {'DRC-01A': existing_snap}
                            except:
                                existing_snap = {}
                        
                        existing_snap['SCN'] = officer_data
                        officer_snapshot = json.dumps(existing_snap)

            # 3. Persist to DB (Routed Logic)
            if self.proceeding_data.get('is_adjudication'):
                # Adjudication Case (Table: adjudication_cases)
                success = self.db.update_adjudication_case(self.proceeding_id, {
                    "additional_details": current_details,
                    "issuing_officer_id": officer_id,
                    "issuing_officer_snapshot": officer_snapshot
                })
            else:
                # Standard Proceeding (Table: proceedings)
                success = self.db.update_proceeding(self.proceeding_id, {
                    "additional_details": current_details,
                    "issuing_officer_id": officer_id,
                    "issuing_officer_snapshot": officer_snapshot
                })
            
            # 4. Update local state
            self.proceeding_data['additional_details'] = current_details
            self.proceeding_data['issuing_officer_id'] = officer_id
            self.proceeding_data['issuing_officer_snapshot'] = officer_snapshot
            
            # 5. Evaluate phase (Enables Step 2 if metadata is now valid)
            self.evaluate_scn_workflow_phase()
            
            # 6. Safeguard: Navigation only if fully valid
            if self.validate_scn_metadata():
                # Navigate to Step 2 (Issue Adoption)
                if hasattr(self, 'scn_side_nav') and len(self.scn_nav_cards) > 1:
                    # trigger navigation by emitting clicked signal
                    self.scn_side_nav.nav_cards[1].clicked.emit(1)
            else:
                QMessageBox.information(self, "Saved", "Reference details saved. Complete all fields to proceed to Issue Adoption.")
                
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to save metadata: {e}")
            print(f"Error saving SCN metadata: {e}")

    def evaluate_scn_workflow_phase(self):
        """Authoritative, data-driven state machine for SCN drafting phases."""
        if not hasattr(self, 'scn_nav_cards'):
            return

        metadata_valid = self.validate_scn_metadata()
        
        # Phase 2 Condition: Metadata OK + At least one issue adopted/added in Step 2
        content_valid = any(issue.is_adopted for issue in self.scn_issue_cards)
        
        # Update Authority Flag
        prev_phase = self.scn_workflow_phase
        self.scn_workflow_phase = "DRAFTING" if (metadata_valid and content_valid) else "METADATA"
            
        print(f"SCN Workflow State: metadata={metadata_valid}, content={content_valid} -> {self.scn_workflow_phase}")

        # --- UI Enforcement ---
        
        # Step 2 (Issue Adoption) depends ONLY on metadata
        self.scn_nav_cards[1].set_enabled(metadata_valid)
        
        # Steps 3-6 depend ONLY on DRAFTING phase flag
        is_drafting = (self.scn_workflow_phase == "DRAFTING")
        for i in range(2, 6):
            self.scn_nav_cards[i].set_enabled(is_drafting)
            
        # Trigger downstream updates if we just entered Drafting phase
        if is_drafting and prev_phase == "METADATA":
             self.sync_demand_tiles()
        
        # Re-lock if metadata fails
        if not metadata_valid:
             for i in range(1, 6):
                 self.scn_nav_cards[i].set_enabled(False)

    def on_scn_page_changed(self, index):
        # [PHASE A] Hard Validation Gate: Progression from Step 1
        if self.active_scn_step == 0 and index > 0:
            if hasattr(self, 'scn_grounds_form'):
                errors = self.scn_grounds_form.validate(show_ui=True)
                if errors:
                    msg = "Please resolve the following legal inconsistencies in Step 1 before proceeding:\n\n" + "\n".join([f"• {e}" for e in errors])
                    QMessageBox.warning(self, "Validation Error", msg)
                    
                    # Force navigation back to Step 1
                    self.scn_side_nav.nav_cards[0].setChecked(True)
                    # Note: We don't update self.active_scn_step here to prevent infinite recursion
                    # switch_page will call this again with index 0
                    return

    def on_scn_page_changed(self, index):
        """
        Handle navigation within SCN Wizard (Stacked Widget).
        Now strictly Stage-Driven.
        """
        current_stage = self.get_current_stage()
        
        # Hard Entry Gate: Step 2 (Issue Adoption)
        if index == 1:
            # [LOGIC] SCN Drafting requires SCN_DRAFT stage.
            # This is already enforced by the main tab lock, but good as a secondary check.
            if current_stage < WorkflowStage.SCN_DRAFT:
                 self._show_blocking_msg("SCN Drafting is not active in this stage.")
                 # Go back to Step 1
                 self.scn_side_nav.nav_cards[0].setChecked(True)
                 return
            
            # Trigger Hydration
            self.hydrate_from_snapshot()
            
        # Hard Financial Gate: Step 6 (Preview & Finalization)
        if index == 5:
            invalid_cards = []
            for card in self.scn_issue_cards:
                if not card.validate_tax_inputs(show_ui=True):
                    invalid_cards.append(card.display_title)
            
            if invalid_cards:
                msg = f"Negative demand values detected in the following issues:\n\n" + "\n".join([f"• {c}" for c in invalid_cards]) + "\n\nSCN generation is blocked until these are corrected."
                QMessageBox.warning(self, "Financial Validation Error", msg)
                # Revert to Step 5 or wherever they were
                self.scn_side_nav.nav_cards[4].setChecked(True)
                return
            
            # If valid, proceed to hydration
            self.hydrate_from_snapshot()
            
        print(f"ProceedingsWorkspace: SCN Page {index} active")
        
        # [NEW] Deterministic Stage-Gated Preview
        # Trigger rendering ONLY when entering the finalization stage for the first time
        # We compare the page widget at the current index directly to our stored container
        if hasattr(self, 'scn_side_nav') and hasattr(self, 'scn_finalization_container'):
             if index == 5: # Definitively 'Actions & Finalize'
                 # Always refresh preview when entering the step to ensure data parity
                 self.render_final_preview()
            
        print(f"SCN: Moved to Step {index+1}.")

    def create_ph_intimation_tab(self):
        """Create Personal Hearing Intimation tab with neutral professional design"""
        main_widget = QWidget()
        main_widget.setStyleSheet(f"background-color: {Theme.NEUTRAL_100};") # Screen Background
        main_layout = QVBoxLayout(main_widget)
        main_layout.setContentsMargins(24, 24, 24, 24)
        main_layout.setSpacing(24)
        
        # Splitter for Form (Left) and Preview (Right)
        self.ph_splitter = QSplitter(Qt.Orientation.Horizontal)
        self.ph_splitter.setStyleSheet("QSplitter::handle { background: transparent; }")
        
        # --- LEFT PANE: FORM ---
        form_scroll = QScrollArea()
        form_scroll.setWidgetResizable(True)
        form_scroll.setStyleSheet("QScrollArea { border: none; background: transparent; }")
        
        form_container = QWidget()
        form_container.setStyleSheet("background: transparent;")
        form_layout = QVBoxLayout(form_container)
        form_layout.setContentsMargins(0, 0, 0, 0)
        form_layout.setSpacing(16)
        
        # 1. Header & Entries List
        header_layout = QHBoxLayout()
        header_label = QLabel("Personal Hearing Intimations")
        header_label.setStyleSheet(f"font-size: 11pt; font-weight: bold; color: {Theme.NEUTRAL_900};")
        
        self.ph_add_btn = QPushButton("+ Add New Entry")
        self.ph_add_btn.setFixedHeight(32)
        self.ph_add_btn.setStyleSheet(f"""
            QPushButton {{ 
                background-color: {Theme.SUCCESS}; color: white; padding: 0 16px; 
                font-weight: bold; border-radius: 6px; font-size: 13px;
            }}
            QPushButton:hover {{ background-color: {Theme.SUCCESS_HOVER}; }}
        """)
        self.ph_add_btn.clicked.connect(self.add_new_ph_entry)
        
        header_layout.addWidget(header_label)
        header_layout.addStretch()
        header_layout.addWidget(self.ph_add_btn)
        form_layout.addLayout(header_layout)
        
        # List of Entries (Cards)
        self.ph_list_container = QWidget()
        self.ph_list_layout = QVBoxLayout(self.ph_list_container)
        self.ph_list_layout.setContentsMargins(0, 0, 0, 0)
        self.ph_list_layout.setSpacing(12)
        form_layout.addWidget(self.ph_list_container)
        
        # 2. Entry Editor (Form) - Hidden by default
        self.ph_editor_card = QFrame()
        self.ph_editor_card.setStyleSheet(f"""
            QFrame {{ 
                background-color: {Theme.SURFACE}; 
                border-radius: 8px; 
                border: 1px solid {Theme.NEUTRAL_200};
            }}
            QLabel {{ border: none; color: {Theme.NEUTRAL_500}; font-size: 13px; }}
            QLineEdit, QDateEdit {{ 
                height: 32px; 
                border: 1px solid {Theme.NEUTRAL_200}; 
                border-radius: 6px; 
                padding: 0 8px;
                background-color: white;
            }}
        """)
        self.ph_editor_card.setVisible(False)
        
        shadow = QGraphicsDropShadowEffect()
        shadow.setBlurRadius(15)
        shadow.setXOffset(0)
        shadow.setYOffset(4)
        shadow.setColor(QColor(0, 0, 0, 40))
        self.ph_editor_card.setGraphicsEffect(shadow)

        editor_layout = QVBoxLayout(self.ph_editor_card)
        editor_layout.setContentsMargins(16, 16, 16, 16)
        editor_layout.setSpacing(16)
        
        editor_title = QLabel("Edit PH Details")
        editor_title.setStyleSheet(f"color: {Theme.NEUTRAL_900}; font-weight: bold; font-size: 14px; border: none;")
        editor_layout.addWidget(editor_title)
        
        grid = QGridLayout()
        grid.setSpacing(12)
        grid.setColumnStretch(1, 1)
        grid.setColumnStretch(3, 1)
        
        # Row 1: OC Details
        grid.addWidget(QLabel("OC No:"), 0, 0)
        
        oc_no_layout = QHBoxLayout()
        oc_no_layout.setSpacing(4)
        
        self.ph_edit_oc = QLineEdit()
        oc_no_layout.addWidget(self.ph_edit_oc)
        
        self.ph_auto_oc_btn = QPushButton("Auto-Generate")
        self.ph_auto_oc_btn.setFixedHeight(28)
        self.ph_auto_oc_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.ph_auto_oc_btn.setStyleSheet("""
            QPushButton {
                background-color: #f8fafc;
                border: 1px solid #cbd5e1;
                color: #475569;
                padding: 0 8px;
                border-radius: 4px;
                font-size: 11px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #f1f5f9;
                border-color: #3498db;
                color: #3498db;
            }
        """)
        self.ph_auto_oc_btn.clicked.connect(self._on_ph_auto_generate_oc)
        oc_no_layout.addWidget(self.ph_auto_oc_btn)
        
        grid.addLayout(oc_no_layout, 0, 1)
        
        grid.addWidget(QLabel("OC Date:"), 0, 2)
        self.ph_edit_oc_date = QDateEdit()
        self.ph_edit_oc_date.setCalendarPopup(True)
        self.ph_edit_oc_date.setDate(QDate.currentDate())
        grid.addWidget(self.ph_edit_oc_date, 0, 3)
        
        # Row 2: PH Details
        grid.addWidget(QLabel("PH Date:"), 1, 0)
        self.ph_edit_date = QDateEdit()
        self.ph_edit_date.setCalendarPopup(True)
        self.ph_edit_date.setDate(QDate.currentDate().addDays(7))
        grid.addWidget(self.ph_edit_date, 1, 1)
        
        grid.addWidget(QLabel("PH Time:"), 1, 2)
        self.ph_edit_time = QLineEdit("11:00 AM")
        grid.addWidget(self.ph_edit_time, 1, 3)
        
        # Row 3: Venue
        grid.addWidget(QLabel("Venue:"), 2, 0)
        self.ph_edit_venue = QLineEdit("Paravur Range Office")
        grid.addWidget(self.ph_edit_venue, 2, 1, 1, 3)
        
        # Row 4: Copy To
        grid.addWidget(QLabel("Copy To:"), 3, 0)
        self.ph_edit_copy_to = QLineEdit("The Assistant Commissioner, Central Tax, Paravur Division")
        grid.addWidget(self.ph_edit_copy_to, 3, 1, 1, 3)
        
        editor_layout.addLayout(grid)
        
        # Editor Buttons
        editor_btn_layout = QHBoxLayout()
        editor_btn_layout.setSpacing(12)
        
        cancel_btn = QPushButton("Cancel")
        cancel_btn.setFixedHeight(32)
        cancel_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        cancel_btn.setStyleSheet(f"color: {Theme.NEUTRAL_500}; background: transparent; border: none; font-weight: bold;")
        cancel_btn.clicked.connect(lambda: self.ph_editor_card.setVisible(False))
        
        save_entry_btn = QPushButton("Save Draft")
        save_entry_btn.setFixedHeight(32)
        save_entry_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        save_entry_btn.setStyleSheet(f"""
            QPushButton {{ 
                background: transparent; border: 1px solid {Theme.NEUTRAL_200}; 
                color: {Theme.NEUTRAL_900}; border-radius: 6px; padding: 0 16px; font-weight: bold;
            }}
            QPushButton:hover {{ background-color: {Theme.NEUTRAL_100}; }}
        """)
        save_entry_btn.clicked.connect(self.save_ph_entry)
        
        finalize_reg_btn = QPushButton("Finalize & Register")
        finalize_reg_btn.setFixedHeight(32)
        finalize_reg_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        finalize_reg_btn.setStyleSheet(f"""
            QPushButton {{ 
                background-color: {Theme.PRIMARY}; color: white; padding: 0 16px; 
                font-weight: bold; border-radius: 6px;
            }}
            QPushButton:hover {{ background-color: {Theme.PRIMARY_HOVER}; }}
        """)
        finalize_reg_btn.clicked.connect(self.register_ph_entry)
        
        editor_btn_layout.addStretch()
        editor_btn_layout.addWidget(cancel_btn)
        editor_btn_layout.addWidget(save_entry_btn)
        editor_btn_layout.addWidget(finalize_reg_btn)
        editor_layout.addLayout(editor_btn_layout)
        
        form_layout.addWidget(self.ph_editor_card)
        form_layout.addStretch()
        
        form_scroll.setWidget(form_container)
        form_scroll.setMinimumWidth(450) # Fix side scroll for PH editor
        self.ph_splitter.addWidget(form_scroll)
        
        # --- RIGHT PANE: PREVIEW ---
        preview_container = QWidget()
        preview_container.setStyleSheet("background: transparent;")
        preview_v_layout = QVBoxLayout(preview_container)
        preview_v_layout.setContentsMargins(0, 0, 0, 0)
        preview_v_layout.setSpacing(0)
        
        # Compact Preview Toolbar (Single Flex Row)
        self.ph_preview_toolbar = QFrame()
        self.ph_preview_toolbar.setFixedHeight(40)
        self.ph_preview_toolbar.setStyleSheet(f"background: transparent; border: none;")
        
        toolbar_layout = QHBoxLayout(self.ph_preview_toolbar)
        toolbar_layout.setContentsMargins(16, 8, 16, 8)
        toolbar_layout.setSpacing(8)
        
        live_preview_label = QLabel("Live Preview")
        live_preview_label.setStyleSheet(f"font-weight: bold; color: {Theme.NEUTRAL_900}; font-size: 14px;")
        toolbar_layout.addWidget(live_preview_label)
        toolbar_layout.addStretch()
        
        # Toolbar Actions
        self.ph_show_lh = QCheckBox("Show Letterhead")
        self.ph_show_lh.setChecked(True)
        self.ph_show_lh.stateChanged.connect(self.render_ph_preview)
        self.ph_show_lh.setStyleSheet(f"font-size: 13px; color: {Theme.NEUTRAL_500};")
        toolbar_layout.addWidget(self.ph_show_lh)
        
        btn_style = f"""
            QPushButton {{ 
                background: transparent; border: 1px solid {Theme.NEUTRAL_200}; 
                color: {Theme.NEUTRAL_900}; padding: 0 12px; border-radius: 6px; 
                font-size: 13px; font-weight: bold;
            }}
            QPushButton:hover {{ background-color: {Theme.NEUTRAL_100}; }}
        """
        
        self.ph_refresh_btn = QPushButton("🔄 Refresh Preview")
        self.ph_refresh_btn.setFixedHeight(32)
        self.ph_refresh_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.ph_refresh_btn.setStyleSheet(btn_style)
        self.ph_refresh_btn.clicked.connect(self.render_ph_preview)
        toolbar_layout.addWidget(self.ph_refresh_btn)
        
        self.ph_pdf_btn = QPushButton("👁 Draft PDF")
        self.ph_pdf_btn.setFixedHeight(32)
        self.ph_pdf_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.ph_pdf_btn.setStyleSheet(btn_style)
        self.ph_pdf_btn.clicked.connect(self.generate_ph_pdf)
        toolbar_layout.addWidget(self.ph_pdf_btn)
        
        self.ph_docx_btn = QPushButton("📝 Draft DOCX")
        self.ph_docx_btn.setFixedHeight(32)
        self.ph_docx_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.ph_docx_btn.setStyleSheet(btn_style)
        self.ph_docx_btn.clicked.connect(self.generate_ph_docx)
        toolbar_layout.addWidget(self.ph_docx_btn)
        
        # [PHASE 3] Conclude Hearing Action
        self.ph_conclude_btn = QPushButton("🏁 Conclude Hearing")
        self.ph_conclude_btn.setFixedHeight(32)
        self.ph_conclude_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.ph_conclude_btn.setStyleSheet(f"""
            QPushButton {{ 
                background-color: #f39c12; color: white; padding: 0 16px; 
                font-weight: bold; border-radius: 6px; font-size: 13px;
            }}
            QPushButton:hover {{ background-color: #e67e22; }}
        """)
        self.ph_conclude_btn.clicked.connect(self.conclude_ph_hearing)
        self.ph_conclude_btn.setVisible(False) # Hidden by default, shown via apply_context_layout
        toolbar_layout.addWidget(self.ph_conclude_btn)
        
        preview_v_layout.addWidget(self.ph_preview_toolbar)
        
        # Restore Phase 8 direct browser placement (fixes hidden preview regression)
        self.ph_browser = QWebEngineView()
        self.ph_browser.setStyleSheet(f"background-color: {Theme.NEUTRAL_100}; border: none;")
        preview_v_layout.addWidget(self.ph_browser)
        
        self.ph_splitter.addWidget(preview_container)
        self.ph_splitter.setStretchFactor(0, 1)
        self.ph_splitter.setStretchFactor(1, 2)
        
        main_layout.addWidget(self.ph_splitter)
        return main_widget
        
        
        
        
        


    def create_order_tab(self):
        print("ProceedingsWorkspace: create_order_tab start")
        """Create Order tab"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        layout.setSpacing(10)
        
        # Title
        title = QLabel("<b>Drafting Order</b>")
        title.setStyleSheet("font-size: 10pt; margin-bottom: 10px;")
        layout.addWidget(title)
        
        # Reference Details (OC)
        ref_layout = QHBoxLayout()
        oc_label = QLabel("OC No:")
        self.order_oc_input = QLineEdit()
        self.order_oc_input.setPlaceholderText("Format: No./Year")
        
        
        order_oc_suggest_btn = QPushButton("Get Next")
        order_oc_suggest_btn.setStyleSheet("padding: 2px 8px; background-color: #3498db; color: white; border-radius: 4px; font-size: 8pt;")
        order_oc_suggest_btn.clicked.connect(lambda: self.suggest_next_oc(self.order_oc_input))
        
        oc_date_label = QLabel("OC Date:")
        self.order_oc_date = QDateEdit()
        self.order_oc_date.setCalendarPopup(True)
        self.order_oc_date.setDate(QDate.currentDate())
        
        
        ref_layout.addWidget(oc_label)
        ref_layout.addWidget(self.order_oc_input)
        ref_layout.addWidget(order_oc_suggest_btn)
        ref_layout.addSpacing(10)
        ref_layout.addWidget(oc_date_label)
        ref_layout.addWidget(self.order_oc_date)
        ref_layout.addStretch()
        layout.addLayout(ref_layout)
        
        # Rich Text Editor
        print("ProceedingsWorkspace: creating order_editor")
        self.order_editor = RichTextEditor("Enter the Order content here...")
        print("ProceedingsWorkspace: order_editor created")
        self.order_editor.setMinimumHeight(150)
        
        layout.addWidget(self.order_editor)
        print("ProceedingsWorkspace: order_editor added")
        
        # Action Buttons
        buttons_layout = QHBoxLayout()
        
        save_btn = QPushButton("Save Draft")
        save_btn.setStyleSheet("background-color: #95a5a6; color: white; padding: 8px 20px; font-weight: bold; border-radius: 4px;")
        save_btn.clicked.connect(lambda: self.save_document("Order"))
        buttons_layout.addWidget(save_btn)
        
        pdf_btn = QPushButton("Generate PDF")
        pdf_btn.setStyleSheet("background-color: #e74c3c; color: white; padding: 8px 20px; font-weight: bold; border-radius: 4px;")
        pdf_btn.clicked.connect(self.generate_pdf)
        buttons_layout.addWidget(pdf_btn)
        
        docx_btn = QPushButton("Generate DOCX")
        docx_btn.setStyleSheet("background-color: #3498db; color: white; padding: 8px 20px; font-weight: bold; border-radius: 4px;")
        docx_btn.clicked.connect(self.generate_docx)
        buttons_layout.addWidget(docx_btn)
        
        self.order_finalize_btn = QPushButton("Finalize & Register")
        self.order_finalize_btn.setStyleSheet("background-color: #27ae60; color: white; padding: 8px 20px; font-weight: bold; border-radius: 4px;")
        self.order_finalize_btn.clicked.connect(self.confirm_order_finalization)
        buttons_layout.addWidget(self.order_finalize_btn)
        
        buttons_layout.addStretch()
        layout.addLayout(buttons_layout)
        
        # Wrap in a centered container to constrain width
        container = QWidget()
        container_layout = QVBoxLayout(container)
        container_layout.setAlignment(Qt.AlignmentFlag.AlignHCenter)
        container_layout.addWidget(widget)
        
        # Set max width for the form content
        widget.setMaximumWidth(850) 
        
        # Wrap in Scroll Area
        main_scroll = QScrollArea()
        main_scroll.setWidgetResizable(True)
        main_scroll.setWidget(container)
        return main_scroll

    def create_scn_finalization_panel(self):
        """Create the SCN Finalization Summary Panel with embedded QWebEngineView (Chromium)"""
        container = QWidget()
        container.setObjectName("SCNFinalizationContainer")
        layout = QVBoxLayout(container)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)
        
        # Add Panel
        self.scn_fin_panel = FinalizationPanel()
        layout.addWidget(self.scn_fin_panel)
        
        # Link reference for code compatibility
        self.scn_final_preview = self.scn_fin_panel.browser
        
        # Consolidated Button Connections
        self.scn_fin_panel.save_btn.clicked.connect(lambda: self.save_document("SCN", is_manual=True))
        self.scn_fin_panel.pdf_btn.clicked.connect(self.generate_pdf)
        self.scn_fin_panel.docx_btn.clicked.connect(self.generate_docx)
        self.scn_fin_panel.refresh_btn.clicked.connect(self.render_final_preview)
        self.scn_fin_panel.finalize_btn.clicked.connect(self.confirm_scn_finalization)
        
        # Navigation
        self.scn_fin_panel.cancel_btn.clicked.connect(lambda: self.scn_side_nav.nav_cards[0].click()) # Back to Step 1
        
        return container

    def render_final_preview(self):
        """
        One-shot or manual refresh of the high-fidelity Chromium preview.
        Gathers in-memory state from all IssueCards.
        """
        print("ProceedingsWorkspace: Generating deterministic SCN preview...")
        # 1. Update data for FinalizationPanel summary
        scn_no = self.scn_no_input.text()
        scn_date = self.scn_date_input.date().toString('dd-MM-yyyy')
        oc_no = self.scn_oc_input.text()
        
        self.scn_fin_panel.load_data(
            proceeding_data=self.proceeding_data,
            issues_list=self.scn_issue_cards,
            doc_type="SCN",
            doc_no=scn_no,
            doc_date=scn_date,
            ref_no=oc_no
        )
        
        # 2. Render HTML and load into WebEngineView
        html = self.render_scn(is_preview=True)
        self.scn_final_preview.setHtml(html)
        
        # Mark as initialized to prevent auto-render on every future tab switch
        self.preview_initialized = True

    def confirm_scn_finalization(self):
        """Commit SCN Finalization"""
        try:
            # 1. Save Document as Final
            self.save_document("SCN", is_manual=True) # Ensure latest draft is saved
            
            # 2. Transition State (Atomic Update)
            self.transition_to(WorkflowStage.SCN_ISSUED)
            
            # 3. Update Register (SCN Register entry - implicitly done via status or explicit call?)
            # Also Add OC Entry for SCN
            oc_data = {
                'OC_Number': self.scn_oc_input.text(),
                'OC_Date': self.scn_date_input.date().toString("yyyy-MM-dd"),
                'OC_Content': f"Show Cause Notice Issued. SCN No: {self.scn_no_input.text()}. {self.scn_fin_panel.fin_scn_remarks.toPlainText()}",
                'OC_To': self.proceeding_data.get('legal_name', '')
            }
            self.db.add_oc_entry(self.proceeding_id, oc_data)
            
            QMessageBox.information(self, "Success", "Show Cause Notice Finalized Successfully.")
            
            # 4. Switch to View Mode
            self.scn_finalization_container.hide()
            self.scn_view_container.show()
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Finalization failed: {e}")
            import traceback
            traceback.print_exc()

    def load_scn_issue_templates(self):
        """Load prioritized issue templates for SCN tab (Strictly Limited for Adoption + Manual SOP)"""
        self.scn_issue_combo.clear()
        
        try:
            # Gather IDs of currently present issues in the draft
            present_issue_ids = set()
            if hasattr(self, 'scn_issue_cards'):
                present_issue_ids = {c.template['issue_id'] for c in self.scn_issue_cards}
            
            # 1. Upstream Adoptions (Scrutiny or Direct)
            adopted_from_upstream_ids = set()
            source_id = self.proceeding_data.get('source_scrutiny_id') or self.proceeding_data.get('scrutiny_id')
            
            if source_id:
                finalized_issues = self.db.get_case_issues(source_id, stage='DRC-01A')
                if finalized_issues:
                    self.scn_issue_combo.addItem("--- ADOPT FROM ASMT-10 ---", None)
                    for record in finalized_issues:
                        issue_id = record['issue_id']
                        adopted_from_upstream_ids.add(issue_id) # Corrected variable name
                        
                        # Skip if already present in draft (Dynamic Availability)
                        if issue_id in present_issue_ids:
                            continue

                        # Use Adapter to resolve SCN template for the dropdown
                        adapted = self.build_scn_issue_from_asmt10(record)
                        issue_name = adapted['template'].get('issue_name', 'Issue')
                        
                        # Structured Payload
                        payload = {
                            'issue_id': issue_id,
                            'origin': 'ASMT10',
                            'template': adapted['template'],
                            'data': adapted['data'],
                            'is_unique': True
                        }
                        self.scn_issue_combo.addItem(issue_name, payload)

            # 2. Manual SOP Addition (The Feature)
            self.scn_issue_combo.addItem("--- MANUAL ADDITION (SOP ISSUES) ---", None)
            
            # Fetch all active issues
            all_issues_meta = self.db.get_all_issues_metadata() # Lightweight fetch
            
            for meta in all_issues_meta:
                issue_id = meta['issue_id']
                
                # Filter: Must be SOP issue, NOT in ASMT-10 source
                if not issue_id.startswith('SOP-'): continue
                if issue_id in adopted_from_upstream_ids: continue # Corrected variable name
                
                # Check existance for SOP issues (Unique)
                if issue_id in present_issue_ids: continue
                
                payload = {
                    'issue_id': issue_id,
                    'issue_name': meta['issue_name'],
                    'origin': 'MANUAL_SOP',
                    'is_unique': True
                }
                self.scn_issue_combo.addItem(f"{issue_id}: {meta['issue_name']}", payload)

            # 3. SCN-Origin / Custom Templates
            self.scn_issue_combo.addItem("--- OTHER TEMPLATES ---", None)
            
            # SCN Master Templates (from templates table, type SCN)
            all_templates = self.db.get_issue_templates()
            for t in all_templates:
                if t.get('type') == 'SCN':
                    payload = {
                        'issue_id': t['issue_id'],
                        'origin': 'SCN',
                        'template': t,
                        'is_unique': False # Custom SCN templates might be re-usable? Let's say yes for now.
                    }
                    self.scn_issue_combo.addItem(f"Template: {t['issue_name']}", payload)
            
            # Blank SCN Issue
            blank_payload = {
                'issue_id': "BLANK_SCN_ISSUE",
                'issue_name': "Blank SCN Issue",
                'origin': "SCN",
                'is_unique': False
            }
            self.scn_issue_combo.addItem("Add Blank SCN Issue", blank_payload)

        except Exception as e:
            print(f"Error loading SCN templates: {e}")
            import traceback
            traceback.print_exc()

    def add_scn_issue_card(self, template, data=None, source_type=None, source_id=None, origin="SCN", status="ACTIVE", trigger_eval=True):
        """Add flexible issue card to SCN layout"""
        from src.ui.issue_card import IssueCard
        card = IssueCard(template, data=data, mode="SCN", content_key="scn_content")
        
        # Apply flexible classification (informational badge)
        card.set_classification(origin=origin, status=status)
        if source_id:
            card.set_source_metadata(source_type or "ASMT10", source_id)
            # Explicitly persist source info on the object for snapshotting
            card.source_issue_id = source_id 
            card.source_proceeding_id = self.proceeding_data.get('source_scrutiny_id')
        
        # Persist origin on object for snapshotting
        card.origin = origin

        if data:
            # load_data handles restoring state from saved drafts
            # But wait, IssueCard usually takes data in constructor.
            # If load_data exists it might override.
            # Inspecting IssueCard.. it uses constructor data.
            # load_data method doesn't exist in IssueCard class shown earlier?
            # IssueCard.__init__ logic:
            # if data and 'table_data' in data ...
            # So data passed to constructor is enough.
            pass
        
        # Connect signals
        card.removeClicked.connect(lambda c_obj=card: self.remove_scn_issue_card(None, c_obj))
        

        

        
        self.scn_issues_layout.addWidget(card)
        self.scn_issue_cards.append(card)
        
        # Trigger Phase Evaluation
        if trigger_eval:
            self.evaluate_scn_workflow_phase()
        
        return card

    def insert_scn_issue(self):
        """Insert a manually selected issue template into the drafting layout"""
        # Guard: Phase-1 Finalization
        if self.is_scn_finalized():
             QMessageBox.warning(self, "Action Blocked", "Show Cause Notice is already finalized.")
             return

        payload = self.scn_issue_combo.currentData()
        if not payload:
             QMessageBox.warning(self, "Invalid Selection", "Please select a valid issue from the list (not a category header).")
             return
        
        # Self-Healing: Handle Legacy String Payload (if dropdown stale)
        if isinstance(payload, str):
             print(f"Warning: Stale dropdown payload (str) detected: {payload}")
             # Force reload
             self.load_scn_issue_templates()
             QMessageBox.warning(self, "Refresh Required", "Issue list was stale. Refreshed. Please select and insert again.")
             return

        try:
            issue_id = payload['issue_id']
            origin = payload['origin']
            issue_name = payload.get('issue_name', issue_id)
            
            # Resolve Template
            # 1. Fresh Canonical Fetch for Manual / SCN Templates
            if origin in ["SCN", "MANUAL_SOP"]:
                # Fetch full authoritative JSON from Master DB (JIT)
                template = self.db.get_issue(issue_id)
                if not template:
                    QMessageBox.warning(self, "Integration Error", 
                                      f"Could not load authoritative schema for issue '{issue_id}'.\nCheck Issue Master integrity.")
                    return
                data = payload.get('data') # Fresh or pre-filled
                source_id = None
                
            # 2. Adoption from ASMT-10 (Pre-adapted via build_scn_issue_from_asmt10)
            elif origin == "ASMT10" and 'template' in payload:
                template = payload['template']
                data = payload.get('data')
                source_id = issue_id
            
            # 3. Blank/Generic
            elif issue_id == "BLANK_SCN_ISSUE":
                import uuid
                # Generate unique ID for this instance
                unique_id = f"SCN_MANUAL_{uuid.uuid4().hex[:8]}"
                template = {
                    'issue_id': unique_id,
                    'issue_name': "New SCN Issue",
                    'brief_facts_scn': "[Enter Facts]",
                    'templates': {'brief_facts_scn': "[Enter Facts]"},
                    'variables': {}
                }
                data = None
                source_id = None
                
            # 4. Enforce Structured Duplicate Check before Insertion
            # Rules:
            # - SCRUTINY/ASMT10: Unique by source_issue_id (already adopted?)
            # - MANUAL_SOP: Unique by issue_id (don't add same SOP point twice)
            # - BLANK/SCN: Multi-instance allowed (though blank usually usually unique until saved with content)
            
            is_duplicate = False
            for card in self.scn_issue_cards:
                 existing_origin = getattr(card, 'origin', 'SCN')
                 existing_template_id = card.template.get('issue_id')
                 existing_source_id = getattr(card, 'source_issue_id', None)
                 
                 # duplicate check
                 if origin in ["ASMT10", "SCRUTINY"] and source_id:
                      if existing_source_id == source_id: 
                           is_duplicate = True; break
                 elif origin == "MANUAL_SOP":
                      if existing_template_id == issue_id:
                           is_duplicate = True; break
                           
            if is_duplicate:
                 QMessageBox.information(self, "Issue Already Added", 
                                       f"The issue '{issue_name}' has already been added to the draft.")
                 return

            # [FIX] JIT Repair for Insertion (Safety Net)
            if template:
                tgt_grid = template.get('grid_data')
                if tgt_grid and isinstance(tgt_grid, dict):
                    cols = tgt_grid.get('columns')
                    if isinstance(cols, list) and cols and not isinstance(cols[0], dict):
                         print(f"  - JIT Repair (Insert): Upgrading string columns for {template.get('issue_id')}")
                         new_cols = [{"id": f"col{i}", "label": str(c)} for i, c in enumerate(cols)]
                         tgt_grid['columns'] = new_cols

            self.add_scn_issue_card(
                template=template,
                data=data,
                origin=origin,
                source_id=source_id
            )
            
            # Immediate Persistence for Manual Issues
            if origin == "MANUAL_SOP":
                 self.persist_scn_issues()

        except Exception as e:
            QMessageBox.critical(self, "Insertion Error", f"Failed to insert issue: {e}")
            import traceback
            traceback.print_exc()

        # Phase-1 Isolation: Guarded at highest call-site
        if not self.is_scn_phase1():
            self.sync_demand_tiles()
            
            self.save_document("SCN")

    def remove_scn_issue_card(self, modern_card, card):
        # Guard: Phase-1 Finalization
        if self.is_scn_finalized():
             QMessageBox.warning(self, "Action Blocked", "Show Cause Notice is already finalized.")
             return

        # Guard: Manual SOP Deletion Warning
        if getattr(card, 'origin', 'SCRUTINY') == "MANUAL_SOP":
             # Requirement: Explicit Warning
             reply = QMessageBox.question(self, "Delete Manual Issue", 
                                        "You are deleting a manually added SOP issue. This cannot be undone.\n\nProceed?",
                                        QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
             if reply == QMessageBox.StandardButton.No:
                  return

        if modern_card:
            self.scn_issues_layout.removeWidget(modern_card)
            modern_card.deleteLater()
        else:
            self.scn_issues_layout.removeWidget(card)
            card.deleteLater()
            
        if card in self.scn_issue_cards:
            self.scn_issue_cards.remove(card)
            
        # Trigger Phase Evaluation
        self.evaluate_scn_workflow_phase()
        # Phase-1 Isolation: Guarded at highest call-site
        if not self.is_scn_phase1():
            self.sync_demand_tiles()
            
            self.save_document("SCN")

    def is_scn_phase1(self):
        """State-Absolute Phase-1 detection: True only if we are still in METADATA phase."""
        return self.scn_workflow_phase == "METADATA"

    def hydrate_from_snapshot(self):
        """
        Snapshot-Only Hydration Strategy for SCN.
        Rules:
        1. Load existing 'SCN' draft if present.
        2. Adopt/Merge 'DRC-01A' issues (from Current or Source proceeding).
           - Supports Direct Adjudication (Current ID)
           - Supports Scrutiny Flow (Source ID)
        """
        # 1. Idempotency Gate
        if self.scn_issues_initialized and self.scn_issue_cards:
            return

        # 2. Hard Clear UI
        while self.scn_issues_layout.count():
            item = self.scn_issues_layout.takeAt(0)
            widget = item.widget()
            if widget: widget.deleteLater()
        self.scn_issue_cards = []

        try:
            # 3. Fetch Existing SCN Draft from DB
            existing_scn_draft = self.db.get_case_issues(self.proceeding_id, stage='SCN')
            
            loaded_ids = set()
            if existing_scn_draft:
                print(f"SCN Load: Restoring {len(existing_scn_draft)} cards from snapshot.")
                self._hydrate_cards_from_records(existing_scn_draft, is_initial_hydration=False)
                loaded_ids = {r.get('issue_id') for r in existing_scn_draft}
                
            # [FIX] SCN Adoption: Fetch source issues from DRC-01A
            # Priority: 1. Current Proceeding (Direct Adjudication) 2. Source Proceeding (Scrutiny)
            source_issues = []
            
            # Check Local (Direct Adjudication or Same-Proceeding Flow)
            # Fetch 'DRC-01A' to adopt. (ASMT-10 is irrelevant for SCN if DRC-01A exists)
            local_drc = self.db.get_case_issues(self.proceeding_id, stage='DRC-01A')
            if local_drc:
                 source_issues = local_drc
                 print(f"SCN Adoption: Found {len(local_drc)} DRC-01A issues in Current Proceeding.")
            else:
                 # Check Source (Classic Flow)
                 src_id = self.proceeding_data.get('source_scrutiny_id') or self.proceeding_data.get('scrutiny_id')
                 if src_id:
                     remote_drc = self.db.get_case_issues(src_id, stage='DRC-01A')
                     # Fallback to ASMT-10 ONLY if legacy/no DRC-01A found (optional, but robust)
                     if not remote_drc:
                          remote_drc = self.db.get_case_issues(src_id, stage='ASMT-10')
                          
                     if remote_drc:
                          source_issues = remote_drc
                          print(f"SCN Adoption: Found {len(remote_drc)} Source issues (SourceID={src_id}).")

            # Identify missing
            missing_records = [r for r in source_issues if r.get('issue_id') not in loaded_ids]
            
            if missing_records:
                print(f"SCN Merge: Adopting {len(missing_records)} missing issues from Source/DRC-01A...")
                # Treat these as "initial hydration" (Use Adapter)
                self._hydrate_cards_from_records(missing_records, is_initial_hydration=True)
                # [PERSISTENCE] Save immediately to lock in the merged state
                self.save_scn_issue_snapshot()
            
            self.scn_issues_initialized = True
            self._persist_scn_init_flag()


        except Exception as e:
            print(f"SCN Hydration Error: {e}")
            import traceback
            traceback.print_exc()

    def persist_scn_issues(self, is_manual_save=False):
        """
        Orchestrate saving the current SCN issue state to DB (case_issues).
        Captures the exact snapshot of the drafting area INCLUDING TEMPLATES.
        Requirement: Snapshot-Only Hydration (Independent of Master)
        """
        # [PHASE C] Structural Integrity Enforcement
        # Ensure totals are recalculated before persistence
        self.calculate_grand_totals()

        # [PHASE A] Financial Integrity Block (Selective for Finalization)
        if not is_manual_save:
            invalid_cards = []
            for card in self.scn_issue_cards:
                if not card.validate_tax_inputs(show_ui=True):
                    invalid_cards.append(card.display_title)
            
            if invalid_cards:
                msg = f"Cannot save SCN. Negative demand values detected in:\n\n" + "\n".join([f"• {c}" for c in invalid_cards])
                QMessageBox.critical(self, "Financial Validation Error", msg)
                return

        try:
            current_snapshot = []
            
            for card in self.scn_issue_cards:
                # Issue Card MUST expose its origin and source info
                
                # Extract Origin
                origin = getattr(card, 'origin', 'SCRUTINY')
                source_pid = getattr(card, 'source_proceeding_id', None)
                
                card_data = card.get_data()
                
                # [PHASE B] Template Traceability
                # Deterministic hash generation from canonical representation
                template_dict = getattr(card, 'template', {})
                template_version = template_dict.get('version', '1.0')
                template_hash = self.db.generate_canonical_hash(template_dict)

                # CRITICAL HARDENING: Save the Template Structure
                # This ensures we don't need to consult Master when re-hydrating.
                # card.template contains the structure used to render the card.
                # We inject it into the 'data_json' blob.
                # NOTE: This increases DB size but guarantees legal immutability.
                
                data_payload = {
                     'values': card_data.get('variables', {}),
                     'table_data': card_data.get('table_data'),
                     'template_snapshot': card.template, # FULL TEMPLATE
                     'template_version': template_version,
                     'template_schema_hash': template_hash,
                     # Persist other state flags if needed
                     'status': getattr(card, 'status', 'ACTIVE')
                }
                
                if data_payload['template_snapshot']:
                    # print(f"  - Snapshot has grid_data? {'grid_data' in data_payload['template_snapshot']}")
                    pass

                snapshot_item = {
                    'issue_id': card.template.get('issue_id'),
                    'data': data_payload, # Contains everything needed to render
                    'origin': origin,
                    'source_proceeding_id': source_pid,
                    'added_by': 'User' 
                }
                current_snapshot.append(snapshot_item)
            
            # [PHASE B] Draft Versioning (Manual Save ONLY)
            if is_manual_save:
                print(f"SCN Governance: Triggering Draft Snapshot for Proceeding {self.proceeding_id}")
                self.db.save_proceeding_draft(self.proceeding_id, current_snapshot)
            
            # Save to DB
            try:
                import json
                # [DEBUG] Pre-flight serialization check to catch Circular References
                for i, item in enumerate(current_snapshot):
                    try:
                        # Test dump to catch cycles early
                        json.dumps(item) 
                    except ValueError as ve:
                        print(f"[CRITICAL] Serialization Failed for Card {i} ({item.get('issue_id')}): {ve}")
                        # Analyze keys
                        for k, v in item.get('data', {}).items():
                            try:
                                json.dumps(v)
                            except ValueError:
                                print(f"  -> BAD KEY: {k} maps to {type(v)}")
                        # Fail unsafe to prevent corrupt save
                        raise ve

                success = self.db.save_scn_issue_snapshot(self.proceeding_id, current_snapshot)
                if success:
                    print(f"SCN Persistence: Saved {len(current_snapshot)} issues with snapshot templates.")
                else:
                    print("SCN Persistence: Failed to save to DB.")
            except Exception as e:
                import traceback
                traceback.print_exc()
                print(f"Error persisting SCN issues (Serialization/DB): {e}")

        except Exception as e:
            print(f"Error persisting SCN issues (Outer): {e}")

    def is_stub_grid(self, grid_data):
        """Authoritative check: Is the grid data empty or just a placeholder stub?
        Supports both list-of-lists (legacy) and Dictionary Schema (Modern).
        """
        if not grid_data:
            return True
            
        # Case 1: Dictionary Schema {"columns": [], "rows": []}
        if isinstance(grid_data, dict):
            rows = grid_data.get('rows', [])
            if not rows: return True
            for row in rows:
                for col_id, cell in row.items():
                    if isinstance(cell, dict):
                        val = cell.get('value')
                        if val not in (None, "", "____"):
                            return False
                    elif cell not in (None, "", "____"):
                        return False
            return True

        # Case 2: List of Lists (Legacy)
        if isinstance(grid_data, list):
            if len(grid_data) <= 1: # Only header or nothing
                return True
            for row in grid_data[1:]:
                for cell in row:
                    if isinstance(cell, dict):
                        val = cell.get('value')
                        if val not in (None, "", "____"):
                            return False
                    elif cell not in (None, "", "____"):
                        return False
            return True
            
        return True

    def _normalize_issue_id(self, issue_id):
        """Strip 'Point X- ' prefixes for robust template lookup"""
        if not issue_id: return issue_id
        import re
        normalized = re.sub(r'^Point \d+- ', '', str(issue_id))
        return normalized

    def _normalize_summary_table(self, summary_table: dict) -> dict:
        """
        Convert Scrutiny summary_table (headers/rows) into a semantic NormalizedTable structure.
        Roles are identified heuristically (e.g., last row is Difference in a 3-row table).
        """
        headers = summary_table.get('headers')
        if not headers:
             # Fallback: Scrutiny summary might use 'columns' (Canonical Schema)
             raw_cols = summary_table.get('columns', [])
             headers = [c.get('label') if isinstance(c, dict) else str(c) for c in raw_cols]
             
        rows_data = summary_table.get('rows', [])
        
        # [ROBUSTNESS] Header Inference from First Row if missing
        if not headers and rows_data and isinstance(rows_data[0], dict):
             headers = list(rows_data[0].keys())
             print(f"[CONV DIAG] Infereed headers from row data: {headers}")

        # [DIAGNOSTIC] Log input structure
        print(f"[CONV DIAG] Headers Final: {headers}")
        print(f"[CONV DIAG] Rows Count: {len(rows_data)}")
        
        # 1. Map Columns to Tax Heads
        columns = []
        for i, h in enumerate(headers):
            h_upper = str(h).upper()
            col = {"id": f"col_{i}", "label": str(h)}
            if i == 0:
                col["id"] = "label"
                col["static"] = True
            elif "IGST" in h_upper: col["tax_head"] = "IGST"; col["id"] = "igst"
            elif "CGST" in h_upper: col["tax_head"] = "CGST"; col["id"] = "cgst"
            elif "SGST" in h_upper: col["tax_head"] = "SGST"; col["id"] = "sgst"
            elif "CESS" in h_upper: col["tax_head"] = "CESS"; col["id"] = "cess"
            columns.append(col)

        # 2. Determine Row Roles
        # Scrutiny summary tables usually have 3 rows: [Base1, Base2, Difference]
        normalized_rows = []
        num_rows = len(rows_data)
        
        num_value_cols = max(0, len(columns) - 1)
        
        for i, row in enumerate(rows_data):
            # 2.1 Calculate Role Deterministically
            role = "BASE"
            if num_rows >= 2 and i == num_rows - 1:
                role = "DIFFERENCE"

            try:
                # Row label handling: Summary table can have rows as lists or dicts (col0, col1...)
                if isinstance(row, dict):
                    # 1. Label handling (Semantic vs col0)
                    label = str(row.get("col0", ""))
                    if not label:
                        # Semantic Fallback: Try 'description', 'desc', 'label'
                        for k in ["description", "desc", "label", "particulars"]:
                            if k in row:
                                label = str(row[k])
                                break
                                
                    # 2. Value Mapping (Semantic vs colX)
                    values = []
                    for j in range(1, len(columns)):
                        col_def = columns[j]
                        header_label = col_def["label"].lower().strip()
                        
                        # A. Try Standard colX
                        val = row.get(f"col{j}")
                        
                        # B. Try Semantic Mapping (Deterministic Aliases)
                        if val is None:
                            # Map Header Label -> Potential Keys
                            aliases = []
                            
                            if "amount" in header_label or "value" in header_label:
                                aliases = ["amount", "value", "amt", "total_value"]
                            elif "tax" in header_label:
                                if "igst" in header_label: aliases = ["igst", "tax_igst", "igst_amt"]
                                elif "cgst" in header_label: aliases = ["cgst", "tax_cgst", "cgst_amt"]
                                elif "sgst" in header_label: aliases = ["sgst", "tax_sgst", "sgst_amt"]
                                elif "cess" in header_label: aliases = ["cess", "tax_cess", "cess_amt"]
                                else: aliases = ["tax", "tax_amt"]
                            elif "turnover" in header_label:
                                aliases = ["turnover", "total_turnover"]
                            elif "rate" in header_label:
                                aliases = ["rate", "tax_rate"]
                            
                            # Attempt alias lookup
                            for alias in aliases:
                                if alias in row:
                                    val = row[alias]
                                    break
                        
                        # Default to 0 if not found
                        if val is None:
                            val = 0
                            
                        values.append(val)
                else:
                    label = str(row[0]) if isinstance(row, list) and len(row) > 0 else f"Row {i+1}"
                    # Handle list rows explicitly
                    raw_values = row[1:] if isinstance(row, list) else []
                    values = []
                    for rv in raw_values:
                        values.append(rv)

                # 2.2 Enforce Rectangular Shape
                # 1. Truncate extra columns
                values = values[:num_value_cols]
                # 2. Pad missing columns with 0
                while len(values) < num_value_cols:
                    values.append(0)
                
                normalized_rows.append({
                    "role": role,
                    "label": label,
                    "values": values
                })

            except Exception as e:
                # [SAFEGUARD] Fallback to zero-row BUT maintain contract (Role)
                print(f"Row {i} normalization failed: {e}")
                normalized_rows.append({
                    "role": role, # CRITICAL: Maintain role (e.g. DIFFERENCE) so grid conversion binds variables
                    "label": "Error Row",
                    "values": [0] * num_value_cols
                })

        return {
            "columns": columns,
            "rows": normalized_rows
        }

    def _convert_normalized_to_grid(self, normalized_table: dict) -> dict:
        """
        Transform NormalizedTable into the SCN-compatible grid_data schema (Canonical Dict).
        Binds DIFFERENCE row cells to tax variables.
        Rules: Deterministic validation, explicit cell schema.
        """
        norm_columns = normalized_table.get("columns", [])
        rows = normalized_table.get("rows", [])
        
        # 1. Assert Invariants
        if not isinstance(norm_columns, list):
             raise ValueError(f"Normalized columns must be a list, got {type(norm_columns)}")
             
        num_cols = len(norm_columns)
        if num_cols <= 0:
             # [FIX] Explicit check to prevent KeyError: 0 downstream
             raise ValueError("MANDATORY: Table must have at least one column (label). Table is empty.")

        # 2. Prepare Columns for GridAdapter Schema
        grid_columns = []
        col_mappings = {} 
        
        for i, col in enumerate(norm_columns):
            grid_id = f"col{i}" 
            grid_columns.append({"id": grid_id, "label": col.get("label", f"Col {i}")})
            col_mappings[i] = grid_id

        # 3. Convert Rows to Canonical Schema
        grid_rows = []

        for r_idx, norm_row in enumerate(rows):
            # Assert row integrity
            assert "role" in norm_row, f"Row {r_idx} missing role"
            assert "label" in norm_row, f"Row {r_idx} missing label"
            assert "values" in norm_row, f"Row {r_idx} missing values"
            
            values = norm_row["values"]
            role = norm_row["role"]
            label = norm_row["label"]
            
            # Assert Rectangular Integrity: num_values == num_cols - 1 (label col excluded from values list)
            assert len(values) == num_cols - 1, f"Row {r_idx} value count ({len(values)}) mismatch with columns ({num_cols-1})"

            grid_row = {}

            # 3.1 Label Cell (col0) - Explicit Schema
            grid_id_0 = col_mappings[0]
            grid_row[grid_id_0] = {
                "value": label, 
                "type": "static"
            }

            # 3.2 Data Cells - Explicit Schema
            for i, val in enumerate(values):
                col_idx = i + 1
                grid_id = col_mappings[col_idx]
                col_meta = norm_columns[col_idx]
                tax_head = col_meta.get("tax_head")
                
                cell = {"value": val}
                
                if role == "DIFFERENCE" and tax_head:
                    # Bound to calculation engine
                    cell["type"] = "input"
                    cell["var"] = f"tax_{tax_head.lower()}"
                else:
                    cell["type"] = "static"
                    
                grid_row[grid_id] = cell
            
            grid_rows.append(grid_row)

        return {
            "columns": grid_columns,
            "rows": grid_rows
        }


    def build_scn_issue_from_asmt10(self, asmt_record: dict) -> dict:
        """
        Final Authoritative Adapter: Resolves SCN-specific templates while carrying forward
        only factual table data from ASMT-10 snapshots.
        """
        issue_id = asmt_record['issue_id']
        asmt_data = asmt_record['data']

        # 1. Identity & Template Resolution (Prefer SCN doc_type if exists)
        # Try exact ID, then normalized ID, then Name-based
        normalized_id = self._normalize_issue_id(issue_id)
        scn_template = self.db.get_issue(issue_id) or \
                       self.db.get_issue(normalized_id) or \
                       self.db.get_issue_by_name(normalized_id) or \
                       self.db.get_issue_by_name(issue_id)
        
        # 2. Strict Template Normalization for SCN
        if not scn_template:
             # Create a minimal scaffold if no master template exists
             scn_template = {
                 'issue_id': issue_id,
                 'issue_name': asmt_data.get('issue', asmt_data.get('category', 'Issue')),
                 'templates': {}
             }
        
        # 2.1 Ensure SCN-specific keys exist or fallback
        t = scn_template.get('templates', {})
        
        # Pull Narration from ASMT-10 if SCN template is new or empty
        asmt_narration = asmt_data.get('brief_facts') or asmt_data.get('description', '')
        
        if not t.get('brief_facts_scn') and not t.get('scn'):
             t['brief_facts_scn'] = asmt_narration
        if not t.get('grounds'):
             t['grounds'] = ""
        if not t.get('legal'):
             t['legal'] = ""
        if not t.get('conclusion'):
             t['conclusion'] = ""
        
        scn_template['templates'] = t

        # 3. Factual Table Carry-Forward (Direct Pass-Through)
        # [USER REQUEST] "The exact tables which were in the asmt-10 should appear"
        # We bypass complex normalization and trust the source structure.
        
        summary_table = asmt_data.get('summary_table')
        grid_data = asmt_data.get('grid_data')
        
        final_grid = None

        # [FIX] Prioritize Analysis Result (summary_table) over Static Template (grid_data)
        # The Master DB has 'grid_data' initialized with Zeros, which masked the real values.
        
        # Priority 1: Summary Table (Convert to Grid if needed)
        if summary_table:
            # Check if it's already in canonical format (rows + columns) OR just needs render-time inference
            # [DIRECT PASS-THROUGH] We now trust the updated renderer (ui_helpers) to handle List-rows
            if hasattr(summary_table, '__iter__') and not isinstance(summary_table, (str, bytes)):
                 # Accept Dict (canonical) AND List (raw rows)
                 if isinstance(summary_table, list):
                      if not summary_table: 
                           final_grid = None
                      else:
                           # [OPTION B] Strict Fidelity: Ensure List conforms to GridAdapter
                           print(f"Direct Pass-Through of List-Based Summary Table for {issue_id}")
                           # Check if first item is simple value (Flat List) -> Wrap as single col
                           if not isinstance(summary_table[0], (dict, list)):
                                # Convert ["A", "B"] -> [{"col0": {"value": "A"}}, ...]
                                wrapped_rows = []
                                for val in summary_table:
                                     wrapped_rows.append({"col0": {"value": val}})
                                final_grid = {"columns": [{"id": "col0", "label": "Value"}], "rows": wrapped_rows}
                           else:
                                # Safe List-of-Lists or List-of-Dicts
                                final_grid = {'rows': summary_table}

                 elif isinstance(summary_table, dict) and 'rows' in summary_table:
                      print(f"Direct Pass-Through of Summary Table for {issue_id}")
                      final_grid = summary_table
                      
                      # [FIX] Enforce Canonical Column Schema (Prevent Metadata Loss)
                      # If columns are strings, upgrade them immediately.
                      cols = final_grid.get('columns')
                      if isinstance(cols, list) and cols and not isinstance(cols[0], dict):
                          print(f"Detected Legacy String Columns for {issue_id} -> Upgrading to Canonical Schema")
                          new_cols = []
                          for i, col_val in enumerate(cols):
                              # Generate ID: col0, col1...
                              # Use value as Label
                              new_cols.append({"id": f"col{i}", "label": str(col_val)})
                          final_grid['columns'] = new_cols

                 else:
                      # [FIX] Trap for Legacy Scrutiny Dicts (without 'rows') -> Force Normalization
                      print(f"Legacy Scrutiny Table detected for {issue_id} -> Normalizing")
                      try:
                          normalized = self._normalize_summary_table(summary_table)
                          final_grid = self._convert_normalized_to_grid(normalized)
                      except Exception as e:
                          print(f"Legacy Table Conversion Failed: {e}")
                          final_grid = None
            else:
                 try:
                     normalized = self._normalize_summary_table(summary_table)
                     final_grid = self._convert_normalized_to_grid(normalized)
                 except: final_grid = None

        # Priority 2: Existing Grid Data (Canonical) - Fallback for Legacy/Static
        elif grid_data and isinstance(grid_data, dict) and 'rows' in grid_data:
            final_grid = grid_data
        
        # Inject Provenance Metadata (If we have a valid grid)
        if final_grid:
            try:
                import datetime
                source_meta = {
                    "origin": "ASMT10",
                    "asmt_id": self.proceeding_id, 
                    "converted_on": datetime.datetime.now().isoformat()
                }
                source_scrutiny_id = self.proceeding_data.get('source_scrutiny_id') or self.proceeding_data.get('scrutiny_id')
                source_meta["asmt_id"] = source_scrutiny_id
                
                # Injection: Find first available cell in first row
                if final_grid.get('rows'):
                    first_row = final_grid['rows'][0]
                    # Robust injection for List or Dict rows
                    if isinstance(first_row, dict):
                         target_col = "col0"
                         if target_col not in first_row and first_row:
                             target_col = list(first_row.keys())[0]
                         if target_col in first_row:
                             first_row[target_col]["source"] = source_meta
                
                # Demand Calculation Contract
                # [FIX] REMOVED blind injection of tax_demand_mapping.
                # If the template doesn't specify a mapping, we should NOT assume standard variables exist.
                # This allows IssueCard to correctly fall back to heuristics (grid summation) for legacy/list-based issues.
                pass
            except Exception as e:
                print(f"Metadata Injection Failed: {e}")

        factual_tables = final_grid or asmt_data.get('tables') or []


        if factual_tables:
            # Force the table structure into the SCN template
            scn_template['grid_data'] = factual_tables
            print(f"[BRIDGE DIAG] build_scn: Attached grid_data to template for {issue_id}. Type: {type(factual_tables)}")
        else:
            print(f"[BRIDGE DIAG] build_scn: NO factual_tables for {issue_id}")
        
        # 4. Narration Initialization
        # We populate variables with factual data for placeholder resolution
        # Carry forward all variables from the ASMT-10 snapshot (findings)
        raw_snapshot = asmt_data.get('snapshot', {}).copy()
        variables = {}
        
        # CLAIM 2 FIX: Deterministic Variable Mapping
        # Bridge Scrutiny internal keys to clean SCN placeholder names
        SCRUTINY_TO_SCN_MAP = {
            "cgst_as_per_3b": "cgst_3b", "cgst_as_per_1": "cgst_1", "cgst_as_per_2b": "cgst_2b",
            "sgst_as_per_3b": "sgst_3b", "sgst_as_per_1": "sgst_1", "sgst_as_per_2b": "sgst_2b",
            "igst_as_per_3b": "igst_3b", "igst_as_per_1": "igst_1", "igst_as_per_2b": "igst_2b",
            "cess_as_per_3b": "cess_3b", "cess_as_per_1": "cess_1", "cess_as_per_2b": "cess_2b",
            "cgst_diff": "diff_cgst", "sgst_diff": "diff_sgst", "igst_diff": "diff_igst", "cess_diff": "diff_cess"
        }
        
        for k, v in raw_snapshot.items():
            variables[k] = v # Keep original
            if k in SCRUTINY_TO_SCN_MAP:
                variables[SCRUTINY_TO_SCN_MAP[k]] = v # Add mapped alias
        
        # Merge SCN-specific context
        variables.update({
            'brief_facts': asmt_data.get('brief_facts', ""),
            'total_shortfall': asmt_data.get('total_shortfall', 0),
            'issue_name': asmt_data.get('issue_name', asmt_data.get('issue', 'Issue'))
        })
        
        # Ensure converted grid variables are also in the variables dict
        if grid_data and isinstance(grid_data, dict) and 'rows' in grid_data:
             for row in grid_data['rows']:
                 # Canonical Row is {"col0": {val, type}, "col1": ...}
                 for cell_key, cell in row.items():
                     if isinstance(cell, dict):
                         var = cell.get('var')
                         if var and 'value' in cell:
                             variables[var] = cell['value']
        
        # [BRIDGE DIAG] Confirm return payload integrity
        
        return {
            'template': scn_template,
            'data': {
                'issue_id': issue_id,
                'origin': 'ASMT10',
                'table_data': grid_data or factual_tables, # Prefer conversion result
                'variables': variables,
                'status': 'ACTIVE',
                'source_issue_id': issue_id
            }
        }

    def is_scn_finalized(self):
        """Check if SCN is in finalized state (Read-Only Mode)"""
        current_stage = self.get_current_stage()
        return current_stage >= WorkflowStage.SCN_ISSUED

    def _hydrate_cards_from_records(self, records, is_initial_hydration=False):
        """Authoritative card factory. No fallbacks, heuristics, or shared help paths."""
        
        # Fetch case-level provenance for validation
        case_source_id = self.proceeding_data.get('source_scrutiny_id') or self.proceeding_data.get('scrutiny_id')
        
        for i, record in enumerate(records):
            if is_initial_hydration:
                try:
                    # Use Adapter for fresh ASMT-10 adoption
                    adapted = self.build_scn_issue_from_asmt10(record)
                    card = self.add_scn_issue_card(
                        template=adapted['template'],
                        data=adapted['data'],
                        origin="ASMT10",
                        source_id=record['issue_id'],
                        trigger_eval=False
                    )
                    if card: card.on_grid_data_adopted()
                except Exception as e:
                    print(f"[HYDRATION ERROR] Failed to load card {i} (ID: {record.get('issue_id')}): {e}")
                    import traceback
                    traceback.print_exc()
            else:
                # Restoration of existing SCN draft (SCN stage) - STRICT SNAPSHOT HYDRATION
                data_payload = record.get('data', {})
                issue_id = record['issue_id']
                
                # Rule: NEVER consult Issue Master for restored issues.
                # Use the template snapshot embedded in the data payload.
                # Extract values from payload structure
                if 'values' in data_payload:
                     # New Structure
                     restore_data = {
                         'variables': data_payload.get('values'),
                         'table_data': data_payload.get('table_data')
                     }
                else:
                     # Legacy Structure (Direct props)
                     restore_data = data_payload
                
                template_snapshot = data_payload.get('template_snapshot')
                     
                # [FIX] Hydration Data Flow: Authoritative Master Merge
                # Rule: Consult Master (if available) but preserve USER VALUES from snapshot.
                master_template = self.db.get_issue(issue_id) or self.db.get_issue_by_name(issue_id)
                
                if master_template:
                     # 1. Base is Master (Identity + Static Text + Logic + LIABILITY CONFIG)
                     template = copy.deepcopy(master_template)

                     # 2. Overlay Snapshot State (The "Values", NOT the "Structure")
                     if template_snapshot:
                          # [DEEP MERGE] Preserve Master Structure (var tags) but Adopt Snapshot Values
                          snapshot_grid = template_snapshot.get('grid_data')
                          if snapshot_grid and template.get('grid_data'):
                               master_grid = template['grid_data']
                               
                               # Polymorphic Normalize: Handle List (Legacy) vs Dict (Modern)
                               s_rows = snapshot_grid.get('rows', []) if isinstance(snapshot_grid, dict) else (snapshot_grid if snapshot_grid else [])
                               m_rows = master_grid.get('rows', []) if isinstance(master_grid, dict) else master_grid
                               m_cols = master_grid.get('columns', []) if isinstance(master_grid, dict) else []
                               
                               # [RESILIENCE] Helper to extract value from dict or primitive
                               def safe_get_value(cell):
                                   if isinstance(cell, dict): return cell.get('value', 0)
                                   return cell if cell is not None else 0
                               

                               
                               row_policy = master_grid.get('row_policy', 'fixed')

                               for i, s_row in enumerate(s_rows):
                                   m_row = None

                                   if i < len(m_rows):
                                       m_row = m_rows[i]
                                   elif row_policy == 'dynamic' and len(m_rows) > 0:
                                       # [FIX] Expand Master for Dynamic Rows
                                       prototype = m_rows[0]
                                       if isinstance(prototype, dict):
                                           try:
                                               new_row = copy.deepcopy(prototype)
                                               new_row['id'] = f"r{i+1}_hydrated"
                                               m_rows.append(new_row)
                                               m_row = new_row
                                           except Exception as e:
                                               print(f"[HYDRATION ERROR] Failed to clone dynamic row: {e}")

                                   if m_row:
                                       
                                       # Scenario A: Modern Dict-based Row
                                       if isinstance(m_row, dict):
                                           if isinstance(s_row, dict):
                                               # Match by Column ID with Positional Fallback
                                               for cid, s_cell in s_row.items():
                                                   target_cid = None
                                                   if cid in m_row:
                                                       target_cid = cid
                                                   elif cid.startswith('col') and cid[3:].isdigit():
                                                       # Fallback: Positional mapping for legacy col0, col1, etc.
                                                       col_idx = int(cid[3:])
                                                       if col_idx < len(m_cols):
                                                           target_cid = m_cols[col_idx].get('id') if isinstance(m_cols[col_idx], dict) else str(m_cols[col_idx]).lower().replace(" ", "_")
                                                   
                                                   if target_cid and target_cid in m_row and isinstance(m_row[target_cid], dict):
                                                       m_row[target_cid]['value'] = safe_get_value(s_cell)
                                           elif isinstance(s_row, list):
                                               # Match Legacy List-Row by Index using Master Column IDs
                                               for col_idx, s_cell in enumerate(s_row):
                                                   if col_idx < len(m_row):
                                                       m_col = m_cols[col_idx]
                                                       cid = m_col.get('id') if isinstance(m_col, dict) else str(m_col).lower().replace(" ", "_")
                                                       if cid in m_row and isinstance(m_row[cid], dict):
                                                            m_row[cid]['value'] = safe_get_value(s_cell)
                                       
                                       # Scenario B: Legacy List-based Row
                                       elif isinstance(m_row, list) and isinstance(s_row, list):
                                           for col_idx, s_cell in enumerate(s_row):
                                               if col_idx < len(m_row):
                                                   if isinstance(m_row[col_idx], dict):
                                                       m_row[col_idx]['value'] = safe_get_value(s_cell)
                               
                               # [CRITICAL] Update restore_data so IssueCard uses the merged structure/values
                               if isinstance(master_grid, dict):
                                   restore_data['table_data'] = copy.deepcopy(master_grid)
                               else:
                                   # Master is list (Legacy fallback)
                                   restore_data['table_data'] = copy.deepcopy(master_grid)

                          # Adoption of variables (if any fresh overrides existed)
                          if 'variables' in template_snapshot:
                               if not template.get('variables'): template['variables'] = {}
                               # Only adopt non-null values
                               for kv, vv in template_snapshot.get('variables', {}).items():
                                   if vv is not None and vv != "":
                                       template['variables'][kv] = vv
                else:
                     # Fallback for ad-hoc issues
                     template = template_snapshot
                     if not template:
                          print(f"WARNING: No template snapshot or master for {issue_id}.")

                if not template:
                     template = {
                        'issue_id': issue_id, 
                        'issue_name': data_payload.get('issue', 'Issue'),
                        'variables': {}
                     }
                
                
                     template = {
                        'issue_id': issue_id, 
                        'issue_name': data_payload.get('issue', 'Issue'),
                        'variables': {}
                     }
                
                # --- GLOBAL PROVENANCE LOGIC (System-Wide) ---
                current_origin = record.get('origin', 'SCRUTINY')
                
                source_id = record.get('source_issue_id') 
                source_proc_id = record.get('source_proceeding_id')
                source_type = record.get('source_proceeding_type', 'SCRUTINY') # Assume scrutiny if typeless legacy
                
                # GLOBAL RULE: Trust Verification Metadata only. 
                # If we have a valid source link (ID + ProcID), we honor it.
                if current_origin == "SCN" and source_id and source_proc_id and source_type in ("SCRUTINY", "ASMT10"):
                     # Double Check: Does source_proc_id match THIS case's lineage?
                     if source_proc_id == case_source_id:
                          print(f"Hydration Repair: Provenance-based Auto-Correction SCN -> {source_type} for {issue_id}")
                          current_origin = source_type
                
                
                # [HYDRATION REPAIR] Self-Healing for "Hollow" Templates
                # Scenario: Snapshot lacks grid_data (e.g., initial draft was saved before table render).
                # Fix: Re-fetch authoritative table from ASMT-10 source if available.
                
                is_hollow = (
                    current_origin in ["ASMT10", "SCRUTINY"] and 
                    not template.get('grid_data') and 
                    not template.get('tables') and
                    not restore_data.get('table_data')
                )
                
                if is_hollow and source_id and source_proc_id:
                     print(f"[HYDRATION REPAIR] Detected HOLLOW template for {issue_id}. Attempting repair from {source_proc_id}...")
                     
                     # 1. Fetch Source Context
                     try:
                         # Performance Note: This fetches case slice. Acceptable for repair scenario.
                         source_records = self.db.get_case_issues(source_proc_id, stage='DRC-01A')
                         source_record = next((r for r in source_records if r['issue_id'] == source_id), None)
                         
                         if source_record:
                             # 2. Regenerate Fresh Template/Data from Source
                             print(f"  - Source record found. Re-building...")
                             repair_package = self.build_scn_issue_from_asmt10(source_record)
                             
                             fresh_template = repair_package['template']
                             fresh_grid = fresh_template.get('grid_data')
                             
                             if fresh_grid:
                                  print(f"  - REPAIR SUCCESS: Injected {len(fresh_grid.get('rows', []))} rows.")
                                  
                                  # 3. Surgical Injection (Preserve User's Draft, Repair Structure)
                                  template['grid_data'] = fresh_grid
                                  
                                  # Also update restore_data to ensure load_data picks it up
                                  restore_data['table_data'] = fresh_grid
                             else:
                                  print("  - Repair Warning: Source also has no grid_data.")
                         else:
                             print("  - Repair Failed: Source record not found.")
                     except Exception as e:
                         print(f"  - Repair Error: {e}")

                print(f"[BRIDGE DIAG] hydrate_scn: Preparing card {issue_id}. restore_data[table_data] type: {type(restore_data.get('table_data'))}")
                
                # [FIX] JIT Repair for Legacy Snapshots (Failed Strictness Check)
                # Existing snapshots in DB may have string-columns. Upgrade them NOW.
                for target in [template, restore_data]:
                    tgt_grid = None
                    if target == template: tgt_grid = target.get('grid_data')
                    else: tgt_grid = target.get('table_data')
                    
                    if tgt_grid and isinstance(tgt_grid, dict):
                        cols = tgt_grid.get('columns')
                        if isinstance(cols, list) and cols and not isinstance(cols[0], dict):
                             print(f"  - JIT Repair: Upgrading string columns for {issue_id}")
                             new_cols = [{"id": f"col{i}", "label": str(c)} for i, c in enumerate(cols)]
                             tgt_grid['columns'] = new_cols

                if template:
                    print(f"  - Template has grid_data? {'grid_data' in template}")
                
                card = self.add_scn_issue_card(
                    template=template,
                    data=restore_data, 
                    origin=current_origin,
                    status=data_payload.get('status', 'ACTIVE'),
                    source_id=source_id,
                    trigger_eval=False
                )
                if card: card.on_grid_data_adopted()

        # Batch Evaluation: Trigger phase check ONLY after all cards are hydrated
        self.evaluate_scn_workflow_phase()

    def reset_scn_adoption(self):
        """
        Reset Adoption Strategy (Redefined).
        Action: Remove ONLY issues with origin='SCRUTINY'.
        Constraint: NEVER delete origin='MANUAL_SOP'.
        Constraint: NEVER re-query ASMT-10. SCN Snapshot is authoritative.
        """
        if not self.proceeding_id: return

        # Guard: Phase-1 Finalization
        if self.is_scn_finalized():
             QMessageBox.warning(self, "Action Blocked", "Show Cause Notice is already finalized.")
             return

        # 1. Confirm Intent
        # We are purging adopted issues. Manual issues stay.
        msg = "This will remove all issues adopted from Scrutiny (Origin: SCRUTINY).\n\n"
        msg += "Manually added SOP issues will be PRESERVED.\n"
        msg += "To restore Scrutiny issues later, use the 'Adoption' dropdown.\n\n"
        msg += "Proceed with removal?"
        
        reply = QMessageBox.question(self, "Remove Scrutiny Issues", msg,
                                   QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
                                   
        if reply != QMessageBox.StandardButton.Yes:
            return

        # 2. Perform Selective Removal
        # Iterate backwards to safely remove from list
        removed_count = 0
        preserved_count = 0
        
        # Snapshot copy of list to avoid modification errors during iteration
        for card in list(self.scn_issue_cards):
            origin = getattr(card, 'origin', 'SCRUTINY')
            
            if origin == 'SCRUTINY' or origin == 'ASMT10':
                # Remove this card
                # Use existing removal logic but suppress heavy save per item if possible
                # For now, standard removal is safer to ensure UI sync
                self.remove_scn_issue_card(None, card)
                removed_count += 1
            else:
                preserved_count += 1
                
        # 3. Final Persistence & Feedback
        # remove_scn_issue_card triggers save, so DB is already updated.
        
        QMessageBox.information(self, "Action Complete", 
                              f"Removed {removed_count} Scrutiny issues.\nPreserved {preserved_count} Manual issues.")
        
        # 4. Refresh Dropdown immediately to make adopted issues available again
        self.load_scn_issue_templates()

    def _persist_scn_init_flag(self):
        """Authoritative, minimal persistence for SCN initialization flag only. UI-only side effect."""
        try:
            # [PHASE 15] Safe Merge
            details = copy.deepcopy(self.proceeding_data.get('additional_details', {}))
            if 'scn_metadata' not in details: details['scn_metadata'] = {}
            
            details['scn_metadata']['scn_issues_initialized'] = True
            # For backward compatibility during migration
            details['scn_issues_initialized'] = True 
            
            # Persist to DB directly
            self.db.update_proceeding(self.proceeding_id, {
                'additional_details': details
            })
            self.proceeding_data['additional_details'] = details
            print(f"ProceedingsWorkspace: SCN initialization flag persisted for {self.proceeding_id}")
        except Exception as e:
            print(f"Error persisting SCN init flag: {e}")

    def sync_demand_tiles(self):
        """Sync Demand Tiles with Issues"""
        try:
            # Clear existing tiles
            while self.demand_tiles_layout.count():
                item = self.demand_tiles_layout.takeAt(0)
                widget = item.widget()
                if widget:
                    widget.deleteLater()
            self.demand_tiles = []
            
            # Iterate through SCN issue cards
            tax_numerals = []
            
            for i, card in enumerate(self.scn_issue_cards, 1):
                template = card.template
                issue_id = template.get('issue_id', 'unknown')
                issue_name = template.get('issue_name', f'Issue {i}')

                # [DIAGNOSTIC]
                print(f"SYNC ISSUE INDEX: {i}")
                print(f"SYNC ISSUE ID: {issue_id}")

                # [REFRESH] Force state synchronization before generating demand text
                # This ensures liability rows are extracted from grid into variables
                card.calculate_values() 
                data = card.get_data() # Fresh snapshot
                print(f"BREAKDOWN SENT TO GENERATOR (Issue {i}):", data.get('tax_breakdown'))

                # Generate Tax Text (Clause i/ii/iii)
                text = self.db.generate_single_issue_demand_text(data, i)
                # Extract numeral (naive but sufficient for now)
                roman = self.db.to_roman(i).lower()
                tax_numerals.append(f"({roman})")
                
                # Create Tile
                tile = ModernCard(f"And whereas - {issue_name} (Clause {roman})", collapsible=True)
                
                # Editor
                editor = RichTextEditor()
                formatted_text = text.replace('\n', '<br>')
                editor.setHtml(formatted_text)
                editor.setMinimumHeight(100)
                

                
                tile.addWidget(editor)
                self.demand_tiles_layout.addWidget(tile)
                
                self.demand_tiles.append({
                    'issue_id': issue_id,
                    'type': 'TAX',
                    'card': tile,
                    'editor': editor
                })
            
            # Create Consolidated Interest Tile
            if tax_numerals:
                refs = ", ".join(tax_numerals[:-1]) + " and " + tax_numerals[-1] if len(tax_numerals) > 1 else tax_numerals[0]
                interest_roman = self.db.to_roman(len(self.scn_issue_cards) + 1).lower()
                
                int_text = f"{interest_roman}. Interest at an appropriate rate under Section 50 of the CGST Act 2017 on the amount demanded at Para No. {refs} as mentioned above should not be demanded and recovered from them under Section 73(1) of the CGST Act 2017 and corresponding Section under Kerala SGST Act, 2017, read with Section 20 of the IGST Act, 2017;"
                
                int_tile = ModernCard(f"Interest Demand (Clause {interest_roman})", collapsible=True)
                int_editor = RichTextEditor()
                int_editor.setHtml(int_text)
                int_editor.setMinimumHeight(80)
                

                int_tile.addWidget(int_editor)
                self.demand_tiles_layout.addWidget(int_tile)
                self.demand_tiles.append({'issue_id': 'INTEREST_GLOBAL', 'type': 'INTEREST', 'card': int_tile, 'editor': int_editor})

                # Create Consolidated Penalty Tile
                pen_roman = self.db.to_roman(len(self.scn_issue_cards) + 2).lower()
                pen_text = f"{pen_roman}. Penalty should not be imposed on them under the provision of Section 73 (1) of CGST Act 2017 read with Section 122 (2) (a) of CGST Act, 2017 and corresponding section under the Kerala SGST Act, 2017 read with section 20 of the IGST Act, 2017, for the contraventions referred hereinabove."
                
                pen_tile = ModernCard(f"Penalty Demand (Clause {pen_roman})", collapsible=True)
                pen_editor = RichTextEditor()
                pen_editor.setHtml(pen_text)
                pen_editor.setMinimumHeight(80)
                

                pen_tile.addWidget(pen_editor)
                self.demand_tiles_layout.addWidget(pen_tile)
                self.demand_tiles.append({'issue_id': 'PENALTY_GLOBAL', 'type': 'PENALTY', 'card': pen_tile, 'editor': int_editor})
                
            
            
        except Exception as e:
            print(f"Error syncing demand tiles: {e}")
            traceback.print_exc()
                

            QMessageBox.warning(self, "Error", f"Failed to generate demand text: {e}")

    def save_document(self, doc_type="SCN", is_manual=False):
        """Save SCN document with authoritative structural integrity"""
        # Phase-1 Isolation: Structurally block save during Step-2 (Adoption)
        if doc_type == "SCN" and self.is_scn_phase1():
             print("ProceedingsWorkspace: Autopersist blocked in Step-2 Adoption phase.")
             return

        if not self.proceeding_id:
            return
            
        print(f"ProceedingsWorkspace: Saving {doc_type} draft (is_manual={is_manual})...")
        
        try:
            if doc_type == "SCN":
                # Aggregate issue data using the authoritative schema (New Persistence)
                self.persist_scn_issues(is_manual_save=is_manual)
                
                # Aggregate Demand Text from Tiles
                full_demand_text = ""
                demand_tiles_data = []
                
                if hasattr(self, 'demand_tiles') and self.demand_tiles:
                    for tile in self.demand_tiles:
                        editor = tile['editor']
                        issue_id = tile['issue_id']
                        content = editor.toHtml()
                        full_demand_text += content + "<br><br>"
                        demand_tiles_data.append({'issue_id': issue_id, 'content': content})
                else:
                    # Fallback if no tiles (shouldn't happen with new UI)
                     full_demand_text = "<p>No demand details generated.</p>"

                metadata = {
                    "scn_number": self.scn_no_input.text(),
                    "scn_oc_number": self.scn_oc_input.text(),
                    "scn_date": self.scn_date_input.date().toString("yyyy-MM-dd"),
                    "reliance_documents": self.reliance_editor.toHtml(),
                    "copy_submitted_to": self.copy_to_editor.toHtml(),
                    "demand_text": full_demand_text,
                    "demand_tiles_data": demand_tiles_data,
                    "scn_issues_initialized": self.scn_issues_initialized,
                    "scn_model_snapshot": self._get_scn_model() # Authoritative point-in-time snapshot
                }
                
                # [PHASE 15] Safe Merge with deep copy
                details = copy.deepcopy(self.proceeding_data.get('additional_details', {}))
                details['scn_metadata'] = metadata
                
                self.db.update_proceeding(self.proceeding_id, {
                    "status": "SCN Draft",
                    "additional_details": details
                })
                
                # Update local data
                self.proceeding_data['additional_details'] = details
                
                QMessageBox.information(self, "Success", "Show Cause Notice draft saved successfully!")
                return True
                
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error saving SCN: {e}")
            print(f"Error saving SCN: {e}")
            return False

    def on_manual_draft_save(self):
        """Centralized manual save handler with non-intrusive feedback"""
        # 1. Structural Validation + Persistence
        self.save_document("SCN", is_manual=True)
        
        # 2. Feedback Mechanism (Status Bar or Toast Overlay)
        from datetime import datetime
        timestamp = datetime.now().strftime("%H:%M:%S")
        msg = f"Draft saved at {timestamp}"
        
        # If we have a status bar, use it. Otherwise, use a transient label.
        if hasattr(self, 'statusBar') and self.statusBar():
            self.statusBar().showMessage(msg, 5000)
        else:
            print(f"SCN Governance: {msg}")
            # Fallback toast if needed (simplified for high-integrity)
            QToolTip.showText(QCursor.pos(), msg, self, QRect(), 3000)

    def _create_save_draft_btn(self):
        """Standardized factory for Phase B Save Draft buttons"""
        btn = QPushButton("Save Draft")
        btn.setObjectName("BtnSaveDraft")
        btn.setCursor(Qt.CursorShape.PointingHandCursor)
        btn.setStyleSheet("""
            QPushButton#BtnSaveDraft {
                background-color: #f8fafc;
                border: 1px solid #cbd5e1;
                color: #475569;
                padding: 8px 15px;
                border-radius: 4px;
                font-weight: 500;
                font-size: 9pt;
            }
            QPushButton#BtnSaveDraft:hover {
                background-color: #f1f5f9;
                border-color: #3498db;
                color: #3498db;
            }
        """)
        btn.clicked.connect(self.on_manual_draft_save)
        return btn

    def _find_template_for_scrutiny_item(self, item, template_map):
        """Heuristic to find a matching template for a scrutiny result item"""
        # 1. Try if template_id is directly in item (for future-proofing)
        if 'template_id' in item and item['template_id'] in template_map:
            return template_map[item['template_id']]
            
        # 2. Try matching by category
        cat = str(item.get('category', '')).lower()
        desc = str(item.get('description', '')).lower()
        
        for tid, t in template_map.items():
            name = str(t.get('issue_name', '')).lower()
            if name == cat or name == desc:
                return t
                
        # 3. Fallback to generic
        return {'issue_id': 'GENERIC_ISSUE', 'issue_name': item.get('description') or item.get('category'), 'variables': {}}

    def _show_blocking_msg(self, message):
        """Show a blocking message in Step 2 for empty states"""
        # Clear existing cards
        while self.scn_issues_layout.count():
            item = self.scn_issues_layout.takeAt(0)
            widget = item.widget()
            if widget: widget.deleteLater()
            
        lbl = QLabel(message)
        lbl.setStyleSheet("color: #e74c3c; font-weight: bold; font-size: 12pt; margin: 50px;")
        lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.scn_issues_layout.addWidget(lbl)


    def open_asmt10_reference(self):
        """Open ASMT-10 in a modal overlay for reference"""
        source_id = self.proceeding_data.get('source_scrutiny_id') or self.proceeding_data.get('scrutiny_id')
        if not source_id:
            QMessageBox.information(self, "Reference", "No source ASMT-10 linked to this case.")
            return
            
        # Generate HTML (Simplified reuse of generator)
        scrutiny_data = self.db.get_scrutiny_case_data(source_id)
        if not scrutiny_data:
            QMessageBox.warning(self, "Reference", "Source scrutiny data could not be retrieved.")
            return
            
        case_info = {
            'case_id': scrutiny_data.get('case_id'),
            'financial_year': scrutiny_data.get('financial_year'),
            'section': scrutiny_data.get('section'),
            'notice_date': scrutiny_data.get('asmt10_finalised_on', '-'),
            'oc_number': scrutiny_data.get('oc_number', 'DRAFT'),
            'last_date_to_reply': scrutiny_data.get('last_date_to_reply', 'N/A')
        }
        taxpayer = scrutiny_data.get('taxpayer_details', {})
        issues = scrutiny_data.get('selected_issues', [])
        
        generator = ASMT10Generator()
        full_data = case_info.copy()
        full_data['taxpayer_details'] = taxpayer
        html = generator.generate_html(full_data, issues, for_preview=True)
        
        dlg = ASMT10ReferenceDialog(self, html)
        dlg.exec()


    def _get_scn_model(self):
        """Build a structured SCN data model for consistent rendering across Preview, PDF, and DOCX."""
        if not self.proceeding_data:
            return None

        # 1. Base Metadata
        model = self.proceeding_data.copy()
        
        # [FIX] Break Circular Reference: scn_model_snapshot -> additional_details -> scn_model_snapshot
        if 'additional_details' in model:
            # Shallow copy the inner dict so we can modify it without affecting self.proceeding_data
            details_copy = model['additional_details'].copy() if isinstance(model['additional_details'], dict) else {}
            # Remove the self-referential snapshot
            details_copy.pop('scn_model_snapshot', None)
            model['additional_details'] = details_copy
        
        # Format Dates
        scn_date = self.scn_date_input.date()
        model['issue_date'] = scn_date.toString("dd/MM/yyyy")
        model['year'] = scn_date.year()
        
        # Dynamic Financial Year (based on SCN Date)
        # If Month >= 4 (April), FY is Year-(Year+1). Else (Year-1)-Year.
        if scn_date.month() >= 4:
            fy_start = scn_date.year()
            fy_end = (scn_date.year() + 1) % 100
        else:
            fy_start = scn_date.year() - 1
            fy_end = scn_date.year() % 100
        model['current_financial_year'] = f"{fy_start}-{fy_end:02d}"
        
        # OC & SCN No
        model['oc_no'] = self.scn_oc_input.text() or "____"
        model['scn_no'] = self.scn_no_input.text() or "____"
        model['initiating_section'] = model.get('adjudication_section') or model.get('initiating_section', '') or "____"
        model['section'] = model['initiating_section'] # Specific key for template
        
        # Taxpayer Details
        tp = model.get('taxpayer_details', {})
        if tp is None: tp = {}
        if isinstance(tp, str): 
            try:
                import json
                tp = json.loads(tp)
            except:
                tp = {}
        
        model['legal_name'] = tp.get('Legal Name', '') or model.get('legal_name', '')
        model['trade_name'] = tp.get('Trade Name', '') or model.get('trade_name', '')
        model['address'] = tp.get('Address', '') or model.get('address', '')
        model['gstin'] = model.get('gstin', '')
        model['constitution_of_business'] = tp.get('Constitution of Business', 'Registered')
        
        # [NEW] Officer Details (Snapshot Priority -> Database Query -> Fallback)
        officer_name = "________________"
        officer_desg = "________________"
        officer_juris = "________________"
        
        # Priority 1: Adjudication Case specific snapshot (if origin is ADJUDICATION)
        # Priority 2: Proceedings base snapshot (if origin is SCRUTINY)
        
        details = model.get('additional_details', {})
        if isinstance(details, str):
            try: details = json.loads(details)
            except: details = {}
            
        snapshot_json = model.get('issuing_officer_snapshot')
        officer_id = model.get('issuing_officer_id')
        
        # Attempt recovery from JSON details if main columns missing (transitional safety)
        if not snapshot_json and 'issuing_officer_snapshot' in details:
            snapshot_json = details['issuing_officer_snapshot']
        if not officer_id and 'issuing_officer_id' in details:
            officer_id = details['issuing_officer_id']
            
        if snapshot_json:
            try:
                import json
                parsed_snap = json.loads(snapshot_json)
                
                # Check for isolated SCN snapshot
                if 'SCN' in parsed_snap:
                    off_data = parsed_snap['SCN']
                elif 'name' in parsed_snap:
                    off_data = parsed_snap # Fallback for old flat format
                else:
                    off_data = {}
                    
                officer_name = off_data.get('name', officer_name)
                officer_desg = off_data.get('designation', officer_desg)
                officer_juris = off_data.get('jurisdiction', officer_juris)
            except json.JSONDecodeError:
                pass
        elif officer_id:
            # Fallback to live query
            off_data = self.db.get_officer_by_id(officer_id)
            if off_data:
                officer_name = off_data.get('name', officer_name)
                officer_desg = off_data.get('designation', officer_desg)
                officer_juris = off_data.get('jurisdiction', officer_juris)

        model['officer_name'] = officer_name
        model['officer_designation'] = officer_desg
        model['designation'] = officer_desg
        model['jurisdiction'] = officer_juris

        # 2. Sequential Issues Handling
        included_issues = []
        if hasattr(self, 'scn_issue_cards'):
            # Filter strictly by is_included
            for card in self.scn_issue_cards:
                if card.get_data().get('is_included', True):
                    included_issues.append(card)

        # Build issue list for rendering
        issues_data = []
        total_tax = 0
        igst_total = 0
        cgst_total = 0
        sgst_total = 0
        
        # Mapping for narrative replacement
        id_to_index = {}
        for idx, card in enumerate(included_issues, start=1):
            id_to_index[card.issue_id] = str(idx)

        import re
        for idx, card in enumerate(included_issues, start=1):
            issue_info = {
                'index': idx,
                'title': card.display_title,
                'issue_id': card.issue_id, # Keep for internal use, but suppress in display
                'paras': [],
                'table_html': card.generate_table_html(card.template, card.variables)
            }
            
            # Narrative Content Handling with Regex Replacement
            editor_part = card.editor.toHtml()
            editor_part = re.sub(r'<p>\s*&nbsp;\s*</p>', '', editor_part)
            editor_part = re.sub(r'<p>\s*</p>', '', editor_part)
            paras = re.findall(r'<p.*?>(.*?)</p>', editor_part, re.DOTALL)
            if not paras and editor_part.strip():
                paras = [editor_part]

            for p_content in paras:
                if p_content.strip():
                    clean_content = re.sub(r'\s+', ' ', p_content).strip()
                    # REGEX REPLACE internal IDs with sequential numbers
                    for internal_id, seq_num in id_to_index.items():
                        # Word boundary regex to avoid partial replacements
                        pattern = r'\b' + re.escape(internal_id) + r'\b'
                        clean_content = re.sub(pattern, f"issue {seq_num}", clean_content)
                    
                    issue_info['paras'].append(clean_content)
            
            issues_data.append(issue_info)
            
            # Totals Aggregation (Structured Model)
            card_data = card.get_data()
            card_breakdown = card.get_tax_breakdown() # Returns integer mapping
            
            for act, vals in card_breakdown.items():
                v_tax = vals.get('tax', 0)
                v_int = vals.get('interest', 0)
                v_pen = vals.get('penalty', 0)
                
                total_tax += v_tax
                if act == 'IGST': igst_total += v_tax
                elif act == 'CGST': cgst_total += v_tax
                elif act == 'SGST': sgst_total += v_tax
        
        # 3. Aggregate Tax Summary for the Table (Act-wise)
        total_breakdown = {}
        for card in included_issues:
            card_breakdown = card.get_tax_breakdown()
            for act, vals in card_breakdown.items():
                if act not in total_breakdown:
                    total_breakdown[act] = {'tax': 0, 'interest': 0, 'penalty': 0, 'total': 0}
                
                t = vals.get('tax', 0)
                i = vals.get('interest', 0)
                p = vals.get('penalty', 0)
                
                total_breakdown[act]['tax'] += t
                total_breakdown[act]['interest'] += i
                total_breakdown[act]['penalty'] += p
                total_breakdown[act]['total'] += (t + i + p)

        # Build tax_rows for the generator
        tax_rows = []
        for act in self.ACT_PRIORITY:
            if act in total_breakdown:
                vals = total_breakdown[act]
                tax_rows.append({
                    'act': act,
                    'period': model.get('financial_year', ''),
                    'tax': vals['tax'],
                    'interest': vals['interest'],
                    'penalty': vals['penalty'],
                    'total': vals['total']
                })

        model['issues'] = issues_data
        model['total_tax_val'] = total_tax
        model['igst_total_val'] = igst_total
        model['cgst_total_val'] = cgst_total
        model['sgst_total_val'] = sgst_total
        
        # Helper: Indian Currency Format (Hardened)
        def format_indian_currency(value):
            if value is None: return "0"
            try:
                val = safe_int(value)
            except (ValueError, TypeError):
                return str(value)
            
            is_negative = val < 0
            val = abs(val)
            
            # Format to 2 decimals first to handle rounding
            s_val = f"{val:.2f}"
            
            # Extract integer and decimal parts
            if "." in s_val:
                integer_part, decimal_part = s_val.split(".")
            else:
                integer_part, decimal_part = s_val, "00"
            
            # Drop decimal if zero
            if decimal_part == "00":
                decimal_suffix = ""
            else:
                decimal_suffix = "." + decimal_part
                
            # Apply commas to integer part
            if len(integer_part) > 3:
                last3 = integer_part[-3:]
                rest = integer_part[:-3]
                # split rest into chunks of 2 from right
                parts = []
                while len(rest) > 2:
                    parts.insert(0, rest[-2:])
                    rest = rest[:-2]
                parts.insert(0, rest)
                formatted_int = ",".join(parts) + "," + last3
            else:
                formatted_int = integer_part
                
            result = formatted_int + decimal_suffix
            if is_negative:
                result = "-" + result
            return result

        # Formatted totals using Indian Format
        model['total_amount'] = format_indian_currency(total_tax)
        model['igst_total'] = format_indian_currency(igst_total)
        model['cgst_total'] = format_indian_currency(cgst_total)
        model['sgst_total'] = format_indian_currency(sgst_total)

        # 3. Dynamic Paragraph Numbering
        # Intro=Para 1, Jurisdiction=Para 2, Issues start at Para 3
        # If we have N issues, they occupy Para 3 to 3+(N-1)
        next_para = 3 + len(issues_data)
        model['para_demand'] = next_para
        model['para_waiver'] = next_para + 1
        model['para_cancellation'] = next_para + 2
        model['para_hearing'] = next_para + 3
        model['para_exparte'] = next_para + 4
        model['para_prejudice'] = next_para + 5
        model['para_amendment'] = next_para + 6
        model['para_reliance'] = next_para + 7

        # 4. Additional Data (Reliance, Copy To)
        rel_text = self.reliance_editor.toPlainText()
        model['reliance_documents'] = [line for line in rel_text.split('\n') if line.strip()]
        
        copy_text = self.copy_to_editor.toPlainText()
        model['copy_submitted_to'] = [line for line in copy_text.split('\n') if line.strip()]
        
        model['show_letterhead'] = self.show_letterhead_cb.isChecked() if hasattr(self, 'show_letterhead_cb') else False
        
        # [NEW] Introductory Narrative (Grounds)
        try:
            from src.utils.scn_generator import generate_intro_narrative
            details = self.proceeding_data.get('additional_details', {})
            # Ensure details is dict
            if isinstance(details, str):
                import json
                try: details = json.loads(details)
                except: details = {}
                
            grounds_data = details.get('scn_grounds')
            
            # [PHASE 18] Generation-Time Linkage (Strict Overlay)
            # Ensure SCN always cites the authoritative ASMT-10 Ref from Case Data
            import copy
            if grounds_data:
                # 1. Deepcopy to prevent mutation of stored JSON
                gen_view = copy.deepcopy(grounds_data)
                
                # 2. Overlay Authoritative Identifiers
                auth_oc = self.proceeding_data.get('oc_number', '')
                auth_date = self.proceeding_data.get('notice_date', '')
                
                if 'data' in gen_view:
                    if 'asmt10_ref' not in gen_view['data']:
                        gen_view['data']['asmt10_ref'] = {}
                        
                    gen_view['data']['asmt10_ref']['oc_no'] = auth_oc
                    gen_view['data']['asmt10_ref']['date'] = auth_date
                
                # 3. Generate Narrative with guaranteed correct IDs
                model['intro_narrative'] = generate_intro_narrative(gen_view)
            else:
                model['intro_narrative'] = ""
        except Exception as e:
            print(f"Narrative Gen Error: {e}")
            model['intro_narrative'] = "<b>[ERROR GENERATING INTRODUCTORY NARRATIVE]</b>"
        
        # Demand Text (Loaded from Tiles)
        if hasattr(self, 'demand_tiles') and self.demand_tiles:
             full_text = ""
             for tile in self.demand_tiles:
                 full_text += tile['editor'].toHtml() + "<br><br>"
             model['demand_text'] = full_text
        else:
             model['demand_text'] = "<p>No demand details generated.</p>"

        # 4. Tax Table HTML (Data-Driven)
        model['tax_table_html'] = self.generate_tax_table_html(tax_rows)

        return model
    def render_scn(self, is_preview=False, for_pdf=False):
        """Render SCN HTML using Jinja2 template"""
        model = self._get_scn_model()
        if not model:
            return "<h3>No Case Data Loaded</h3>"
            
        model['is_preview'] = is_preview
        
        try:
            # 1. Format Issues HTML from the centralized model
            issues_html = ""
            for issue in model['issues']:
                current_para_num = issue['index'] + 2 # Paras 1&2 are fixed, Issues start at 3
                title = issue['title']
                
                # No wrapper for the whole issue to avoid pagination truncation of multi-block content
                
                # Issue Heading
                issues_html += f"""
                <div class="issue-header">
                    <table class="para-table">
                        <tr>
                            <td class="para-num">{current_para_num}.</td>
                            <td class="para-content"><h3>Issue No. {issue['index']}: {title}</h3></td>
                        </tr>
                    </table>
                </div>
                """
                
                # Sub Paragraphs
                sub_para_count = 1
                for p_content in issue['paras']:
                    num_str = f"{current_para_num}.{sub_para_count}"
                    issues_html += f"""
                    <table class="para-table">
                        <tr>
                            <td class="para-num">{num_str}</td>
                            <td class="para-content">
                                <p class="legal-para">{p_content}</p>
                            </td>
                        </tr>
                    </table>
                    """
                    sub_para_count += 1
                
                # Table as Sub-Para (Unwrapped for pagination splitting)
                if issue['table_html'] and issue['table_html'].strip():
                    num_str = f"{current_para_num}.{sub_para_count}"
                    issues_html += f"""
                    <table class="para-table">
                        <tr>
                            <td class="para-num">{num_str}</td>
                            <td class="para-content">
                                <p class="legal-para">The details of the discrepancies are as follows:</p>
                            </td>
                        </tr>
                    </table>
                    {issue['table_html']}
                    """
                
                # End issue
            
            model['issues_content'] = issues_html
            model['issues_templates'] = issues_html
            
            # Letterhead
            import base64
            model['letter_head'] = ""
            if model['show_letterhead']:
                try:
                    from src.utils.config_manager import ConfigManager
                    config = ConfigManager()
                    lh_filename = config.get_pdf_letterhead()
                    lh_path = config.get_letterhead_path('pdf')
                    
                    if lh_path and os.path.exists(lh_path):
                        ext = os.path.splitext(lh_path)[1][1:].lower()
                        
                        # Fetch visual adjustments for this specific letterhead
                        adj = config.get_letterhead_adjustments(lh_filename)
                        width_val = min(adj.get('width', 100), 100)
                        
                        lh_style = f"width: {width_val}%; padding-top: {adj.get('padding_top', 0)}px; margin-bottom: {adj.get('margin_bottom', 20)}px;"
                        
                        if ext == "html":
                            # Case 1: HTML letterhead - read as text
                            with open(lh_path, 'r', encoding='utf-8') as f:
                                lh_full = f.read()
                                # Extract body content if present, otherwise use full content
                                match = re.search(r"<body[^>]*>(.*?)</body>", lh_full, re.DOTALL | re.IGNORECASE)
                                inner_html = match.group(1) if match else lh_full
                                model['letter_head'] = f'<div style="{lh_style} text-align: center;">{inner_html}</div>'
                        else:
                            # Case 2: Image letterhead - read as binary and base64 encode
                            with open(lh_path, "rb") as image_file:
                                encoded_string = base64.b64encode(image_file.read()).decode('utf-8')
                            if ext == "jpg": ext = "jpeg"
                            model['letter_head'] = f'<div style="{lh_style} text-align: center;"><img src="data:image/{ext};base64,{encoded_string}" alt="Letterhead" style="max-width: 100%; height: auto;"></div>'
                except Exception as e:
                    print(f"Letterhead failed: {e}")
                    model['letter_head'] = ""

            # Standard SCN context variables
            model['section'] = model.get('initiating_section', '')
            model['for_pdf'] = for_pdf
            
            # Load CSS explicitly for SCN injection (since TemplateEngine doesn't do this yet)
            template_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "templates")
            css_dir = os.path.join(template_dir, 'css')
            
            def read_css(filename):
                path = os.path.join(css_dir, filename)
                if os.path.exists(path):
                    with open(path, 'r', encoding='utf-8') as f:
                        return f.read()
                return ""

            model['base_css'] = read_css('doc_base.css')
            renderer_mode = 'pdf' if not is_preview else 'qt'
            if renderer_mode == 'pdf':
                model['renderer_css'] = ""
            else:
                model['renderer_css'] = read_css('doc_qt.css')

            model['full_styles_html'] = f"<style>\n{model['base_css']}\n{model['renderer_css']}\n</style>"
            
            # Use Centralized TemplateEngine
            from src.utils.template_engine import TemplateEngine
            return TemplateEngine.render_document("scn.html", model)
            
        except Exception as e:
            print(f"Error rendering SCN: {e}")
            import traceback
            traceback.print_exc()
            return f"<h3>Render Error: {str(e)}</h3>"

    def save_drc01a_metadata(self):
        """Save DRC-01A Metadata (OC No, Dates, etc.) to DB"""
        # 1. Regenerate Model for aggregated periods
        model = self._get_drc01a_model()
        
        metadata = {
            "oc_number": self.oc_number_input.text(),
            "oc_date": self.oc_date_input.date().toString("yyyy-MM-dd"),
            "reply_date": self.reply_deadline_input.date().toString("yyyy-MM-dd"),
            "payment_date": self.payment_deadline_input.date().toString("yyyy-MM-dd"),
            "financial_year": self.proceeding_data.get('financial_year', ''),
            "initiating_section": self.proceeding_data.get('initiating_section', ''),
            "sections_violated_html": self.sections_editor.toHtml(),
            "tax_period_from": model.get('tax_period_from', ''),
            "tax_period_to": model.get('tax_period_to', '')
        }
        
        # [PHASE 15] Safe Merge
        details = copy.deepcopy(self.proceeding_data.get('additional_details', {}))
        details['drc01a_metadata'] = metadata
        
        # Fetch snapshot for DB
        officer_id = None
        officer_snapshot = None
        if hasattr(self, 'officer_combo'):
            officer_id = self.officer_combo.currentData()
            if officer_id:
                officer_data = self.db.get_officer_by_id(officer_id)
                if officer_data:
                    import json
                    existing_snap_str = self.proceeding_data.get('issuing_officer_snapshot')
                    existing_snap = {}
                    if existing_snap_str:
                        try:
                            existing_snap = json.loads(existing_snap_str)
                            if 'name' in existing_snap and 'SCN' not in existing_snap and 'DRC-01A' not in existing_snap:
                                existing_snap = {'DRC-01A': existing_snap}
                        except:
                            existing_snap = {}
                    
                    existing_snap['DRC-01A'] = officer_data
                    officer_snapshot = json.dumps(existing_snap)
        
        self.db.update_proceeding(self.proceeding_id, {
            "initiating_section": metadata['initiating_section'],
            "last_date_to_reply": metadata['reply_date'],
            "additional_details": details,
            "issuing_officer_id": officer_id,
            "issuing_officer_snapshot": officer_snapshot
        })
        
        # Update local data
        self.proceeding_data.update(metadata)
        self.proceeding_data['additional_details'] = details
        self.proceeding_data['issuing_officer_id'] = officer_id
        self.proceeding_data['issuing_officer_snapshot'] = officer_snapshot

    def save_drc01a(self):
        """Save DRC-01A Draft"""
        # 1. Aggregate HTML for Document
        issues_html = ""
        issues_list = []
        
        for card in self.issue_cards:
            issues_html += card.generate_html()
            issues_html += "<br><hr><br>"
            
            # Collect structured data for restoration
            issues_list.append({
                'issue_id': card.template.get('issue_id', 'UNKNOWN'),
                'data': card.get_data()
            })
            
        # [PHASE 5] Single-Source Persistence (DB Only)
        # Issues are now strictly stored in 'case_issues' table.
        # 'additional_details' is reserved for Metadata and legacy migration flags.
        
        # 2. Save Document HTML (for View Mode - Canonical Snapshot)
        doc_data = {
            "proceeding_id": self.proceeding_id,
            "doc_type": "DRC-01A",
            "content_html": issues_html,
            "is_final": 0
        }
        self.db.save_document(doc_data)
        
        # 3. Save Structured Draft Data to case_issues table (Authoritative Data)
        # Verify Not Issued using exact check (though DB Trigger also protects this)
        if self.get_current_stage() < WorkflowStage.DRC01A_ISSUED:
             self.db.save_case_issues(self.proceeding_id, issues_list, stage='DRC-01A')
        
        # 4. Save Metadata (Dates, Section)
        self.save_drc01a_metadata()
        
        self.db.update_proceeding(self.proceeding_id, {
            "status": "DRC-01A Draft"
        })

        QMessageBox.information(self, "Success", "DRC-01A draft saved successfully!")

    def generate_tax_table_html(self, tax_rows):
        """
        PURE FUNCTION: Generate HTML table rows from structured data.
        Accepts list of dicts: [{'act', 'period', 'tax', 'interest', 'penalty', 'total'}, ...]
        """
        rows_html = ""
        if not tax_rows:
            return "<tr><td colspan='7' style='text-align: center;'>No tax details found.</td></tr>"

        from src.utils.formatting import format_indian_number
        
        total_tax = 0
        total_int = 0
        total_pen = 0
        total_grand = 0

        for row in tax_rows:
            act = row.get('act', '')
            period = row.get('period', '')
            tax = row.get('tax', 0)
            interest = row.get('interest', 0)
            penalty = row.get('penalty', 0)
            total = row.get('total', 0)
            
            total_tax += tax
            total_int += interest
            total_pen += penalty
            total_grand += total

            # DRC-01A expects 6 columns: Act | Period | Tax | Interest | Penalty | Total
            rows_html += f"""
            <tr>
                <td style="border: 1px solid black; padding: 5px; text-align: center;">{act}</td>
                <td style="border: 1px solid black; padding: 5px; text-align: center;">{period}</td>
                <td style="border: 1px solid black; padding: 5px; text-align: right;">{format_indian_number(tax)}</td>
                <td style="border: 1px solid black; padding: 5px; text-align: right;">{format_indian_number(interest)}</td>
                <td style="border: 1px solid black; padding: 5px; text-align: right;">{format_indian_number(penalty)}</td>
                <td style="border: 1px solid black; padding: 5px; text-align: right;"><b>{format_indian_number(total)}</b></td>
            </tr>
            """
        
        # Add a Grand Total row for the DRC-01A/SCN Table inject
        rows_html += f"""
        <tr style="background-color: #f2f2f2; font-weight: bold;">
            <td colspan="2" style="border: 1px solid black; padding: 5px; text-align: right;">Total</td>
            <td style="border: 1px solid black; padding: 5px; text-align: right;">{format_indian_number(total_tax)}</td>
            <td style="border: 1px solid black; padding: 5px; text-align: right;">{format_indian_number(total_int)}</td>
            <td style="border: 1px solid black; padding: 5px; text-align: right;">{format_indian_number(total_pen)}</td>
            <td style="border: 1px solid black; padding: 5px; text-align: right;">{format_indian_number(total_grand)}</td>
        </tr>
        """
        return rows_html

    def generate_pdf(self):
        """Generate PDF for current tab"""
        try:
            current_index = self.content_stack.currentIndex()
            
            # Determine content based on tab
            html_content = ""
            filename_prefix = "Document"
            doc_type = "Document"
            
            if current_index == 1: # DRC-01A
                # AUTO-SAVE: Ensure DB is up to date before generation
                self.save_drc01a_metadata()
                
                html_content = self.generate_drc01a_html()
                oc_no = self.oc_number_input.text() or "DRAFT"
                # Sanitize OC No for filename
                safe_oc = "".join([c for c in oc_no if c.isalnum() or c in ('-','_')])
                default_filename = f"DRC-01A_{safe_oc}.pdf"
                doc_type = "DRC 01A"
            elif current_index == 2: # ASMT-10
                # Reuse the ASMT-10 HTML generation logic
                scrutiny_id = self.proceeding_data.get('source_scrutiny_id') or self.proceeding_data.get('scrutiny_id')
                if scrutiny_id:
                    scrutiny_data = self.db.get_scrutiny_case_data(scrutiny_id)
                    if scrutiny_data:
                        case_info = {
                            'case_id': scrutiny_data.get('case_id'),
                            'financial_year': scrutiny_data.get('financial_year'),
                            'section': scrutiny_data.get('section'),
                            'notice_date': scrutiny_data.get('asmt10_finalised_on', '-'),
                            'oc_number': scrutiny_data.get('oc_number', 'DRAFT')
                        }
                        taxpayer = scrutiny_data.get('taxpayer_details', {})
                        issues = scrutiny_data.get('selected_issues', [])
                        generator = ASMT10Generator()
                        full_data = case_info.copy()
                        full_data['taxpayer_details'] = taxpayer
                        html_content = generator.generate_html(full_data, issues, for_pdf=True)
                        filename_prefix = f"ASMT-10_{case_id}"
                        default_filename = f"{filename_prefix}.pdf"
                        doc_type = "ASMT-10"
                    else:
                        QMessageBox.warning(self, "Error", "Source scrutiny data could not be retrieved.")
                        return
                else:
                    QMessageBox.warning(self, "Error", "No source ASMT-10 linked to this case.")
                    return
            elif current_index == 3: # SCN
                if self.is_scn_phase1():
                    QMessageBox.warning(self, "Locked", "SCN PDF generation is locked during Phase-1.")
                    return
                html_content = self.render_scn(for_pdf=True)
                case_id = self.proceeding_data.get('case_issue_id', 'DRAFT').replace('/', '_')
                default_filename = f"SCN_{case_id}.pdf"
                doc_type = "Show Cause Notice"
            else:
                QMessageBox.warning(self, "Error", f"PDF generation not supported for tab index {current_index} yet.")
                return
            
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "Save PDF As",
                default_filename,
                "PDF Files (*.pdf)"
            )
            
            if file_path:
                # Generate PDF using the correct method
                import shutil
                import os
                from src.utils.document_generator import DocumentGenerator
                
                # [ULTRA-SAFE] UI Mitigation
                from PyQt6.QtWidgets import QApplication
                QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)
                
                filename_only = os.path.splitext(os.path.basename(file_path))[0]
                doc_gen = DocumentGenerator()
                try:
                    generated_path = doc_gen.generate_pdf_from_html(html_content, filename_only)
                    
                    if generated_path and os.path.exists(generated_path):
                        # Move the file to user-selected location
                        shutil.move(generated_path, file_path)
                        
                        # 3. Save Document Record to DB
                        doc_data = {
                            'proceeding_id': self.proceeding_data.get('id'),
                            'doc_type': doc_type,
                            'content_html': html_content,
                            'template_id': None,
                            'template_version': 1,
                            'version_no': 1,
                            'is_final': 1,
                            'snapshot_path': file_path
                        }
                        self.db.save_document(doc_data)
                        
                        QMessageBox.information(self, "Success", f"PDF generated and Registered successfully!\n\nSaved to: {file_path}")
                    else:
                        raise RuntimeError("File generation failed without error message.")
                        
                except Exception as e:
                    err_str = str(e)
                    if "MISSING_DEPENDENCY" in err_str:
                        msg = ("PDF Generation Unavailable.\n\n"
                               "The system is missing required rendering libraries (GTK3).\n"
                               "Please install the dependencies or use 'Draft Word' instead.")
                        QMessageBox.critical(self, "Dependency Error", msg)
                    elif "TIMEOUT" in err_str:
                        QMessageBox.warning(self, "Timeout", "PDF generation took too long (20s limit) and was cancelled for safety.")
                    else:
                        QMessageBox.critical(self, "Error", f"Failed to generate {doc_type} PDF: {err_str}")
                finally:
                    QApplication.restoreOverrideCursor()
                
                # Open the file if it exists and was successfully moved
                if os.path.exists(file_path):
                    try:
                        os.startfile(file_path)
                    except Exception as e:
                        print(f"Could not open file: {e}")
                        
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error during export: {str(e)}")
            import traceback
            traceback.print_exc()

    def generate_docx(self):
        """Generate DOCX document for DRC-01A"""
        try:
            # Check Tab
            current_index = self.content_stack.currentIndex()
            if current_index == 1: # DRC-01A
                pass
            elif current_index == 3: # SCN
                if self.is_scn_phase1():
                    QMessageBox.warning(self, "Locked", "SCN DOCX generation is locked during Phase-1.")
                    return
                # SCN Logic
                model = self._get_scn_model()
                if not model:
                     QMessageBox.warning(self, "Error", "Failed to build SCN model.")
                     return
                
                from script_docx import Document # Using a helper if available, or docx.Document
                from docx import Document
                from docx.shared import Pt, Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                case_id = model.get('case_id', 'DRAFT').replace('/', '_')
                default_filename = f"SCN_{case_id}.docx"
                
                from PyQt6.QtWidgets import QFileDialog
                file_path, _ = QFileDialog.getSaveFileName(self, "Save SCN DOCX As", default_filename, "Word Documents (*.docx)")
                if not file_path: return

                doc = Document()
                for section in doc.sections:
                    section.top_margin = Inches(0.8)
                    section.bottom_margin = Inches(0.8)
                    section.left_margin = Inches(1)
                    section.right_margin = Inches(1)

                # Title
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run("SHOW CAUSE NOTICE")
                run.bold = True
                run.font.size = Pt(16)
                
                # Header
                doc.add_paragraph(f"GSTIN: {model.get('gstin', '')}")
                doc.add_paragraph(f"Legal Name: {model.get('legal_name', '')}")
                doc.add_paragraph(f"OC Number: {model.get('oc_no', '')}")
                doc.add_paragraph(f"Date: {model.get('issue_date', '')}")
                doc.add_paragraph()

                # Paragraph 1: Intro
                p = doc.add_paragraph()
                p.add_run("1. ").bold = True
                p.add_run("A brief of the case and the grounds for initiating proceedings are as follows:")
                
                # Issues
                for issue in model['issues']:
                    p_num = issue['index'] + 2 # Starts at 3
                    title = issue['title']
                    
                    p = doc.add_paragraph()
                    p.add_run(f"{p_num}. Issue No. {issue['index']}: {title}").bold = True
                    
                    for s_idx, para in enumerate(issue['paras'], start=1):
                        sp = doc.add_paragraph()
                        sp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        sp.add_run(f"{p_num}.{s_idx} ").bold = True
                        sp.add_run(para)
                
                # Demands
                num = model['para_demand']
                p = doc.add_paragraph()
                p.add_run(f"{num}. Summary of Tax Liability:").bold = True
                doc.add_paragraph(f"The total tax liability is determined as ₹{model['total_amount']}.")

                # Proper Officer
                doc.add_paragraph()
                sig = doc.add_paragraph()
                sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                sig.add_run("Proper Officer").bold = True
                
                doc.save(file_path)
                import os
                os.startfile(file_path)
                return

            # 1. Regenerate model & Validate (DRC-01A Only)
            if current_index == 1:
                drc_model = self._get_drc01a_model()
                valid, msg = self.validate_drc01a_model(drc_model)
                if not valid:
                    QMessageBox.warning(self, "Export Validation", msg)
                    return
            else:
                drc_model = {}

            from PyQt6.QtWidgets import QFileDialog
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import os
            
            # Ask user for save location
            case_id = self.proceeding_data.get('case_id', 'DRAFT').replace('/', '_')
            default_filename = f"DRC-01A_{case_id}.docx"
            
            file_path, _ = QFileDialog.getSaveFileName(self, "Save DOCX As", default_filename, "Word Documents (*.docx)")
            if not file_path: return

            # Create DOCX document
            doc = Document()
            for section in doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Title
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title.add_run("FORM DRC-01A")
            run.bold = True
            run.font.size = Pt(16)
            
            # Case details
            doc.add_paragraph(f"Case ID: {drc_model.get('case_id', 'N/A')}")
            doc.add_paragraph(f"GSTIN: {drc_model.get('gstin', 'N/A')}")
            doc.add_paragraph(f"Legal Name: {drc_model.get('legal_name', 'N/A')}")
            doc.add_paragraph(f"Address: {drc_model.get('address', 'N/A')}")
            doc.add_paragraph()
            
            # Tax Demand Table (Data-Driven)
            tax_rows = drc_model.get('tax_rows', [])
            if tax_rows:
                heading = doc.add_paragraph()
                run = heading.add_run("Tax Demand Details:")
                run.bold = True
                run.font.size = Pt(14)
                
                # Create table
                table = doc.add_table(rows=len(tax_rows) + 1, cols=7)
                table.style = 'Light Grid Accent 1'
                
                # Headers
                headers = ["Act", "Tax Period From", "Tax Period To", "Tax (₹)", "Interest (₹)", "Penalty (₹)", "Total (₹)"]
                for col, header in enumerate(headers):
                    cell = table.rows[0].cells[col]
                    cell.text = header
                    cell.paragraphs[0].runs[0].bold = True
                
                # Data
                for r_idx, row in enumerate(tax_rows, start=1):
                    table.rows[r_idx].cells[0].text = row.get('Act', '')
                    table.rows[r_idx].cells[1].text = row.get('Period', '')
                    table.rows[r_idx].cells[2].text = row.get('Period', '')
                    table.rows[r_idx].cells[3].text = row.get('Tax', '0')
                    table.rows[r_idx].cells[4].text = row.get('Interest', '0')
                    table.rows[r_idx].cells[5].text = row.get('Penalty', '0')
                    table.rows[r_idx].cells[6].text = row.get('Total', '0')
                
                doc.add_paragraph()
            
            # Issues & Sections (Simplified for now)
            heading = doc.add_paragraph()
            run = heading.add_run("Details of Discrepancies:")
            run.bold = True
            run.font.size = Pt(14)
            doc.add_paragraph("Please refer to the SCN preview/PDF for full issue narratives.")
            doc.add_paragraph()
            
            # Advice Text
            advice_text = drc_model.get('AdviceText', "Please pay the amount as ascertained.")
            p = doc.add_paragraph()
            p.add_run(advice_text)
            
            # Signatures
            doc.add_paragraph()
            sig = doc.add_paragraph()
            sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = sig.add_run("Proper Officer\n(Signature)\nName: ____________________\nDesignation: ____________________\nJurisdiction: ____________________")
            run.bold = True
            
            doc.save(file_path)
            QMessageBox.information(self, "Success", f"DOCX generated successfully!\n\nSaved to: {file_path}")
            os.startfile(file_path)
                    
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error generating DOCX: {str(e)}")

    def check_existing_documents(self):
        """Check workflow stage and toggle view mode"""
        # [PHASE 2] Source of Truth is Workflow Stage, not document existence
        current_stage = self.get_current_stage()
        
        drc01a_done = current_stage >= WorkflowStage.DRC01A_ISSUED
        scn_done = current_stage >= WorkflowStage.SCN_ISSUED
        
        self.toggle_view_mode("drc01a", drc01a_done)
        self.toggle_view_mode("scn", scn_done)

    def toggle_view_mode(self, doc_type, show_view):
        """Toggle between Draft and View containers"""
        if doc_type == "drc01a":
            if show_view:
                self.drc01a_draft_container.hide()
                self.drc01a_view_container.show()
            else:
                self.drc01a_view_container.hide()
                self.drc01a_draft_container.show()
                
        elif doc_type == "scn":
            if show_view:
                self.scn_draft_container.hide()
                self.scn_view_container.show()
            else:
                self.scn_view_container.hide()
                self.scn_draft_container.show()

    def _on_auto_generate_oc(self):
        """Helper for SCN Auto-Generate utility"""
        self._auto_generating_oc = True
        self.suggest_next_oc(self.scn_oc_input)
        if self.scn_oc_input.text():
            self.oc_provenance_lbl.setText("✔ Generated from O.C. Register")
            self.oc_provenance_lbl.show()
            # Tooltip safeguard
            self.scn_oc_input.setToolTip("Auto-generated value. You can still edit manually.")
        self._auto_generating_oc = False

    def _on_ph_auto_generate_oc(self):
        """Helper for PH Auto-Generate OC utility"""
        self.suggest_next_oc(self.ph_edit_oc)
        
    def _on_scn_oc_changed(self):
        """Handle OC text changes for SCN"""
        if not getattr(self, '_auto_generating_oc', False):
            # Manually edited
            if hasattr(self, 'oc_provenance_lbl'):
                self.oc_provenance_lbl.hide()
        
        
        self.evaluate_scn_workflow_phase()

    def toggle_view_mode(self, doc_type, show_view):
        """Toggle between Draft and View containers with [PHASE 5] Hard Lock Enforcement"""
        if doc_type == "drc01a":
            if not show_view:
                # REQUESTING EDIT MODE
                # [HARD LOCK] Block if already issued
                if self.get_current_stage() >= WorkflowStage.DRC01A_ISSUED:
                    from PyQt6.QtWidgets import QMessageBox
                    QMessageBox.warning(self, "Locked", "DRC-01A has been issued and cannot be edited.")
                    return # Silent Reject
                
                self.drc01a_view_container.hide()
                self.drc01a_draft_container.show()
            else:
                # REQUESTING VIEW MODE
                self.drc01a_draft_container.hide()
                self.drc01a_view_container.show()
                
                # [PHASE 5] Snapshot Freeze Logic
                if self.get_current_stage() >= WorkflowStage.DRC01A_ISSUED:
                    try:
                        oc_no = self.oc_number_input.text().strip()
                        safe_oc = "".join([c for c in oc_no if c.isalnum() or c in ('-','_')])
                        filename = f"DRC-01A_{safe_oc}.html"
                        
                        base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
                        file_path = os.path.join(base_dir, "data", "generated_documents", filename)
                        
                        if os.path.exists(file_path):
                            from PyQt6.QtCore import QUrl
                            self.drc01a_browser.setUrl(QUrl.fromLocalFile(file_path))
                        else:
                            self.drc01a_browser.setHtml(f"<h3 style='color:red; text-align:center'>Snapshot File Not Found: {filename}</h3>")
                    except Exception as e:
                        self.drc01a_browser.setHtml(f"<h3 style='color:red; text-align:center'>Error loading snapshot: {str(e)}</h3>")
                
        elif doc_type == "scn":
             # SCN Logic (Placeholder for future hardening)
             if show_view:
                 self.scn_draft_container.hide()
                 self.scn_view_container.show()
             else:
                 self.scn_view_container.hide()
                 self.scn_draft_container.show()

    def confirm_ph_finalization(self):

        """Finalize PH and Register OC"""
        if not self.ph_oc_input.text().strip():
            QMessageBox.warning(self, "Validation Error", "OC Number is mandatory for registration.")
            return
            
        try:
            oc_data = {
                'OC_Number': self.ph_oc_input.text(),
                'OC_Date': self.ph_oc_date.date().toString("yyyy-MM-dd"),
                'OC_Content': f"Personal Hearing Intimation issued for GSTIN {self.proceeding_data.get('gstin','')}.",
                'OC_To': self.proceeding_data.get('legal_name', '')
            }
            self.db.add_oc_entry(self.proceeding_id, oc_data)
            
            # Update status
            self.db.update_proceeding(self.proceeding_id, {
                "status": "PH Intimated",
                "workflow_stage": WorkflowStage.PH_SCHEDULED
            })
            
            QMessageBox.information(self, "Success", "PH Intimation Finalized and OC Registered.")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to finalize PH: {e}")

    def add_new_ph_entry(self):
        """Add a new PH entry slot (Max 3)"""
        if self.is_scn_finalized():
            QMessageBox.warning(self, "Locked", "Case finalized. Cannot add new PH Intimations.")
            return
            
        if len(self.ph_entries) >= 3:
            QMessageBox.warning(self, "Limit Reached", "Maximum 3 Personal Hearing Intimations allowed per case.")
            return

        # Prepare default data
        next_val = "DRAFT"
        self.ph_edit_oc.setText(next_val)
        self.ph_edit_oc_date.setDate(QDate.currentDate())
        self.ph_edit_date.setDate(QDate.currentDate().addDays(7))
        self.ph_edit_time.setText("11:00 AM")
        self.ph_edit_venue.setText("Paravur Range Office")
        self.ph_edit_copy_to.setText("The Assistant Commissioner, Central Tax, Paravur Division")
        
        # Track that we are adding a NEW entry (index = -1)
        self.ph_editing_index = -1
        self.ph_editor_card.setVisible(True)

    def save_ph_entry(self):
        """Save form data into ph_entries list and persist"""
        entry = {
            "oc_no": self.ph_edit_oc.text(),
            "issue_date": self.ph_edit_oc_date.date().toString("dd/MM/yyyy"),
            "ph_date": self.ph_edit_date.date().toString("dd/MM/yyyy"),
            "ph_time": self.ph_edit_time.text(),
            "venue": self.ph_edit_venue.text(),
            "copy_to": self.ph_edit_copy_to.text(),
            "show_letterhead": self.ph_show_lh.isChecked()
        }
        
        print(f"DEBUG: save_ph_entry captured: {entry}")
        
        if self.ph_editing_index == -1:
            self.ph_entries.append(entry)
        else:
            self.ph_entries[self.ph_editing_index] = entry
            
        self.ph_editor_card.setVisible(False)
        self.refresh_ph_list()
        self.render_ph_preview()
        
        # Persist to additional_details
        self.save_ph_data()

    def register_ph_entry(self):
        """Finalize, Register OC, and Save entry"""
        if self.is_scn_finalized(): return
        
        oc_no = self.ph_edit_oc.text().strip()
        if not oc_no or oc_no == "DRAFT":
            QMessageBox.warning(self, "Validation Error", "Please enter a valid OC Number for registration.")
            return
            
        try:
            # 1. Register in OC Register
            oc_data = {
                'OC_Number': oc_no,
                'OC_Date': self.ph_edit_oc_date.date().toString("yyyy-MM-dd"),
                'OC_Content': f"Personal Hearing Intimation for Case {self.proceeding_id}",
                'OC_To': self.proceeding_data.get('legal_name', '')
            }
            self.db.add_oc_entry(self.proceeding_id, oc_data)
            
            # 2. Mark entry as registered
            entry = {
                "oc_no": oc_no,
                "issue_date": self.ph_edit_oc_date.date().toString("dd/MM/yyyy"),
                "ph_date": self.ph_edit_date.date().toString("dd/MM/yyyy"),
                "ph_time": self.ph_edit_time.text(),
                "venue": self.ph_edit_venue.text(),
                "copy_to": self.ph_edit_copy_to.text(),
                "show_letterhead": self.ph_show_lh.isChecked(),
                "is_registered": True
            }
            
            if self.ph_editing_index == -1:
                self.ph_entries.append(entry)
            else:
                self.ph_entries[self.ph_editing_index] = entry
                
            # 3. Update status (cumulative)
            self.db.update_proceeding(self.proceeding_id, {"status": "PH Intimated"})
            
            self.ph_editor_card.setVisible(False)
            self.refresh_ph_list()
            self.render_ph_preview()
            self.save_ph_data()
            
            QMessageBox.information(self, "Success", f"PH Intimation registered successfully with OC {oc_no}.")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to register PH: {e}")

    def refresh_ph_list(self):
        """Rebuild the list of PH entry cards with modern neutralized style"""
        while self.ph_list_layout.count():
            item = self.ph_list_layout.takeAt(0)
            if item.widget(): item.widget().deleteLater()
            
        is_scn_fin = self.is_scn_finalized()
        
        for idx, entry in enumerate(self.ph_entries):
            card = QFrame()
            is_reg = entry.get('is_registered', False)
            card.setStyleSheet(f"""
                QFrame {{ 
                    background-color: {Theme.SURFACE}; 
                    border-radius: 8px; 
                    border: 1px solid {Theme.NEUTRAL_200};
                }}
            """)
            
            l = QHBoxLayout(card)
            l.setContentsMargins(12, 12, 12, 12)
            
            # Info Section
            info_layout = QVBoxLayout()
            info_layout.setSpacing(4)
            
            title = QLabel(f"PH {idx+1}: {entry['ph_date']} @ {entry['ph_time']}")
            title.setStyleSheet(f"font-weight: bold; color: {Theme.NEUTRAL_900}; font-size: 14px; border: none;")
            
            status_row = QHBoxLayout()
            status_row.setSpacing(8)
            
            badge_text = "REGISTERED" if is_reg else "DRAFT"
            badge_bg = Theme.SUCCESS if is_reg else Theme.NEUTRAL_100
            badge_color = "white" if is_reg else Theme.NEUTRAL_500
            
            badge = QLabel(badge_text)
            badge.setStyleSheet(f"""
                padding: 2px 8px; border-radius: 4px; font-size: 10px; font-weight: bold;
                background-color: {badge_bg}; color: {badge_color}; border: none;
            """)
            
            oc_info = QLabel(f"OC: {entry['oc_no']}")
            oc_info.setStyleSheet(f"color: {Theme.NEUTRAL_500}; font-size: 12px; border: none;")
            
            status_row.addWidget(badge)
            status_row.addWidget(oc_info)
            status_row.addStretch()
            
            info_layout.addWidget(title)
            info_layout.addLayout(status_row)
            l.addLayout(info_layout)
            l.addStretch()
            
            # Actions
            edit_btn = QPushButton("View" if (is_scn_fin or is_reg) else "Edit")
            edit_btn.setFixedHeight(32)
            edit_btn.setFixedWidth(80)
            edit_btn.setStyleSheet(f"""
                QPushButton {{ 
                    background: transparent; border: 1px solid {Theme.NEUTRAL_200}; 
                    color: {Theme.NEUTRAL_900}; border-radius: 6px; font-weight: bold;
                }}
                QPushButton:hover {{ background-color: {Theme.NEUTRAL_100}; }}
            """)
            edit_btn.clicked.connect(lambda _, i=idx: self.edit_ph_entry(i))
            l.addWidget(edit_btn)
            
            del_btn = QPushButton("Delete")
            del_btn.setFixedHeight(32)
            del_btn.setFixedWidth(80)
            del_btn.setStyleSheet(f"""
                QPushButton {{ 
                    background: transparent; border: none; 
                    color: {Theme.DANGER}; font-weight: bold;
                }}
                QPushButton:hover {{ color: {Theme.DANGER_HOVER}; }}
            """)
            del_btn.clicked.connect(lambda _, i=idx: self.delete_ph_entry(i))
            if is_scn_fin or is_reg: del_btn.setEnabled(False)
            l.addWidget(del_btn)
            
            self.ph_list_layout.addWidget(card)
        
        self.ph_add_btn.setEnabled(not (len(self.ph_entries) >= 3 or is_scn_fin))

    def edit_ph_entry(self, index):
        """Load entry into form for editing"""
        try:
            entry = self.ph_entries[index]
            self.ph_editing_index = index
            is_reg = entry.get('is_registered', False)
            is_scn_fin = self.is_scn_finalized()
            
            print(f"DEBUG: Editing PH Entry {index}: {entry}")

            # Safe Get with Defaults
            self.ph_edit_oc.setText(entry.get('oc_no', ''))
            
            # Handle Dates (Robust parsing)
            i_date = entry.get('issue_date', '')
            p_date = entry.get('ph_date', '')
            
            if i_date: self.ph_edit_oc_date.setDate(QDate.fromString(i_date, "dd/MM/yyyy"))
            else: self.ph_edit_oc_date.setDate(QDate.currentDate())
                
            if p_date: self.ph_edit_date.setDate(QDate.fromString(p_date, "dd/MM/yyyy"))
            else: self.ph_edit_date.setDate(QDate.currentDate())
                
            self.ph_edit_time.setText(entry.get('ph_time', ''))
            self.ph_edit_venue.setText(entry.get('venue', ''))
            self.ph_edit_copy_to.setText(entry.get('copy_to', ''))
            
            # Restore Checkbox (Missing in previous version)
            self.ph_show_lh.setChecked(entry.get('show_letterhead', False))
            
            # Lock fields if registered or SCN finalized
            locked = is_reg or is_scn_fin
            self.ph_edit_oc.setReadOnly(locked)
            self.ph_edit_oc_date.setReadOnly(locked)
            self.ph_edit_date.setReadOnly(locked)
            self.ph_edit_time.setReadOnly(locked)
            self.ph_edit_venue.setReadOnly(locked)
            self.ph_edit_copy_to.setReadOnly(locked)
            self.ph_show_lh.setEnabled(not locked)
            
            # Find buttons in editor card layout and toggle
            for i in range(self.ph_editor_card.layout().count()):
                item = self.ph_editor_card.layout().itemAt(i)
                # Both grid layout and horizontal button layout need button locking
                if isinstance(item.layout(), QHBoxLayout):
                    for j in range(item.layout().count()):
                        w = item.layout().itemAt(j).widget()
                        if isinstance(w, QPushButton) and w.text() != "Cancel":
                            w.setEnabled(not locked)
            
            self.ph_editor_card.setVisible(True)
            self.render_ph_preview()
        except Exception as e:
            print(f"Error loading PH entry for edit: {e}")
            traceback.print_exc()
            QMessageBox.warning(self, "Error", "Failed to load entry details.")

    def delete_ph_entry(self, index):
        """Remove a PH entry"""
        if self.is_scn_finalized(): return
        
        res = QMessageBox.question(self, "Confirm Delete", "Delete this Personal Hearing Intimation permanentely?", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if res == QMessageBox.StandardButton.Yes:
            self.ph_entries.pop(index)
            self.refresh_ph_list()
            self.save_ph_data()
            if self.ph_browser: self.ph_browser.setHtml("<h3>Select or Add a PH Entry to Preview</h3>")

    def render_ph_preview(self):
        """Render the highlighted/active PH entry into WebEngine with high-fidelity A4 simulation"""
        if not self.ph_entries:
            self.ph_browser.setHtml("<div style='background: #f1f3f4; height: 100vh; display: flex; align-items: center; justify-content: center; font-family: sans-serif; color: #7f8c8d;'><h3>Select or Add a PH Entry to Preview</h3></div>")
            return
            
        # Preview the current entry being edited OR the last one
        idx = self.ph_editing_index if self.ph_editing_index != -1 else (len(self.ph_entries) - 1)
        entry = self.ph_entries[idx].copy()
        
        # Reactive State: Override entry data with current form values for live preview
        entry['show_letterhead'] = self.ph_show_lh.isChecked()
        
        if self.ph_editor_card.isVisible():
            entry.update({
                "oc_no": self.ph_edit_oc.text(),
                "issue_date": self.ph_edit_oc_date.date().toString("dd/MM/yyyy"),
                "ph_date": self.ph_edit_date.date().toString("dd/MM/yyyy"),
                "ph_time": self.ph_edit_time.text(),
                "venue": self.ph_edit_venue.text(),
                "copy_to": self.ph_edit_copy_to.text()
            })
            
        # Gather case data
        case_data = {
            'gstin': self.proceeding_data.get('gstin', ''),
            'legal_name': self.proceeding_data.get('legal_name', ''),
            'address': self.proceeding_data.get('address', ''),
            'scn_no': self.scn_no_input.text() or "____",
            'scn_date': self.scn_date_input.date().toString("dd/MM/yyyy")
        }
        
        html = self.ph_generator.generate_html(case_data, entry, for_preview=True)
        self.ph_browser.setHtml(html)

    def save_ph_data(self):
        """Persist ph_entries to additional_details using safe merge"""
        try:
            details = copy.deepcopy(self.proceeding_data.get('additional_details', {}))
            details['ph_entries'] = self.ph_entries
            
            # [DEBUG]
            print(f"DEBUG: save_ph_data [Pre-DB] Entries Count: {len(self.ph_entries)}")
            print(f"DEBUG: save_ph_data [Pre-DB] Dumping details... (Keys: {list(details.keys())})")
            
            success = self.db.update_proceeding(self.proceeding_id, {
                "additional_details": details
            })
            
            if success:
                print("DEBUG: save_ph_data [DB Update Success]")
                self.proceeding_data['additional_details'] = details
            else:
                print("DEBUG: save_ph_data [DB Update FAILED]")
                
        except Exception as e:
            print(f"DEBUG: save_ph_data CRITICAL FAILURE: {e}")
            traceback.print_exc()

    def generate_ph_pdf(self):
        """Finalize and Generate PDF for active PH entry"""
        if not self.ph_entries: return
        
        idx = self.ph_editing_index if self.ph_editing_index != -1 else (len(self.ph_entries) - 1)
        entry = self.ph_entries[idx]
        
        case_data = {
            'gstin': self.proceeding_data.get('gstin', ''),
            'legal_name': self.proceeding_data.get('legal_name', ''),
            'address': self.proceeding_data.get('address', ''),
            'scn_no': self.scn_no_input.text() or "____",
            'scn_date': self.scn_date_input.date().toString("dd/MM/yyyy")
        }
        
        html = self.ph_generator.generate_html(case_data, entry, for_preview=False, for_pdf=True)
        
        safe_oc = "".join([c for c in entry['oc_no'] if c.isalnum() or c in ('-','_')])
        default_filename = f"PH_Intimation_{safe_oc}.pdf"
        
        file_path, _ = QFileDialog.getSaveFileName(self, "Save PH Intimation PDF", default_filename, "PDF Files (*.pdf)")
        if file_path:
            from src.utils.document_generator import DocumentGenerator
            import shutil
            import os
            from PyQt6.QtWidgets import QApplication
            
            filename_only = os.path.splitext(os.path.basename(file_path))[0]
            doc_gen = DocumentGenerator()
            
            # [ULTRA-SAFE] UI Mitigation
            QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)
            # Disable triggering button if we had a direct reference, 
            # but since this is a general handler we rely on the modal wait cursor.
            
            try:
                generated_path = doc_gen.generate_pdf_from_html(html, filename_only)
                
                if generated_path and os.path.exists(generated_path):
                    shutil.move(generated_path, file_path)
                    QMessageBox.information(self, "Success", f"PH Intimation PDF saved to: {file_path}")
                    os.startfile(file_path)
                else:
                    raise RuntimeError("File generation failed without error message.")
                    
            except Exception as e:
                err_str = str(e)
                if "MISSING_DEPENDENCY" in err_str:
                    msg = ("PDF Generation Unavailable.\n\n"
                           "The system is missing required rendering libraries (GTK3).\n"
                           "Please install the dependencies or use 'Draft DOCX' instead.")
                    QMessageBox.critical(self, "Dependency Error", msg)
                elif "TIMEOUT" in err_str:
                    QMessageBox.warning(self, "Timeout", "PDF generation took too long (20s limit) and was cancelled for safety.")
                else:
                    QMessageBox.critical(self, "Error", f"Failed to generate PDF: {err_str}")
            finally:
                QApplication.restoreOverrideCursor()

    def generate_ph_docx(self):
        """Generate DOCX Document for the active PH Intimation entry"""
        if not self.ph_entries: return
        
        idx = self.ph_editing_index if self.ph_editing_index != -1 else (len(self.ph_entries) - 1)
        entry = self.ph_entries[idx]
        
        # Build filename
        safe_oc = "".join([c for c in entry['oc_no'] if c.isalnum() or c in ('-','_')])
        default_filename = f"PH_Intimation_{safe_oc}.docx"
        
        file_path, _ = QFileDialog.getSaveFileName(self, "Save PH Intimation DOCX", default_filename, "Word Documents (*.docx)")
        if not file_path: return
        
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Margins
            for section in doc.sections:
                section.top_margin = Inches(0.8)
                section.bottom_margin = Inches(0.8)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
                
            # Header Info
            p = doc.add_paragraph()
            p.add_run(f"O.C. No. {entry['oc_no']}").bold = True
            tab_run = p.add_run("\t\t\t\t\t") # Basic spacing simulation
            p.add_run(f"Date: {entry['issue_date']}").bold = True
            
            doc.add_paragraph()
            
            # To Address
            doc.add_paragraph("To,")
            p = doc.add_paragraph()
            p.add_run(self.proceeding_data.get('legal_name', '')).bold = True
            doc.add_paragraph(f"GSTIN: {self.proceeding_data.get('gstin', '')}")
            doc.add_paragraph(self.proceeding_data.get('address', ''))
            
            doc.add_paragraph()
            doc.add_paragraph("Gentlemen/Sir/Madam,")
            doc.add_paragraph()
            
            # Subject
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run("Subject: Intimation of Personal Hearing – reg")
            run.bold = True
            
            doc.add_paragraph()
            
            # Reference
            p = doc.add_paragraph()
            p.add_run("References: ").bold = True
            p.add_run(f"1. SCN reference number: {self.scn_no_input.text()} dated {self.scn_date_input.date().toString('dd/MM/yyyy')}")
            
            doc.add_paragraph()
            
            # Paragraphs
            p1 = doc.add_paragraph("1. Please refer to the above mentioned SCN number issued by Office of the Superintendent Paravur Range.")
            p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            p2 = doc.add_paragraph(f"2. In this connection, it is to inform you that personal hearing in this case will be held at {entry['ph_time']} on {entry.get('ph_date', '')} before the Superintendent of Central Tax, Paravur Range Office.")
            p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            p3 = doc.add_paragraph("3. You may therefore appear in person or through an authorized representative for the personal hearing on the above mentioned date and time as per your convenience, at the above mentioned address, without fail along with records/documents/evidences you wish to rely upon in support of your case.")
            p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            doc.add_paragraph()
            doc.add_paragraph()
            
            # Signature
            sig = doc.add_paragraph()
            sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sig.add_run("VISHNU V\nSuperintendent").bold = True
            
            doc.add_paragraph()
            
            # Copy To
            doc.add_paragraph("Copy submitted to:").bold = True
            doc.add_paragraph(f"1. {entry['copy_to']}")
            
            doc.save(file_path)
            QMessageBox.information(self, "Success", f"Word document saved to: {file_path}")
            os.startfile(file_path)
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to generate Word document: {e}")
            traceback.print_exc()

    def confirm_order_finalization(self):
        """Finalize Order and Register OC"""
        if not self.order_oc_input.text().strip():
            QMessageBox.warning(self, "Validation Error", "OC Number is mandatory for registration.")
            return
            
        try:
            oc_data = {
                'OC_Number': self.order_oc_input.text(),
                'OC_Date': self.order_oc_date.date().toString("yyyy-MM-dd"),
                'OC_Content': f"Final Order (DRC-07) issued for GSTIN {self.proceeding_data.get('gstin','')}.",
                'OC_To': self.proceeding_data.get('legal_name', '')
            }
            self.db.add_oc_entry(self.proceeding_id, oc_data)
            
            # 2. Transition State (Atomic Update)
            self.transition_to(WorkflowStage.ORDER_ISSUED)
            
            QMessageBox.information(self, "Success", "Order Finalized and OC Registered.")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to finalize Order: {e}")

    def suggest_next_oc(self, input_field: QLineEdit):
        """Fetch next available OC number and set it to input"""
        try:
            import datetime
            current_year = datetime.date.today().year
            
            # Fetch next number from DB
            # We assume format XXX/YEAR
            next_num = self.db.get_next_oc_number(str(current_year))
            
            formatted_oc = f"{next_num}/{current_year}"
            input_field.setText(formatted_oc)
            
        except Exception as e:
            print(f"Error suggesting OC: {e}")

    # [PHASE 3] Restored Finalization Methods & New PH Logic

    def confirm_drc01a_finalization(self):
        """Commit DRC-01A Finalization"""
        try:
            # 1. Save Document as Final
            self.save_drc01a_data() # Ensure latest draft is saved
            
            # 2. Transition State (Atomic Update)
            self.transition_to(WorkflowStage.DRC01A_ISSUED)
            
            # 3. Add OC Entry for DRC-01A
            oc_data = {
                'OC_Number': self.oc_number_input.text(),
                'OC_Date': self.oc_date_input.date().toString("yyyy-MM-dd"),
                'OC_Content': f"DRC-01A Issued. Case ID: {self.proceeding_data.get('case_id')}",
                'OC_To': self.proceeding_data.get('legal_name', '')
            }
            self.db.add_oc_entry(self.proceeding_id, oc_data)
            
            QMessageBox.information(self, "Success", "DRC-01A Finalized Successfully.")
            
            # 4. Switch to View Mode
            self.drc01a_finalization_container.hide()
            self.drc01a_view_container.show()
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Finalization failed: {e}")

    def confirm_scn_finalization(self):
        """Commit SCN Finalization"""
        try:
            # 1. Save Document as Final
            self.save_document("SCN", is_manual=True) # Ensure latest draft is saved
            
            # 2. Transition State (Atomic Update)
            self.transition_to(WorkflowStage.SCN_ISSUED)
            
            # 3. Add OC Entry for SCN
            oc_data = {
                'OC_Number': self.scn_oc_input.text(),
                'OC_Date': self.scn_date_input.date().toString("yyyy-MM-dd"),
                'OC_Content': f"Show Cause Notice Issued. SCN No: {self.scn_no_input.text()}. {self.scn_fin_panel.fin_scn_remarks.toPlainText()}",
                'OC_To': self.proceeding_data.get('legal_name', '')
            }
            self.db.add_oc_entry(self.proceeding_id, oc_data)
            
            QMessageBox.information(self, "Success", "Show Cause Notice Finalized Successfully.")
            
            # 4. Switch to View Mode
            self.scn_finalization_container.hide()
            self.scn_view_container.show()
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Finalization failed: {e}")
            import traceback
            traceback.print_exc()

    def confirm_ph_finalization(self):
        """Commit PH Intimation Finalization"""
        try:
            # 1. Transition State (Atomic Update) (PH Scheduled)
            self.transition_to(WorkflowStage.PH_SCHEDULED)
            
            # 2. Add OC Entry for PH
            oc_data = {
                'OC_Number': self.ph_edit_oc.text(),
                'OC_Date': self.ph_date.date().toString("yyyy-MM-dd"),
                'OC_Content': f"PH Intimation Issued. Hearing on {self.ph_date.date().toString('dd/MM/yyyy')} at {self.ph_time.time().toString('hh:mm A')}",
                'OC_To': self.proceeding_data.get('legal_name', '')
            }
            self.db.add_oc_entry(self.proceeding_id, oc_data)
            
            QMessageBox.information(self, "Success", "Personal Hearing Intimated Successfully.")
            
            self.contact_card.hide()
            self.refresh_ph_list()
            
            # Show 'Conclude Hearing' button logic would be refreshed by UI update
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Finalization failed: {e}")

    def conclude_ph_hearing(self):
        """[Phase 3] Conclude PH Hearing Lifecycle Step -> Allows Order Generation"""
        # 1. Dialog for Outcome
        dialog = QDialog(self)
        dialog.setWindowTitle("Conclude Personal Hearing")
        layout = QVBoxLayout(dialog)
        
        attended_cb = QCheckBox("Taxpayer Attended Hearing?")
        layout.addWidget(attended_cb)
        
        remarks_edit = QTextEdit()
        remarks_edit.setPlaceholderText("Enter outcome remarks...")
        layout.addWidget(remarks_edit)
        
        btn_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        btn_box.accepted.connect(dialog.accept)
        btn_box.rejected.connect(dialog.reject)
        layout.addWidget(btn_box)
        
        if dialog.exec() == QDialog.Accepted:
            try:
                # 2. Persist Outcome
                outcome_data = {
                    'attended': attended_cb.isChecked(),
                    'remarks': remarks_edit.toPlainText(),
                    'concluded_at': datetime.datetime.now().isoformat()
                }
                
                # Update additional details
                details = self.proceeding_data.get('additional_details', {})
                if isinstance(details, str): import json; details = json.loads(details)
                
                details['ph_outcome'] = outcome_data
                self.proceeding_data['additional_details'] = details
                
                # DB Update for details (Separate from transition to be safe, or just rely on transition if it handled data?)
                # Transition uses update_proceeding which accepts partials.
                # But here we want to update 'additional_details' column specifically.
                import json
                self.db.update_proceeding(self.proceeding_id, {'additional_details': json.dumps(details)})
                
                # 3. Transition to PH_COMPLETED
                self.transition_to(WorkflowStage.PH_COMPLETED)
                
                QMessageBox.information(self, "Success", "Hearing Concluded. Order Generation Unlocked.")
                
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to conclude hearing: {e}")
